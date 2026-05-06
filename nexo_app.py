
import os
import re
import json
import math
import calendar
import sqlite3
import hashlib
from io import BytesIO
from datetime import date, datetime, timedelta

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================================
# NEXO — STARTUP UI + PLANOS + ADMIN + LOGO + WORD/PDF
# ==========================================================
# Instalar:
# pip install streamlit pandas matplotlib python-docx openpyxl reportlab requests
#
# Rodar:
# streamlit run nexo_app_premium.py
#
# Mercado Pago:
# 1) crie uma variável de ambiente MERCADO_PAGO_ACCESS_TOKEN
# 2) opcional: defina APP_PUBLIC_URL com a URL pública do app
# 3) a integração cria uma preferência de pagamento e libera o acesso
#    após confirmação manual/webhook. No Streamlit local, use "Confirmar
#    pagamento manualmente" no admin ou implemente webhook em servidor externo.
# ==========================================================

st.set_page_config(
    page_title="NEXO",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed",
)

APP_NAME = "NEXO"
APP_SLOGAN = "Plataforma clínica inteligente"
DB_PATH = "nexo.db"
UPLOAD_DIR = "uploads_nexo"
os.makedirs(UPLOAD_DIR, exist_ok=True)

ADMIN_EMAIL = "admin@nexo.app"
ADMIN_SENHA = "NexoMaster2026!"

PROFISSOES = [
    "Musicoterapeuta", "Fonoaudiólogo", "Terapeuta Ocupacional", "Psicólogo",
    "Pet Terapeuta", "Educador Físico", "Psicomotricista", "Psicopedagogo",
    "Aplicador ABA", "Outro",
]


# ==========================================================
# MOTOR MULTIÁREA — PERFIS PROFISSIONAIS
# ==========================================================
PERFIS_PROFISSIONAIS = {
    "Musicoterapeuta": {
        "rotulo_relatorio": "Relatório Musicoterapêutico",
        "foco": "experiência sonoro-musical, comunicação musical, responsividade, vínculo terapêutico e organização expressiva",
        "areas": ["comunicação musical", "interação sonoro-musical", "regulação emocional", "atenção compartilhada", "expressão vocal/corporal", "organização rítmica"],
        "termos": ["improvisação clínica", "recriação musical", "escuta terapêutica", "expressividade musical", "reciprocidade sonora"],
        "avaliacoes": ["Nordoff-Robbins", "IAPS", "DEMUCA", "MEL", "Histórico sonoro-musical"],
        "prescricao": "intervenções sonoro-musicais estruturadas e improvisacionais, com foco em comunicação, regulação, vínculo, expressividade e ampliação de repertórios funcionais.",
        "objetivos_base": ["Ampliar iniciativa e responsividade sonoro-musical", "Favorecer comunicação expressiva mediada pela música", "Fortalecer regulação emocional por meio de experiências musicais organizadoras"]
    },
    "Fonoaudiólogo": {
        "rotulo_relatorio": "Relatório Fonoaudiológico",
        "foco": "linguagem receptiva e expressiva, comunicação funcional, pragmática, articulação, processamento auditivo e interação comunicativa",
        "areas": ["linguagem receptiva", "linguagem expressiva", "pragmática", "articulação", "processamento auditivo", "comunicação funcional"],
        "termos": ["intencionalidade comunicativa", "troca comunicativa", "compreensão verbal", "produção oral", "funções comunicativas"],
        "avaliacoes": ["Triagem de linguagem", "Protocolo pragmático", "Avaliação miofuncional", "Observação comunicativa", "Avaliação auditiva funcional"],
        "prescricao": "intervenções fonoaudiológicas com foco em comunicação funcional, compreensão, expressão, pragmática, articulação e generalização dos repertórios comunicativos.",
        "objetivos_base": ["Ampliar comunicação funcional", "Favorecer compreensão e expressão verbal/não verbal", "Desenvolver uso pragmático da linguagem em contextos naturais"]
    },
    "Terapeuta Ocupacional": {
        "rotulo_relatorio": "Relatório de Terapia Ocupacional",
        "foco": "desempenho ocupacional, autonomia, integração sensorial, coordenação motora, participação e atividades de vida diária",
        "areas": ["AVDs", "coordenação motora", "integração sensorial", "autonomia", "brincar funcional", "participação ocupacional"],
        "termos": ["desempenho ocupacional", "funcionalidade", "modulação sensorial", "planejamento motor", "independência nas rotinas"],
        "avaliacoes": ["Perfil sensorial", "Avaliação funcional", "AVDs", "Observação ocupacional", "Coordenação motora"],
        "prescricao": "intervenções ocupacionais voltadas à funcionalidade, participação, autonomia, modulação sensorial, coordenação motora e desempenho nas rotinas diárias.",
        "objetivos_base": ["Ampliar autonomia em atividades de vida diária", "Favorecer modulação sensorial e organização corporal", "Desenvolver participação funcional nas rotinas"]
    },
    "Psicólogo": {
        "rotulo_relatorio": "Relatório Psicológico",
        "foco": "aspectos emocionais, comportamentais, cognitivos, vinculares, regulação emocional e repertório adaptativo",
        "areas": ["regulação emocional", "comportamento", "cognição", "vínculo", "habilidades sociais", "repertório adaptativo"],
        "termos": ["autorregulação", "expressão emocional", "flexibilidade cognitiva", "repertório comportamental", "manejo emocional"],
        "avaliacoes": ["Entrevista clínica", "Observação comportamental", "Escalas emocionais", "Inventários psicológicos", "Anamnese"],
        "prescricao": "intervenções psicológicas orientadas à regulação emocional, ampliação de repertório adaptativo, manejo comportamental e desenvolvimento de recursos cognitivos e relacionais.",
        "objetivos_base": ["Favorecer reconhecimento e expressão emocional", "Ampliar repertórios adaptativos", "Desenvolver estratégias de autorregulação"]
    },
    "Aplicador ABA": {
        "rotulo_relatorio": "Relatório de Intervenção ABA",
        "foco": "comportamentos-alvo, repertórios de aprendizagem, antecedentes, respostas, consequências, generalização e redução de barreiras comportamentais",
        "areas": ["comportamentos-alvo", "habilidades de aprendizagem", "comunicação funcional", "generalização", "autonomia", "redução de comportamentos interferentes"],
        "termos": ["ABC", "reforçamento", "modelagem", "prompt", "generalização", "comportamento-alvo"],
        "avaliacoes": ["VB-MAPP", "ABLLS-R", "AFLS", "Registro ABC", "Análise funcional"],
        "prescricao": "programas comportamentais com definição de comportamentos-alvo, critérios mensuráveis, estratégias de ensino, reforçamento, generalização e monitoramento de dados.",
        "objetivos_base": ["Aumentar repertórios funcionais mensuráveis", "Reduzir comportamentos interferentes por estratégias baseadas em função", "Promover generalização das habilidades aprendidas"]
    },
    "Psicomotricista": {
        "rotulo_relatorio": "Relatório Psicomotor",
        "foco": "organização corporal, tônus, lateralidade, equilíbrio, coordenação, planejamento motor e relação corpo-espaço",
        "areas": ["coordenação global", "motricidade fina", "equilíbrio", "lateralidade", "esquema corporal", "planejamento motor"],
        "termos": ["esquema corporal", "organização tônico-postural", "coordenação visomotora", "orientação espacial", "ritmo corporal"],
        "avaliacoes": ["Avaliação psicomotora", "Coordenação motora", "Esquema corporal", "Lateralidade", "Equilíbrio"],
        "prescricao": "intervenções psicomotoras voltadas à organização corporal, coordenação, equilíbrio, lateralidade, planejamento motor e integração corpo-espaço.",
        "objetivos_base": ["Ampliar organização corporal", "Desenvolver coordenação e planejamento motor", "Favorecer integração espaço-temporal"]
    },
    "Psicopedagogo": {
        "rotulo_relatorio": "Relatório Psicopedagógico",
        "foco": "processos de aprendizagem, atenção, memória, funções executivas, leitura, escrita, raciocínio e estratégias cognitivas",
        "areas": ["atenção", "memória", "funções executivas", "leitura", "escrita", "raciocínio lógico"],
        "termos": ["mediação da aprendizagem", "estratégias cognitivas", "processamento acadêmico", "organização de estudo", "habilidades escolares"],
        "avaliacoes": ["Sondagem pedagógica", "Avaliação de leitura/escrita", "Atenção", "Memória", "Funções executivas"],
        "prescricao": "intervenções psicopedagógicas orientadas à aprendizagem, organização cognitiva, leitura, escrita, atenção, memória e estratégias de estudo.",
        "objetivos_base": ["Ampliar estratégias de aprendizagem", "Fortalecer atenção e organização cognitiva", "Desenvolver habilidades acadêmicas funcionais"]
    },
    "Educador Físico": {
        "rotulo_relatorio": "Relatório de Desenvolvimento Motor",
        "foco": "condicionamento, coordenação motora, equilíbrio, força, mobilidade, consciência corporal e participação em atividades físicas",
        "areas": ["coordenação", "equilíbrio", "força", "mobilidade", "resistência", "consciência corporal"],
        "termos": ["capacidades físicas", "controle motor", "habilidades motoras", "progressão de carga", "participação corporal"],
        "avaliacoes": ["Avaliação motora", "Capacidades físicas", "Coordenação", "Mobilidade", "Resistência"],
        "prescricao": "programas de atividade física adaptada com foco em coordenação, equilíbrio, força, mobilidade, consciência corporal e participação segura.",
        "objetivos_base": ["Melhorar coordenação e controle motor", "Ampliar participação em atividades físicas", "Desenvolver força, equilíbrio e mobilidade"]
    },
    "Pet Terapeuta": {
        "rotulo_relatorio": "Relatório de Intervenção Assistida por Animais",
        "foco": "vínculo, engajamento, comunicação, regulação emocional, interação social e participação mediada pelo animal",
        "areas": ["vínculo", "engajamento", "interação social", "regulação emocional", "comunicação", "participação"],
        "termos": ["intervenção assistida por animais", "mediação animal", "vínculo terapêutico", "engajamento socioemocional", "participação mediada"],
        "avaliacoes": ["Observação de engajamento", "Interação mediada", "Regulação emocional", "Participação", "Comunicação funcional"],
        "prescricao": "intervenções assistidas por animais com foco em vínculo, engajamento, comunicação, regulação emocional, segurança e participação social.",
        "objetivos_base": ["Ampliar engajamento mediado pelo animal", "Favorecer regulação emocional", "Desenvolver interação social em contexto terapêutico"]
    },
    "Outro": {
        "rotulo_relatorio": "Relatório Clínico-Funcional",
        "foco": "desenvolvimento funcional, participação, comunicação, autonomia, regulação e acompanhamento evolutivo",
        "areas": ["funcionalidade", "comunicação", "autonomia", "participação", "regulação", "desenvolvimento"],
        "termos": ["planejamento individualizado", "evolução funcional", "objetivos mensuráveis", "participação", "monitoramento clínico"],
        "avaliacoes": ["Avaliação clínica", "Observação funcional", "Anamnese", "Plano terapêutico", "Registro evolutivo"],
        "prescricao": "intervenções individualizadas com objetivos mensuráveis, acompanhamento evolutivo e foco no desenvolvimento funcional.",
        "objetivos_base": ["Ampliar funcionalidade", "Acompanhar evolução por indicadores mensuráveis", "Favorecer participação e autonomia"]
    }
}

def perfil_profissional(profissao: str):
    return PERFIS_PROFISSIONAIS.get(profissao, PERFIS_PROFISSIONAIS["Outro"])

def sugestoes_avaliacoes_por_profissao(profissao: str):
    base = ["Nordoff-Robbins", "IAPS", "DEMUCA"]
    perfil = perfil_profissional(profissao)
    # Mantém as escalas já implementadas e mostra sugestões coerentes da área como referência clínica.
    return base, perfil.get("avaliacoes", [])

def texto_justificativa_avaliacao(avaliacao: str, profissao: str):
    perfil = perfil_profissional(profissao)
    return (
        f"A avaliação {avaliacao} foi considerada relevante para a área de {profissao}, pois auxilia na organização de dados clínicos relacionados a "
        f"{perfil['foco']}. A análise estruturada permite transformar observações em indicadores funcionais, justificar objetivos terapêuticos "
        f"e acompanhar a evolução do paciente com maior clareza técnica."
    )

PLANOS = {
    "Starter": {
        "valor": 59.00,
        "valor_anual_mensal": 49.00,
        "pacientes": 10,
        "relatorios_mes": 20,
        "uploads_mes": 30,
        "profissionais": 1,
        "modo_clinica": False,
        "descricao": "Para profissionais iniciando a organização clínica com até 10 pacientes ativos.",
        "destaques": ["10 pacientes ativos", "Relatórios automáticos", "Agenda básica", "Gráficos automáticos"]
    },
    "Professional": {
        "valor": 149.00,
        "valor_anual_mensal": 129.00,
        "pacientes": 30,
        "relatorios_mes": 80,
        "uploads_mes": 150,
        "profissionais": 1,
        "modo_clinica": False,
        "descricao": "Para profissionais com rotina clínica consolidada e até 30 pacientes ativos.",
        "destaques": ["30 pacientes ativos", "Upload inteligente avançado", "Objetivos terapêuticos", "Evolução comparativa"]
    },
    "Clinic": {
        "valor": 449.00,
        "valor_anual_mensal": 399.00,
        "pacientes": 60,
        "relatorios_mes": 200,
        "uploads_mes": 400,
        "profissionais": 5,
        "modo_clinica": True,
        "descricao": "Para clínicas e equipes pequenas com gestão de até 60 pacientes ativos.",
        "destaques": ["60 pacientes ativos", "Modo Clínica", "Gestão de equipe", "Agenda integrada"]
    },
    "Enterprise": {
        "valor": 899.00,
        "valor_anual_mensal": 799.00,
        "pacientes": 999999,
        "relatorios_mes": 999999,
        "uploads_mes": 999999,
        "profissionais": 999999,
        "modo_clinica": True,
        "descricao": "Para redes, instituições e uso intensivo com limites operacionais ampliados.",
        "destaques": ["Pacientes ilimitados", "Profissionais ilimitados", "Suporte prioritário", "Customização institucional"]
    },
}

PLANOS_LEGADOS = {
    "Básico": "Starter",
    "Basico": "Starter",
    "Médio Básico": "Professional",
    "Medio Basico": "Professional",
    "Médio": "Professional",
    "Medio": "Professional",
    "Premium": "Clinic",
    "Empresarial Ilimitado": "Enterprise",
    "Empresarial": "Enterprise",
}

def normalizar_plano_nome(plano):
    plano = str(plano or "Starter").strip()
    return PLANOS_LEGADOS.get(plano, plano if plano in PLANOS else "Starter")

def plano_config(plano):
    return PLANOS.get(normalizar_plano_nome(plano), PLANOS["Starter"])


TIPOS_ACESSO = ["admin", "profissional_pagante", "supervisionando", "aluno_mentoria", "cortesia", "teste_gratuito"]
STATUS_ACESSO = ["ativo", "bloqueado", "expirado", "aguardando_pagamento"]

# ==========================================================
# SVG ICONS — estilo próximo ao layout enviado
# ==========================================================
def svg_icon(name: str, size: int = 22, color: str = "#9fb7d8") -> str:
    icons = {
        "brain": '<path d="M9 3a3 3 0 0 0-3 3v.2A3 3 0 0 0 4 9v1a3 3 0 0 0 2 2.82V15a3 3 0 0 0 3 3h1V3H9Z"/><path d="M15 3a3 3 0 0 1 3 3v.2A3 3 0 0 1 20 9v1a3 3 0 0 1-2 2.82V15a3 3 0 0 1-3 3h-1V3h1Z"/><path d="M10 7H7m3 4H6m8-4h3m-3 4h4"/>',
        "home": '<path d="m3 11 9-8 9 8"/><path d="M5 10v10h14V10"/><path d="M9 20v-6h6v6"/>',
        "user": '<circle cx="12" cy="8" r="4"/><path d="M4 21a8 8 0 0 1 16 0"/>',
        "chart": '<path d="M4 19h16"/><path d="M7 16V9"/><path d="M12 16V5"/><path d="M17 16v-3"/><path d="M7 9l5-4 5 8"/>',
        "file": '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8Z"/><path d="M14 2v6h6"/><path d="M8 13h8M8 17h6"/>',
        "upload": '<path d="M12 16V4"/><path d="m7 9 5-5 5 5"/><path d="M20 16v4H4v-4"/>',
        "spark": '<path d="M12 2v5M12 17v5M4.22 4.22l3.54 3.54M16.24 16.24l3.54 3.54M2 12h5M17 12h5M4.22 19.78l3.54-3.54M16.24 7.76l3.54-3.54"/>',
        "credit": '<rect x="3" y="5" width="18" height="14" rx="2"/><path d="M3 10h18"/><path d="M7 15h4"/>',
        "users": '<path d="M16 21v-2a4 4 0 0 0-4-4H6a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M22 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/>',
        "settings": '<path d="M12 15a3 3 0 1 0 0-6 3 3 0 0 0 0 6Z"/><path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06A1.65 1.65 0 0 0 15 19.4a1.65 1.65 0 0 0-1 .6 1.65 1.65 0 0 0-.4 1.05V21a2 2 0 1 1-4 0v-.09A1.65 1.65 0 0 0 8.6 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06A1.65 1.65 0 0 0 4.6 15a1.65 1.65 0 0 0-.6-1H4a2 2 0 1 1 0-4h.09A1.65 1.65 0 0 0 4.6 8a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06A1.65 1.65 0 0 0 9 4.6a1.65 1.65 0 0 0 1-.6V4a2 2 0 1 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06A1.65 1.65 0 0 0 19.4 9c.2.36.4.68.6 1H20a2 2 0 1 1 0 4h-.09a1.65 1.65 0 0 0-.51 1Z"/>',
        "plus": '<path d="M12 5v14M5 12h14"/>',
        "crown": '<path d="m2 6 5 5 5-9 5 9 5-5-3 13H5L2 6Z"/>',
        "bell": '<path d="M18 8a6 6 0 1 0-12 0c0 7-3 7-3 9h18c0-2-3-2-3-9"/><path d="M13.73 21a2 2 0 0 1-3.46 0"/>',
        "search": '<circle cx="11" cy="11" r="8"/><path d="m21 21-4.35-4.35"/>',
        "logout": '<path d="M9 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h4"/><path d="m16 17 5-5-5-5"/><path d="M21 12H9"/>',
        "help": '<circle cx="12" cy="12" r="10"/><path d="M9.1 9a3 3 0 1 1 5.8 1c0 2-3 2-3 4"/><path d="M12 17h.01"/>',
        "clipboard": '<path d="M9 2h6v4H9z"/><path d="M9 4H5a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h14a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2h-4"/>',
    }
    body = icons.get(name, icons["spark"])
    return f'<svg width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">{body}</svg>'

def brain_logo(size: int = 54) -> str:
    return f"""
    <div class="logo-wrap" style="width:{size}px;height:{size}px;">
        {svg_icon('brain', int(size*0.75), '#60a5fa')}
    </div>
    """

# ==========================================================
# DATABASE
# ==========================================================
def db():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def hash_senha(senha: str) -> str:
    return hashlib.sha256(senha.encode("utf-8")).hexdigest()

def colunas_tabela(tabela):
    conn = db()
    c = conn.cursor()
    c.execute(f"PRAGMA table_info({tabela})")
    cols = [r[1] for r in c.fetchall()]
    conn.close()
    return cols

def add_col_if_missing(tabela, coluna, sql):
    conn = db()
    c = conn.cursor()
    cols = [r[1] for r in c.execute(f"PRAGMA table_info({tabela})").fetchall()]
    if coluna not in cols:
        try:
            c.execute(sql)
            conn.commit()
        except sqlite3.OperationalError:
            pass
    conn.close()

def init_db():
    conn = db()
    c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT,
        email TEXT UNIQUE,
        senha_hash TEXT,
        profissao TEXT,
        registro TEXT,
        formacao TEXT,
        especializacoes TEXT,
        assinatura TEXT,
        plano TEXT DEFAULT '',
        status TEXT DEFAULT 'aguardando_pagamento',
        role TEXT DEFAULT 'user',
        tipo_acesso TEXT DEFAULT 'profissional_pagante',
        data_inicio TEXT,
        data_expiracao TEXT,
        limite_pacientes INTEGER DEFAULT 0,
        limite_relatorios_mes INTEGER DEFAULT 0,
        limite_uploads_mes INTEGER DEFAULT 0,
        observacao_admin TEXT,
        logo_path TEXT,
        logo_posicao TEXT DEFAULT 'Topo esquerdo',
        avaliacoes_padrao_autorizado INTEGER DEFAULT 0,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS pacientes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        nome TEXT,
        idade INTEGER,
        nascimento TEXT,
        diagnostico TEXT,
        escolaridade TEXT,
        responsaveis TEXT,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS evolucoes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        paciente_id INTEGER,
        data TEXT,
        titulo TEXT,
        descricao TEXT,
        humor TEXT,
        engajamento INTEGER,
        progresso INTEGER
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS uploads (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        paciente_id INTEGER,
        nome_arquivo TEXT,
        caminho TEXT,
        tipo TEXT,
        categoria TEXT,
        avaliacao_detectada TEXT,
        justificativa TEXT,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS avaliacao_modelos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        nome TEXT,
        profissao TEXT,
        avaliacao_detectada TEXT,
        descricao TEXT,
        campos_json TEXT,
        secoes_json TEXT,
        arquivo_origem TEXT,
        caminho_origem TEXT,
        grafico_json TEXT,
        fonte_referencia TEXT,
        ativo INTEGER DEFAULT 1,
        criado_em TEXT,
        atualizado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS avaliacao_respostas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        paciente_id INTEGER,
        modelo_id INTEGER,
        avaliacao_nome TEXT,
        respostas_json TEXT,
        interpretacao TEXT,
        justificativa TEXT,
        grafico_valores_json TEXT,
        usar_no_relatorio INTEGER DEFAULT 1,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS usage_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        tipo TEXT,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS pagamentos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        plano TEXT,
        valor REAL,
        status TEXT,
        mercado_pago_preference_id TEXT,
        mercado_pago_payment_id TEXT,
        assinatura_url TEXT,
        criado_em TEXT,
        atualizado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS agenda_config (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER UNIQUE,
        dias_semana TEXT DEFAULT '0,1,2,3,4',
        hora_inicio TEXT DEFAULT '08:00',
        hora_fim TEXT DEFAULT '18:00',
        almoco_inicio TEXT DEFAULT '12:00',
        almoco_fim TEXT DEFAULT '13:00',
        duracao_padrao INTEGER DEFAULT 50,
        intervalo_padrao INTEGER DEFAULT 10,
        atualizado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS agenda_eventos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        paciente_id INTEGER,
        data TEXT,
        hora_inicio TEXT,
        hora_fim TEXT,
        titulo TEXT,
        modalidade TEXT DEFAULT 'Presencial',
        status TEXT DEFAULT 'previsto',
        origem TEXT DEFAULT 'manual',
        observacao TEXT,
        criado_em TEXT,
        atualizado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS agenda_bloqueios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        data TEXT,
        hora_inicio TEXT,
        hora_fim TEXT,
        motivo TEXT,
        criado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS clinica_setores (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        gestor_id INTEGER,
        nome TEXT,
        area TEXT,
        descricao TEXT,
        ativo INTEGER DEFAULT 1,
        criado_em TEXT,
        atualizado_em TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS clinica_profissionais (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        gestor_id INTEGER,
        nome TEXT,
        email TEXT,
        profissao TEXT,
        registro TEXT,
        setor_id INTEGER,
        status TEXT DEFAULT 'ativo',
        observacoes TEXT,
        criado_em TEXT,
        atualizado_em TEXT
    )
    """)
    conn.commit()
    conn.close()

    migrations = {
        "users": {
            "logo_path": "ALTER TABLE users ADD COLUMN logo_path TEXT",
            "logo_posicao": "ALTER TABLE users ADD COLUMN logo_posicao TEXT DEFAULT 'Topo esquerdo'",
            "avaliacoes_padrao_autorizado": "ALTER TABLE users ADD COLUMN avaliacoes_padrao_autorizado INTEGER DEFAULT 0",
        },
        "uploads": {
            "caminho": "ALTER TABLE uploads ADD COLUMN caminho TEXT",
            "categoria": "ALTER TABLE uploads ADD COLUMN categoria TEXT",
        },
        "pacientes": {
            "dia_semana": "ALTER TABLE pacientes ADD COLUMN dia_semana INTEGER DEFAULT NULL",
            "hora_atendimento": "ALTER TABLE pacientes ADD COLUMN hora_atendimento TEXT DEFAULT ''",
            "duracao_sessao": "ALTER TABLE pacientes ADD COLUMN duracao_sessao INTEGER DEFAULT 50",
            "frequencia": "ALTER TABLE pacientes ADD COLUMN frequencia TEXT DEFAULT 'Semanal'",
            "modalidade": "ALTER TABLE pacientes ADD COLUMN modalidade TEXT DEFAULT 'Presencial'",
            "agenda_status": "ALTER TABLE pacientes ADD COLUMN agenda_status TEXT DEFAULT 'ativo'",
        },
        "evolucoes": {
            "event_id": "ALTER TABLE evolucoes ADD COLUMN event_id INTEGER DEFAULT NULL",
            "status_atendimento": "ALTER TABLE evolucoes ADD COLUMN status_atendimento TEXT DEFAULT 'realizado'",
        },
        "avaliacao_modelos": {
            "grafico_json": "ALTER TABLE avaliacao_modelos ADD COLUMN grafico_json TEXT",
            "fonte_referencia": "ALTER TABLE avaliacao_modelos ADD COLUMN fonte_referencia TEXT",
            "gestor_id": "ALTER TABLE avaliacao_modelos ADD COLUMN gestor_id INTEGER DEFAULT NULL",
            "setor_id": "ALTER TABLE avaliacao_modelos ADD COLUMN setor_id INTEGER DEFAULT NULL",
            "escopo": "ALTER TABLE avaliacao_modelos ADD COLUMN escopo TEXT DEFAULT 'perfil'",
            "area_alvo": "ALTER TABLE avaliacao_modelos ADD COLUMN area_alvo TEXT DEFAULT ''",
        },
        "avaliacao_respostas": {
            "grafico_valores_json": "ALTER TABLE avaliacao_respostas ADD COLUMN grafico_valores_json TEXT",
        },
    }
    for tabela, cols in migrations.items():
        for col, sql in cols.items():
            add_col_if_missing(tabela, col, sql)

    conn = db()
    c = conn.cursor()
    c.execute("SELECT id FROM users WHERE email=?", (ADMIN_EMAIL,))
    admin = c.fetchone()
    if admin:
        c.execute("""
            UPDATE users SET nome=?, senha_hash=?, profissao=?, registro=?, formacao=?, especializacoes=?, assinatura=?,
            plano=?, status=?, role=?, tipo_acesso=?, data_inicio=?, data_expiracao=?, limite_pacientes=?,
            limite_relatorios_mes=?, limite_uploads_mes=?, observacao_admin=?, avaliacoes_padrao_autorizado=1 WHERE email=?
        """, (
            "Administrador NEXO", hash_senha(ADMIN_SENHA), "Administrador", "ADMIN", "Gestão da plataforma NEXO",
            "Controle de usuários, planos e acessos", "Administrador NEXO", "Enterprise", "ativo", "admin", "admin",
            date.today().strftime("%Y-%m-%d"), "2099-12-31", 999999, 999999, 999999,
            "Administrador principal da plataforma.", ADMIN_EMAIL
        ))
    else:
        c.execute("""
            INSERT INTO users (nome, email, senha_hash, profissao, registro, formacao, especializacoes, assinatura, plano,
            status, role, tipo_acesso, data_inicio, data_expiracao, limite_pacientes, limite_relatorios_mes,
            limite_uploads_mes, observacao_admin, avaliacoes_padrao_autorizado, criado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            "Administrador NEXO", ADMIN_EMAIL, hash_senha(ADMIN_SENHA), "Administrador", "ADMIN", "Gestão da plataforma NEXO",
            "Controle de usuários, planos e acessos", "Administrador NEXO", "Enterprise", "ativo", "admin", "admin",
            date.today().strftime("%Y-%m-%d"), "2099-12-31", 999999, 999999, 999999,
            "Administrador principal da plataforma.", 1, datetime.now().isoformat(timespec="seconds")
        ))
    conn.commit()
    conn.close()

def row_to_user(row):
    if not row:
        return None
    keys = [x[1] for x in sqlite3.connect(DB_PATH).execute("PRAGMA table_info(users)").fetchall()]
    return {keys[i]: row[i] if i < len(row) else None for i in range(len(keys))}

def buscar_usuario(user_id):
    conn = db()
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE id=?", (user_id,))
    user = row_to_user(c.fetchone())
    conn.close()
    return user

def autenticar(email, senha):
    conn = db()
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE email=? AND senha_hash=?", (email.lower().strip(), hash_senha(senha)))
    user = row_to_user(c.fetchone())
    conn.close()
    return user

def criar_usuario(nome, email, senha, profissao, registro, formacao, especializacoes):
    conn = db()
    c = conn.cursor()
    try:
        c.execute("""
        INSERT INTO users (nome, email, senha_hash, profissao, registro, formacao, especializacoes, assinatura,
        plano, status, role, tipo_acesso, data_inicio, data_expiracao, limite_pacientes, limite_relatorios_mes,
        limite_uploads_mes, observacao_admin, criado_em)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            nome, email.lower().strip(), hash_senha(senha), profissao, registro, formacao, especializacoes, nome,
            "", "aguardando_pagamento", "user", "profissional_pagante", "", "", 0, 0, 0,
            "Conta criada pelo formulário público. Precisa escolher plano e pagar.", datetime.now().isoformat(timespec="seconds")
        ))
        conn.commit()
        return True, "Conta criada. Escolha um plano para liberar o modo de produção."
    except sqlite3.IntegrityError:
        return False, "Este e-mail já está cadastrado."
    finally:
        conn.close()

def atualizar_plano_usuario(user_id, plano, status="ativo"):
    plano = normalizar_plano_nome(plano)
    p = PLANOS[plano]
    conn = db()
    c = conn.cursor()
    inicio = date.today()
    exp = inicio + timedelta(days=30)
    c.execute("""
    UPDATE users SET plano=?, status=?, data_inicio=?, data_expiracao=?, limite_pacientes=?,
    limite_relatorios_mes=?, limite_uploads_mes=? WHERE id=?
    """, (
        plano, status, inicio.strftime("%Y-%m-%d"), exp.strftime("%Y-%m-%d"),
        p["pacientes"], p["relatorios_mes"], p["uploads_mes"], user_id
    ))
    conn.commit()
    conn.close()

def atualizar_perfil(user_id, profissao, registro, formacao, especializacoes, assinatura, logo_path=None, logo_posicao="Topo esquerdo"):
    conn = db()
    c = conn.cursor()
    if logo_path:
        c.execute("""UPDATE users SET profissao=?, registro=?, formacao=?, especializacoes=?, assinatura=?,
        logo_path=?, logo_posicao=? WHERE id=?""",
        (profissao, registro, formacao, especializacoes, assinatura, logo_path, logo_posicao, user_id))
    else:
        c.execute("""UPDATE users SET profissao=?, registro=?, formacao=?, especializacoes=?, assinatura=?,
        logo_posicao=? WHERE id=?""",
        (profissao, registro, formacao, especializacoes, assinatura, logo_posicao, user_id))
    conn.commit()
    conn.close()

def listar_pacientes(user_id):
    conn = db()
    c = conn.cursor()
    c.execute("SELECT id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis FROM pacientes WHERE user_id=? ORDER BY nome", (user_id,))
    rows = c.fetchall()
    conn.close()
    return rows

def contar_pacientes(user_id):
    conn = db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM pacientes WHERE user_id=?", (user_id,))
    n = c.fetchone()[0]
    conn.close()
    return n

def salvar_paciente(user_id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis,
                    dia_semana=None, hora_atendimento="", duracao_sessao=50, frequencia="Semanal", modalidade="Presencial"):
    conn = db()
    c = conn.cursor()
    c.execute("""INSERT INTO pacientes (user_id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis,
              dia_semana, hora_atendimento, duracao_sessao, frequencia, modalidade, agenda_status, criado_em)
              VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (user_id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis,
               dia_semana, hora_atendimento, duracao_sessao, frequencia, modalidade, "ativo",
               datetime.now().isoformat(timespec="seconds")))
    conn.commit()
    conn.close()

def salvar_evolucao(user_id, paciente_id, data_ev, titulo, descricao, humor, engajamento, progresso, event_id=None, status_atendimento="realizado"):
    conn = db()
    c = conn.cursor()
    c.execute("""INSERT INTO evolucoes (user_id, paciente_id, data, titulo, descricao, humor, engajamento, progresso, event_id, status_atendimento)
              VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (user_id, paciente_id, data_ev, titulo, descricao, humor, engajamento, progresso, event_id, status_atendimento))
    if event_id:
        c.execute("UPDATE agenda_eventos SET status=?, atualizado_em=? WHERE id=? AND user_id=?",
                  ("evolucao_ok", datetime.now().isoformat(timespec="seconds"), event_id, user_id))
    conn.commit()
    conn.close()

def listar_evolucoes(user_id, paciente_id):
    conn = db()
    c = conn.cursor()
    c.execute("""SELECT data, titulo, descricao, humor, engajamento, progresso FROM evolucoes
              WHERE user_id=? AND paciente_id=? ORDER BY id DESC""", (user_id, paciente_id))
    rows = c.fetchall()
    conn.close()
    return rows

def registrar_upload(user_id, paciente_id, nome_arquivo, caminho, tipo, categoria, avaliacao_detectada, justificativa):
    conn = db()
    c = conn.cursor()
    c.execute("""INSERT INTO uploads (user_id, paciente_id, nome_arquivo, caminho, tipo, categoria,
              avaliacao_detectada, justificativa, criado_em) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (user_id, paciente_id, nome_arquivo, caminho, tipo, categoria, avaliacao_detectada, justificativa,
               datetime.now().isoformat(timespec="seconds")))
    conn.commit()
    conn.close()

def listar_uploads(user_id):
    conn = db()
    c = conn.cursor()
    c.execute("""SELECT nome_arquivo, categoria, avaliacao_detectada, criado_em FROM uploads
              WHERE user_id=? ORDER BY id DESC LIMIT 50""", (user_id,))
    rows = c.fetchall()
    conn.close()
    return rows

def texto_arquivo_upload(arquivo):
    """Extrai texto quando possível, mas nunca impede o upload.

    Suporta TXT/MD/CSV, DOCX, XLS/XLSX e PDF quando PyPDF2 estiver instalado.
    Imagens e arquivos proprietários continuam sendo aceitos e salvos; nesses casos
    o sistema cria o modelo a partir do nome do arquivo, da área profissional e da
    validação manual do usuário.
    """
    nome = (getattr(arquivo, "name", "") or "").lower()
    try:
        data = arquivo.getvalue()
        try:
            arquivo.seek(0)
        except Exception:
            pass
    except Exception:
        return ""

    if not data:
        return ""

    try:
        if nome.endswith((".txt", ".md", ".csv", ".json")):
            return data.decode("utf-8", errors="ignore")[:20000]

        if nome.endswith(".docx"):
            doc = Document(BytesIO(data))
            partes = []
            partes.extend([p.text for p in doc.paragraphs if p.text and p.text.strip()])
            for table in doc.tables[:8]:
                for row in table.rows[:80]:
                    partes.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            return "\n".join(partes)[:20000]

        if nome.endswith((".xlsx", ".xls")):
            xls = pd.ExcelFile(BytesIO(data))
            partes = []
            for sheet in xls.sheet_names[:6]:
                try:
                    df = pd.read_excel(xls, sheet_name=sheet, nrows=80)
                    partes.append(f"Planilha: {sheet}")
                    partes.append("Colunas: " + "; ".join([str(c) for c in df.columns if str(c) != "nan"]))
                    partes.append(df.head(25).astype(str).to_string(index=False))
                except Exception as e:
                    partes.append(f"Planilha: {sheet} não pôde ser lida ({e})")
            return "\n".join(partes)[:20000]

        if nome.endswith(".pdf"):
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(data))
                partes = []
                for page in reader.pages[:8]:
                    partes.append(page.extract_text() or "")
                return "\n".join(partes)[:20000]
            except Exception:
                return ""
    except Exception:
        return ""

    return ""


def normalizar_nome_campo(txt):
    txt = re.sub(r"\s+", " ", str(txt or "")).strip(" :-–—\t")
    return txt[:140]

def slug_streamlit_key(txt, limite=80):
    """Gera uma chave segura e curta para widgets do Streamlit."""
    txt = str(txt or "campo")
    txt = re.sub(r"[^0-9a-zA-Z_]+", "_", txt)
    txt = re.sub(r"_+", "_", txt).strip("_")
    return (txt or "campo")[:limite]

def campos_unicos(campos, limite=120):
    """Remove duplicidades sem perder a ordem e evita campos vazios."""
    saida, vistos = [], set()
    for campo in campos or []:
        nome = normalizar_nome_campo(campo)[:limite]
        if not nome:
            continue
        chave = re.sub(r"\W+", "", nome.lower())
        if chave in vistos:
            continue
        vistos.add(chave)
        saida.append(nome)
    return saida


def aviso_direitos_autorais_modelo(nome_avaliacao):
    return (
        f"Modelo clínico automatizado derivado de referência enviada pelo profissional para {nome_avaliacao}. "
        "O NEXO não reproduz integralmente itens, escalas, manuais, tabelas normativas ou conteúdo protegido do instrumento original. "
        "A estrutura preserva a coerência clínica geral do arquivo enviado por meio de rótulos resumidos, domínios e critérios editáveis. "
        "Revise e valide antes do uso clínico."
    )

def nome_arquivo_seguro(nome):
    base = os.path.basename(str(nome or "arquivo"))
    base = re.sub(r"[^0-9a-zA-ZÀ-ÿ._ -]+", "_", base).strip()
    return base or "arquivo"

def dataframe_modelos_seguro(modelos):
    """Converte modelos do banco para DataFrame sem quebrar quando a query retornar mais/menos colunas."""
    colunas = [
        "ID", "Nome", "Área", "Avaliação", "Descrição", "Criado em",
        "Escopo", "Setor ID", "Gestor ID", "Área alvo"
    ]
    if not modelos:
        return pd.DataFrame(columns=["ID", "Nome", "Área", "Avaliação", "Descrição", "Criado em"])
    largura = max(len(tuple(r)) for r in modelos)
    cols = colunas[:largura] + [f"Extra {i}" for i in range(len(colunas), largura)]
    df = pd.DataFrame([tuple(r) for r in modelos], columns=cols)
    visiveis = ["ID", "Nome", "Área", "Avaliação", "Descrição", "Criado em"]
    return df[[c for c in visiveis if c in df.columns]]

def extensao_arquivo(nome):
    return os.path.splitext(str(nome or ""))[1].lower().replace(".", "")

# ==========================================================
# MOTOR DE LEITURA FIEL/DERIVADA DE ESCALAS
# ==========================================================
PALAVRAS_DOMINIO = re.compile(
    r"(dom[ií]nio|subescala|área|area|categoria|fator|dimens[aã]o|habilidade|compet[eê]ncia|se[cç][aã]o|parte)",
    re.I,
)
PALAVRAS_ITEM = re.compile(
    r"(pontua|pontuação|score|escore|nota|marque|assinale|classifique|freq[uü]ência|frequencia|nunca|raramente|às vezes|as vezes|sempre|observa|realiza|responde|inicia|mant[eé]m|imita|compreende|identifica|tolera|evita|necessita)",
    re.I,
)


def limpar_linha_escala(linha):
    linha = str(linha or "")
    linha = linha.replace("\x00", " ")
    linha = re.sub(r"\s+", " ", linha).strip(" •▪▫■□●○-–—\t")
    linha = re.sub(r"^(item\s*)?\d+[\).:-]?\s*", "", linha, flags=re.I)
    linha = re.sub(r"^[a-zA-Z][\).:-]\s*", "", linha)
    linha = re.sub(r"\b(0|1|2|3|4|5)\s*[=:-]\s*", "", linha)
    linha = re.sub(r"\s*\|\s*", " — ", linha)
    return linha.strip()


def parece_titulo_dominio(linha):
    l = limpar_linha_escala(linha)
    if not l or len(l) > 90:
        return False
    if PALAVRAS_DOMINIO.search(l):
        return True
    # Título curto em caixa alta/título, sem verbo instrucional.
    palavras = l.split()
    if 1 <= len(palavras) <= 7 and not PALAVRAS_ITEM.search(l):
        letras = re.sub(r"[^A-Za-zÀ-ÿ]", "", l)
        if letras and (l.isupper() or sum(1 for p in palavras if p[:1].isupper()) >= max(1, len(palavras)-1)):
            return True
    return False


def parece_item_escala(linha):
    original = str(linha or "").strip()
    l = limpar_linha_escala(original)
    if not l or len(l) < 6 or len(l) > 260:
        return False
    if re.match(r"^(item\s*)?\d+[\).:-]", original, flags=re.I):
        return True
    if re.match(r"^[•\-*]\s+", original):
        return True
    if PALAVRAS_ITEM.search(l):
        return True
    if "?" in l:
        return True
    return False


def resumir_enunciado_seguro(linha, max_palavras=16):
    """Mantém coerência do enunciado sem copiar bloco longo do instrumento."""
    l = limpar_linha_escala(linha)
    l = re.sub(r"\b(nunca|raramente|às vezes|as vezes|frequentemente|sempre)\b", "", l, flags=re.I)
    l = re.sub(r"\b(0|1|2|3|4|5)\b", "", l)
    l = re.sub(r"\s+", " ", l).strip(" -–—:;,.|")
    palavras = l.split()
    if len(palavras) > max_palavras:
        l = " ".join(palavras[:max_palavras]).rstrip(".,;:") + "…"
    if not l:
        return "Indicador observado"
    # Evita rótulo idêntico ao texto integral quando ele for longo.
    return l[:135]


def detectar_intervalo_pontuacao(texto):
    t = str(texto or "")
    nums = []
    for a, b in re.findall(r"\b([+-]?\d+)\s*(?:a|até|-)\s*([+-]?\d+)\b", t, flags=re.I):
        try:
            nums.extend([int(a), int(b)])
        except Exception:
            pass
    for n in re.findall(r"\b(?:nota|score|escore|pontuaç[aã]o)\s*(?:de)?\s*([+-]?\d+)\b", t, flags=re.I):
        try:
            nums.append(int(n))
        except Exception:
            pass
    if nums:
        return min(nums), max(nums)
    # Likert típico sem numeração explícita
    if re.search(r"nunca.*sempre|raramente.*sempre|0.*4|0.*5", t, re.I | re.S):
        return 0, 4
    return 0, 10


def extrair_estrutura_escala(texto, avaliacao, profissao):
    """Extrai domínios e itens com fidelidade mínima aos enunciados.

    Estratégia:
    - reconhece cabeçalhos/domínios;
    - reconhece itens numerados, bullets, perguntas e frases de pontuação;
    - cria rótulos resumidos/editáveis, evitando cópia extensa;
    - preserva a ordem do arquivo.
    """
    linhas_brutas = [l for l in str(texto or "").splitlines()]
    linhas = [normalizar_nome_campo(l) for l in linhas_brutas if normalizar_nome_campo(l)]
    if not linhas:
        return [], []

    secoes = []
    dominio_atual = "Itens observáveis da escala"
    itens_atual = []
    dominios_detectados = []
    contador = 1

    def fechar_secao():
        nonlocal itens_atual, dominio_atual
        if itens_atual:
            secoes.append({"titulo": dominio_atual[:80], "campos": campos_unicos(itens_atual, limite=135)})
            dominios_detectados.append(dominio_atual[:70])
            itens_atual = []

    for linha in linhas[:450]:
        limpa = limpar_linha_escala(linha)
        if not limpa:
            continue
        if parece_titulo_dominio(limpa):
            fechar_secao()
            dominio_atual = limpa[:80]
            continue
        if parece_item_escala(limpa):
            resumo = resumir_enunciado_seguro(limpa)
            itens_atual.append(f"Item {contador:02d} — {resumo}")
            contador += 1

    fechar_secao()

    # Caso o arquivo seja planilha/tabular com colunas importantes, usa colunas como itens coerentes.
    if not secoes:
        candidatos = []
        for linha in linhas[:120]:
            partes = [limpar_linha_escala(p) for p in re.split(r"\s{2,}|;|\t|\|", linha) if limpar_linha_escala(p)]
            for p in partes:
                if 3 <= len(p) <= 80 and re.search(r"(comunica|aten|intera|regula|motor|sensor|linguagem|autonomia|score|pontua|resultado|observa|item|domínio|dominio)", p, re.I):
                    candidatos.append(f"Indicador — {resumir_enunciado_seguro(p, 10)}")
        if candidatos:
            secoes = [{"titulo": "Indicadores extraídos da planilha", "campos": campos_unicos(candidatos, limite=120)[:30]}]
            dominios_detectados = ["Indicadores extraídos da planilha"]

    return secoes[:12], campos_unicos(dominios_detectados, limite=80)[:12]


def campos_fallback_por_profissao(profissao):
    perfil = perfil_profissional(profissao)
    areas = perfil.get("areas", []) or []
    return [f"Indicador de {a}" for a in areas[:8]] + [
        "Resposta funcional observada",
        "Nível de suporte necessário",
        "Pontuação/resultado",
        "Síntese clínica",
        "Conduta sugerida",
    ]


def inferir_campos_por_texto(texto, avaliacao, profissao):
    secoes, _dominios = extrair_estrutura_escala(texto, avaliacao, profissao)
    campos = []
    for sec in secoes:
        campos.extend(sec.get("campos", []))
    if campos:
        return campos_unicos(campos, limite=135)[:36]

    perfil = perfil_profissional(profissao)
    candidatos = []
    for linha in (texto or "").splitlines():
        l = normalizar_nome_campo(linha)
        if not l or len(l) < 3 or len(l) > 140:
            continue
        if re.search(r"(nome|idade|data|diagn[oó]stico|queixa|observa|objetivo|dom[ií]nio|item|pontua|score|resultado|conduta|aten[cç][aã]o|comunica|intera|regula|motric|linguagem|autonomia|sensorial|comport)", l, re.I):
            candidatos.append(f"Indicador — {resumir_enunciado_seguro(l)}")
    for termo in perfil.get("areas", []) + perfil.get("termos", []) + perfil.get("avaliacoes", []):
        candidatos.append(str(termo).capitalize())
    if len(candidatos) < 5:
        candidatos += campos_fallback_por_profissao(profissao)
    return campos_unicos(candidatos, limite=135)[:36]


def montar_modelo_automatizado(nome_arquivo, texto, profissao):
    avaliacao = reconhecer_avaliacao_por_nome(nome_arquivo, profissao)
    secoes_extraidas, dominios = extrair_estrutura_escala(texto, avaliacao, profissao)

    if secoes_extraidas:
        secoes_resultados = secoes_extraidas
    else:
        campos = inferir_campos_por_texto(texto, avaliacao, profissao)
        secoes_resultados = [{"titulo": "Resultados e indicadores", "campos": campos}]

    secoes = [
        {"titulo":"Identificação da avaliação", "campos":["Data da avaliação", "Contexto de aplicação", "Responsável pelo preenchimento"]},
        *secoes_resultados,
        {"titulo":"Síntese clínica", "campos":["Interpretação clínica", "Impacto funcional", "Prioridades terapêuticas", "Conduta sugerida"]},
    ]
    return avaliacao, secoes


# ==========================================================
# MOTOR DE GRÁFICOS FIÉIS POR AVALIAÇÃO
# ==========================================================
AVALIACOES_REFERENCIA = {
    "Nordoff-Robbins": {
        "tipo": "barras_percentual",
        "fonte": "Modelo clínico interno: comunicabilidade musical / responsividade / engajamento / autonomia musical. Revise conforme o protocolo usado pelo profissional.",
        "dominios": [
            {"dominio":"Comunicabilidade musical", "maximo":10},
            {"dominio":"Responsividade", "maximo":10},
            {"dominio":"Iniciativa", "maximo":10},
            {"dominio":"Sustentação da interação", "maximo":10},
            {"dominio":"Autonomia musical", "maximo":10},
        ],
    },
    "IAPS": {
        "tipo": "radar_percentual",
        "fonte": "Modelo clínico interno: domínios de improvisação, recriação, composição e escuta musical. Revise conforme a versão institucional usada.",
        "dominios": [
            {"dominio":"Improvisação", "maximo":20},
            {"dominio":"Recriação", "maximo":20},
            {"dominio":"Composição", "maximo":20},
            {"dominio":"Escuta Musical", "maximo":20},
        ],
    },
    "DEMUCA": {
        "tipo": "barras_percentual",
        "fonte": "Modelo interno baseado nos domínios já usados no app. Revise pesos/itens se a versão do instrumento for diferente.",
        "dominios": [
            {"dominio":"Comportamentos Restritivos", "maximo":14},
            {"dominio":"Interação Social e Cognição", "maximo":18},
            {"dominio":"Percepção e Exploração Rítmica", "maximo":16},
            {"dominio":"Percepção e Exploração Sonora", "maximo":14},
            {"dominio":"Exploração Vocal", "maximo":14},
            {"dominio":"Movimentação Corporal com a Música", "maximo":14},
        ],
    },
    "MEL": {
        "tipo": "barras_percentual",
        "fonte": "Modelo interno de organização por habilidades observáveis. Ajuste para o protocolo MEL utilizado pelo serviço.",
        "dominios": [
            {"dominio":"Engajamento", "maximo":10},
            {"dominio":"Comunicação", "maximo":10},
            {"dominio":"Interação", "maximo":10},
            {"dominio":"Regulação", "maximo":10},
            {"dominio":"Participação", "maximo":10},
        ],
    },
    "Escala GAS": {
        "tipo": "linha_objetivos",
        "fonte": "Escala GAS vinculada a metas -2, -1, 0, +1, +2. Usada automaticamente para MIG/TREINI.",
        "dominios": [
            {"dominio":"Meta 1", "maximo":2, "minimo":-2},
            {"dominio":"Meta 2", "maximo":2, "minimo":-2},
            {"dominio":"Meta 3", "maximo":2, "minimo":-2},
        ],
    },
    "Escala de Objetivos Mensuráveis": {
        "tipo": "linha_objetivos",
        "fonte": "Escala opcional de objetivos mensuráveis para métodos não MIG/TREINI.",
        "dominios": [
            {"dominio":"Objetivo 1", "maximo":2, "minimo":-2},
            {"dominio":"Objetivo 2", "maximo":2, "minimo":-2},
            {"dominio":"Objetivo 3", "maximo":2, "minimo":-2},
        ],
    },
}

def especificacao_grafico_avaliacao(avaliacao_nome, texto="", profissao=""):
    """Cria schema de gráfico coerente com domínios extraídos do arquivo.

    Para escalas conhecidas usa referência interna. Para uploads, tenta usar os domínios reais
    ou cabeçalhos encontrados no documento, em vez de cair sempre em indicadores genéricos.
    """
    nome = avaliacao_nome or "Avaliação clínica"
    for chave, spec in AVALIACOES_REFERENCIA.items():
        if chave.lower() in nome.lower() or nome.lower() in chave.lower():
            return json.loads(json.dumps(spec, ensure_ascii=False))

    secoes_extraidas, dominios_extraidos = extrair_estrutura_escala(texto or "", nome, profissao)
    minimo, maximo = detectar_intervalo_pontuacao(texto or "")

    dominios = []
    if dominios_extraidos:
        for d in dominios_extraidos[:12]:
            dominios.append({"dominio": d[:55], "minimo": minimo, "maximo": maximo})
    elif secoes_extraidas:
        for sec in secoes_extraidas[:12]:
            dominios.append({"dominio": sec.get("titulo", "Domínio")[:55], "minimo": minimo, "maximo": maximo})
    else:
        campos = inferir_campos_por_texto(texto or "", nome, profissao)
        for c in campos:
            if re.search(r"(dom[ií]nio|subescala|score|pontua|resultado|aten[cç][aã]o|comunica|intera|motric|regula|sensorial|linguagem|autonomia|comport)", c, re.I):
                dominios.append({"dominio": re.sub(r"^Item \d+ — ", "", c)[:55], "minimo": minimo, "maximo": maximo})

    if len(dominios) < 3:
        for campo in campos_fallback_por_profissao(profissao)[:6]:
            dominios.append({"dominio": campo.replace("Indicador de ", "")[:55], "minimo": minimo, "maximo": maximo})

    return {
        "tipo": "barras_percentual",
        "fonte": "Modelo inferido do arquivo enviado: domínios e itens foram resumidos para preservar coerência clínica sem reproduzir integralmente conteúdo protegido. Valide pontuação, máximos e sentido do escore.",
        "dominios": campos_unicos([d.get("dominio", "") for d in dominios], limite=55) and [
            {"dominio": d, "minimo": minimo, "maximo": maximo} for d in campos_unicos([x.get("dominio", "") for x in dominios], limite=55)[:12]
        ],
    }

def normalizar_grafico_spec(df, tipo="barras_percentual", fonte=""):
    dominios = []
    for _, row in df.iterrows():
        nome = normalizar_nome_campo(row.get("Domínio", ""))
        if not nome:
            continue
        try:
            maximo = float(row.get("Máximo", 10) or 10)
        except Exception:
            maximo = 10
        try:
            minimo = float(row.get("Mínimo", 0) or 0)
        except Exception:
            minimo = 0
        dominios.append({"dominio": nome, "maximo": maximo, "minimo": minimo})
    if not dominios:
        dominios = [{"dominio":"Resultado global", "maximo":10, "minimo":0}]
    return {"tipo": tipo, "fonte": fonte or "Modelo validado pelo profissional.", "dominios": dominios}

def gerar_grafico_avaliacao_modelo(titulo, spec, valores):
    dominios = spec.get("dominios", []) if spec else []
    if not dominios:
        return None
    labels = [d.get("dominio", "Domínio") for d in dominios]
    raw_vals = [float(valores.get(lbl, 0) or 0) for lbl in labels]
    maxs = [float(d.get("maximo", 10) or 10) for d in dominios]
    mins = [float(d.get("minimo", 0) or 0) for d in dominios]
    tipo = (spec.get("tipo") or "barras_percentual")

    if tipo == "radar_percentual" and len(labels) >= 3:
        vals = []
        for v, mi, ma in zip(raw_vals, mins, maxs):
            denom = ma - mi if ma != mi else ma or 1
            vals.append(max(0, min(100, ((v - mi) / denom) * 100)))
        angles = [n / float(len(labels)) * 2 * math.pi for n in range(len(labels))]
        values = vals + vals[:1]
        angles += angles[:1]
        fig = plt.figure(figsize=(7.5, 7.5))
        fig.patch.set_facecolor("#020617")
        ax = plt.subplot(111, polar=True)
        ax.set_facecolor("#0f172a")
        ax.plot(angles, values, linewidth=2)
        ax.fill(angles, values, alpha=.22)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels([x[:18] for x in labels], color="white")
        ax.set_ylim(0, 100)
        ax.grid(color="#334155")
        ax.set_title(titulo, color="white", pad=20, fontsize=14, fontweight="bold")
    elif tipo == "linha_objetivos":
        fig, ax = plt.subplots(figsize=(10, 4.8))
        fig.patch.set_facecolor("#020617")
        ax.set_facecolor("#0f172a")
        ax.plot(labels, raw_vals, marker="o", linewidth=2)
        ax.axhline(0, linewidth=1, alpha=.45)
        ax.set_ylim(min(mins)-.25, max(maxs)+.25)
        ax.set_title(titulo, color="white", fontsize=14, fontweight="bold")
        ax.tick_params(axis="x", rotation=18, colors="white")
        ax.tick_params(axis="y", colors="white")
        ax.grid(axis="y", color="#334155", alpha=.6)
    else:
        vals = []
        for v, mi, ma in zip(raw_vals, mins, maxs):
            denom = ma - mi if ma != mi else ma or 1
            vals.append(max(0, min(100, ((v - mi) / denom) * 100)))
        fig, ax = plt.subplots(figsize=(10.5, 5.2))
        fig.patch.set_facecolor("#020617")
        ax.set_facecolor("#0f172a")
        ax.bar(labels, vals)
        ax.set_ylim(0, 100)
        ax.set_ylabel("% do escore máximo", color="white")
        ax.set_title(titulo, color="white", fontsize=14, fontweight="bold")
        ax.tick_params(axis="x", rotation=25, colors="white")
        ax.tick_params(axis="y", colors="white")
        ax.grid(axis="y", color="#334155", alpha=.35)
    for sp in ["top", "right"]:
        if sp in ax.spines:
            ax.spines[sp].set_visible(False)
    if "bottom" in ax.spines: ax.spines["bottom"].set_color("#475569")
    if "left" in ax.spines: ax.spines["left"].set_color("#475569")
    plt.tight_layout()
    return fig

def texto_grafico_fiel(avaliacao_nome, spec):
    fonte = (spec or {}).get("fonte", "")
    return (
        f"Gráfico configurado para {avaliacao_nome}. O NEXO usa domínios e limites de pontuação salvos no perfil do profissional; "
        f"para avaliações proprietárias, a fidelidade depende da validação clínica do terapeuta e da autorização de uso do instrumento. {fonte}"
    )


AVALIACOES_PADRAO_SISTEMA = ["Nordoff-Robbins", "IAPS", "DEMUCA", "MEL"]


def usuario_pode_usar_avaliacoes_padrao(user):
    """Avaliações padrão internas ficam restritas ao administrador ou a usuários autorizados por ele."""
    if not user:
        return False
    return user.get("role") == "admin" or int(user.get("avaliacoes_padrao_autorizado") or 0) == 1


def avaliacoes_padrao_disponiveis(user, profissao=""):
    if usuario_pode_usar_avaliacoes_padrao(user):
        return AVALIACOES_PADRAO_SISTEMA[:]
    return []


def usuario_deve_usar_modelos_proprios(user):
    return not usuario_pode_usar_avaliacoes_padrao(user)


def fig_to_buffer(fig):
    if fig is None:
        return None
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    buffer.seek(0)
    return buffer


def listar_respostas_relatorio(user_id, paciente_id=None):
    conn = db(); c = conn.cursor()
    if paciente_id:
        c.execute("""SELECT r.id,r.avaliacao_nome,r.interpretacao,r.justificativa,r.grafico_valores_json,r.criado_em,
                          m.nome,m.grafico_json,m.profissao
                   FROM avaliacao_respostas r
                   LEFT JOIN avaliacao_modelos m ON m.id=r.modelo_id
                   WHERE r.user_id=? AND r.usar_no_relatorio=1 AND (r.paciente_id=? OR r.paciente_id IS NULL)
                   ORDER BY r.id DESC LIMIT 12""", (user_id, paciente_id))
    else:
        c.execute("""SELECT r.id,r.avaliacao_nome,r.interpretacao,r.justificativa,r.grafico_valores_json,r.criado_em,
                          m.nome,m.grafico_json,m.profissao
                   FROM avaliacao_respostas r
                   LEFT JOIN avaliacao_modelos m ON m.id=r.modelo_id
                   WHERE r.user_id=? AND r.usar_no_relatorio=1
                   ORDER BY r.id DESC LIMIT 12""", (user_id,))
    rows = c.fetchall(); conn.close()
    out = []
    for row in rows:
        try:
            valores = json.loads(row[4] or "{}")
        except Exception:
            valores = {}
        try:
            spec = json.loads(row[7] or "{}")
        except Exception:
            spec = {}
        if not spec:
            spec = especificacao_grafico_avaliacao(row[1], "", row[8] or "Outro")
        out.append({
            "id": row[0], "avaliacao_nome": row[1], "interpretacao": row[2], "justificativa": row[3],
            "grafico_valores": valores, "criado_em": row[5], "modelo_nome": row[6] or row[1], "grafico_spec": spec,
        })
    return out

def salvar_modelo_avaliacao(user_id, nome, profissao, avaliacao_detectada, descricao, secoes, arquivo_origem="", caminho_origem="", grafico_spec=None, fonte_referencia="", gestor_id=None, setor_id=None, escopo="perfil", area_alvo=""):
    """Salva um modelo de avaliação.

    escopo='perfil'  -> avaliação privada do profissional.
    escopo='clinica' -> avaliação padrão da clínica, criada pelo gestor do Modo Clínica.
    """
    if grafico_spec is None:
        grafico_spec = especificacao_grafico_avaliacao(avaliacao_detectada, "", profissao)
    conn = db(); c = conn.cursor()
    campos = []
    for sec in secoes:
        campos.extend(sec.get("campos", []))
    now = datetime.now().isoformat(timespec="seconds")
    c.execute("""INSERT INTO avaliacao_modelos
        (user_id,nome,profissao,avaliacao_detectada,descricao,campos_json,secoes_json,arquivo_origem,caminho_origem,grafico_json,fonte_referencia,ativo,criado_em,atualizado_em,gestor_id,setor_id,escopo,area_alvo)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (user_id,nome,profissao,avaliacao_detectada,descricao,json.dumps(campos,ensure_ascii=False),json.dumps(secoes,ensure_ascii=False),arquivo_origem,caminho_origem,json.dumps(grafico_spec,ensure_ascii=False),fonte_referencia or grafico_spec.get("fonte",""),1,now,now,gestor_id,setor_id,escopo,area_alvo or profissao))
    conn.commit(); mid = c.lastrowid; conn.close(); return mid

def listar_modelos_avaliacao(user_id, profissao=None, gestor_id=None, setor_id=None, incluir_clinica=True):
    """
    Lista modelos próprios e, quando aplicável, modelos padrão da clínica.

    Correção importante:
    - A versão anterior montava a query como ``WHERE AND (...)`` em alguns fluxos,
      gerando sqlite3.OperationalError: near "AND": syntax error.
    - Agora a query é construída por blocos explícitos: condições obrigatórias,
      escopos permitidos e filtros opcionais.
    """
    conn = db()
    c = conn.cursor()

    condicoes = ["ativo = 1"]
    params = []

    escopos = ["((escopo IS NULL OR escopo = 'perfil') AND user_id = ?)"]
    params.append(user_id)

    if incluir_clinica and gestor_id:
        clinica_condicoes = ["escopo = 'clinica'", "gestor_id = ?"]
        params.append(gestor_id)
        if setor_id is not None and str(setor_id).strip() != "":
            clinica_condicoes.append("(setor_id = ? OR setor_id IS NULL)")
            params.append(setor_id)
        escopos.append("(" + " AND ".join(clinica_condicoes) + ")")

    condicoes.append("(" + " OR ".join(escopos) + ")")

    if profissao and str(profissao).strip():
        condicoes.append("(profissao = ? OR profissao = 'Outro' OR COALESCE(area_alvo,'') = ? OR COALESCE(area_alvo,'') = '')")
        params.extend([profissao, profissao])

    sql = """
        SELECT id, nome, profissao, avaliacao_detectada, descricao, criado_em,
               escopo, setor_id, gestor_id, area_alvo
        FROM avaliacao_modelos
        WHERE {where_sql}
        ORDER BY CASE WHEN escopo = 'clinica' THEN 0 ELSE 1 END, id DESC
    """.format(where_sql=" AND ".join(condicoes))

    try:
        c.execute(sql, params)
        rows = c.fetchall()
    except sqlite3.OperationalError as e:
        conn.close()
        raise sqlite3.OperationalError(f"Erro ao listar modelos de avaliação. SQL: {sql} | params={params} | erro={e}")

    conn.close()
    return rows

def buscar_modelo_avaliacao(user_id, modelo_id, gestor_id=None):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT id,nome,profissao,avaliacao_detectada,descricao,campos_json,secoes_json,grafico_json,fonte_referencia,escopo,setor_id,gestor_id,area_alvo
                 FROM avaliacao_modelos
                 WHERE id=? AND ativo=1 AND (user_id=? OR (escopo='clinica' AND gestor_id=?))""", (modelo_id, user_id, gestor_id or user_id))
    row=c.fetchone(); conn.close()
    if not row: return None
    grafico = json.loads(row[7] or "{}") if len(row) > 7 else {}
    if not grafico:
        grafico = especificacao_grafico_avaliacao(row[3], "", row[2])
    return {"id":row[0],"nome":row[1],"profissao":row[2],"avaliacao_detectada":row[3],"descricao":row[4],"campos":json.loads(row[5] or "[]"),"secoes":json.loads(row[6] or "[]"),"grafico":grafico,"fonte_referencia":row[8] if len(row)>8 else "", "escopo": row[9] if len(row)>9 else "perfil", "setor_id": row[10] if len(row)>10 else None, "gestor_id": row[11] if len(row)>11 else None, "area_alvo": row[12] if len(row)>12 else ""}

def salvar_resposta_avaliacao(user_id, paciente_id, modelo_id, avaliacao_nome, respostas, interpretacao, justificativa, usar_no_relatorio=True, grafico_valores=None):
    if grafico_valores is None:
        grafico_valores = {}
    conn=db(); c=conn.cursor()
    c.execute("""INSERT INTO avaliacao_respostas
        (user_id,paciente_id,modelo_id,avaliacao_nome,respostas_json,interpretacao,justificativa,grafico_valores_json,usar_no_relatorio,criado_em)
        VALUES (?,?,?,?,?,?,?,?,?,?)""",
        (user_id,paciente_id,modelo_id,avaliacao_nome,json.dumps(respostas,ensure_ascii=False),interpretacao,justificativa,json.dumps(grafico_valores,ensure_ascii=False),1 if usar_no_relatorio else 0,datetime.now().isoformat(timespec="seconds")))
    conn.commit(); conn.close()

def gerar_interpretacao_modelo(profissao, avaliacao_nome, respostas):
    perfil = perfil_profissional(profissao)
    preenchidos = {k:v for k,v in respostas.items() if str(v).strip()}
    linhas = [f"A avaliação {avaliacao_nome} foi preenchida considerando o perfil profissional de {profissao}."]
    linhas.append(f"A leitura prioriza {perfil['foco']}.")
    if preenchidos:
        linhas.append("Indicadores informados: " + "; ".join([f"{k}: {str(v)[:120]}" for k,v in list(preenchidos.items())[:8]]) + ".")
    linhas.append("A partir desses dados, o sistema recomenda vincular os resultados a objetivos mensuráveis, acompanhamento evolutivo e revisão periódica da conduta.")
    return "\n\n".join(linhas)

def log_uso(user_id, tipo):
    conn = db()
    c = conn.cursor()
    c.execute("INSERT INTO usage_logs (user_id, tipo, criado_em) VALUES (?, ?, ?)",
              (user_id, tipo, datetime.now().isoformat(timespec="seconds")))
    conn.commit()
    conn.close()

def uso_mes(user_id, tipo):
    conn = db()
    c = conn.cursor()
    prefix = datetime.now().strftime("%Y-%m")
    c.execute("SELECT COUNT(*) FROM usage_logs WHERE user_id=? AND tipo=? AND criado_em LIKE ?",
              (user_id, tipo, f"{prefix}%"))
    n = c.fetchone()[0]
    conn.close()
    return n

def listar_usuarios():
    conn = db()
    c = conn.cursor()
    c.execute("SELECT * FROM users ORDER BY id DESC")
    rows = [row_to_user(r) for r in c.fetchall()]
    conn.close()
    return rows

init_db()

# ==========================================================
# MERCADO PAGO — ASSINATURAS RECORRENTES
# ==========================================================
APP_PUBLIC_URL_PADRAO = "https://nexoclinic.streamlit.app/"

def criar_preferencia_mercado_pago(user_id, plano):
    """
    Cria um link de ASSINATURA recorrente mensal no Mercado Pago.

    Importante:
    - Usa a API /preapproval, própria para assinaturas.
    - Não usa auto_return nem assinatura/preferences, que são de pagamento único.
    - O token deve ficar fora do código, em variável de ambiente:
      MERCADO_PAGO_ACCESS_TOKEN=seu_access_token
    - O retorno usa a URL pública do Streamlit por padrão.
    """
    token = os.getenv("MERCADO_PAGO_ACCESS_TOKEN", "").strip()
    app_url = os.getenv("APP_PUBLIC_URL", APP_PUBLIC_URL_PADRAO).strip() or APP_PUBLIC_URL_PADRAO
    plano = normalizar_plano_nome(plano)
    valor = float(PLANOS[plano]["valor"])

    conn = db()
    c = conn.cursor()

    if not token:
        c.execute("""INSERT INTO pagamentos (user_id, plano, valor, status, assinatura_url, criado_em, atualizado_em)
                  VALUES (?, ?, ?, ?, ?, ?, ?)""",
                  (user_id, plano, valor, "pendente_sem_token", "",
                   datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")))
        conn.commit()
        conn.close()
        return False, "Configure MERCADO_PAGO_ACCESS_TOKEN para gerar assinatura real.", ""

    try:
        import requests

        user = buscar_usuario(user_id) or {}
        payer_email = (user.get("email") or "").strip().lower()
        if not payer_email:
            conn.close()
            return False, "Usuário sem e-mail cadastrado. Não foi possível criar a assinatura.", ""

        payload = {
            "reason": f"NEXO - Plano {plano}",
            "external_reference": f"nexo_user_{user_id}_plano_{plano}",
            "payer_email": payer_email,
            "back_url": app_url,
            "auto_recurring": {
                "frequency": 1,
                "frequency_type": "months",
                "transaction_amount": valor,
                "currency_id": "BRL"
            },
            "status": "pending"
        }

        r = requests.post(
            "https://api.mercadopago.com/preapproval",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json=payload,
            timeout=25
        )

        try:
            data = r.json()
        except Exception:
            data = {"raw_response": r.text}

        if r.status_code >= 300:
            conn.close()
            return False, f"Erro Mercado Pago: {data}", ""

        assinatura_id = data.get("id", "")
        assinatura_url = data.get("init_point") or data.get("sandbox_init_point") or data.get("link", "")

        c.execute("""INSERT INTO pagamentos (user_id, plano, valor, status, mercado_pago_preference_id,
                  assinatura_url, criado_em, atualizado_em) VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                  (user_id, plano, valor, "assinatura_pendente", assinatura_id, assinatura_url,
                   datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")))
        conn.commit()
        conn.close()

        if not assinatura_url:
            return False, f"Assinatura criada, mas o Mercado Pago não retornou link. ID: {assinatura_id}", ""

        return True, "Assinatura criada.", assinatura_url

    except Exception as e:
        conn.close()
        return False, f"Falha ao criar assinatura: {e}", ""

# ==========================================================
# CSS PREMIUM
# ==========================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

html, body, [class^="st"], [data-testid="stAppViewContainer"] {
    font-family:'Inter','Segoe UI',sans-serif !important;
}

/* Mantém a fonte correta dos ícones internos do Streamlit/Material Icons */
.material-icons,
.material-icons-outlined,
.material-icons-round,
.material-icons-sharp,
.material-symbols-outlined,
.material-symbols-rounded,
.material-symbols-sharp,
span[class*="material-icons"],
span[class*="material-symbols"],
i[class*="material-icons"],
i[class*="material-symbols"] {
    font-family:'Material Icons','Material Symbols Outlined','Material Symbols Rounded' !important;
    font-weight:normal !important;
    font-style:normal !important;
    font-size:inherit;
    line-height:1 !important;
    letter-spacing:normal !important;
    text-transform:none !important;
    display:inline-flex !important;
    white-space:nowrap !important;
    word-wrap:normal !important;
    direction:ltr !important;
    -webkit-font-feature-settings:'liga' !important;
    -webkit-font-smoothing:antialiased !important;
}

#MainMenu, footer, header, [data-testid="stDecoration"] { visibility:hidden; height:0; }
[data-testid="stAppViewContainer"] {
    background: radial-gradient(circle at 24% 0%, rgba(79,70,229,.16), transparent 30%),
                radial-gradient(circle at 82% 8%, rgba(14,165,233,.13), transparent 30%),
                linear-gradient(135deg,#050b18 0%,#07111f 46%,#0b1120 100%);
    color:#eef6ff;
}
.block-container { padding-top:1.5rem; padding-bottom:3rem; max-width:1500px; }
[data-testid="stSidebar"] {
    background:linear-gradient(180deg,#040a15,#07101d 56%,#060b17);
    border-right:1px solid rgba(148,163,184,.18);
}
[data-testid="stSidebar"] * { color:#cbd7ee; }
[data-testid="stSidebar"] .stRadio > label { display:none; }
[data-testid="stSidebar"] .stRadio [role="radiogroup"] { gap:8px; }
[data-testid="stSidebar"] .stRadio [role="radio"] {
    background:transparent !important;
    border-radius:13px !important;
    padding:12px 14px !important;
    min-height:42px !important;
    border:1px solid transparent !important;
}
[data-testid="stSidebar"] .stRadio [role="radio"][aria-checked="true"] {
    background:linear-gradient(90deg,#4f46e5,#7c3aed) !important;
    border:1px solid rgba(255,255,255,.14) !important;
    box-shadow:0 14px 36px rgba(79,70,229,.34);
}
[data-testid="stSidebar"] .stRadio [role="radio"] > div:first-child { display:none !important; }
h1,h2,h3 { color:#f8fafc !important; letter-spacing:-.04em; }
p, label, div { line-height:1.35 !important; }
span:not([class*='material-icons']):not([class*='material-symbols']) { line-height:1.35 !important; }
.stTextInput input, .stNumberInput input, .stTextArea textarea,
.stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"], .stDateInput input {
    border-radius:14px !important;
    border:1px solid rgba(148,163,184,.18) !important;
    background:rgba(11,18,32,.72) !important;
    color:#f8fafc !important;
}
.stButton > button, .stDownloadButton > button {
    width:100%;
    min-height:3rem;
    border-radius:14px;
    border:1px solid rgba(255,255,255,.12);
    background:linear-gradient(90deg,#7c3aed,#0ea5e9);
    color:#fff;
    font-weight:800;
    box-shadow:0 16px 40px rgba(59,130,246,.25);
}
.stButton > button:hover, .stDownloadButton > button:hover {
    transform:translateY(-1px);
    color:#fff;
}
.stTabs [data-baseweb="tab-list"] { gap:8px; border-bottom:1px solid rgba(148,163,184,.12); }
.stTabs [data-baseweb="tab"] {
    background:rgba(15,23,42,.62);
    border-radius:14px 14px 0 0;
    border:1px solid rgba(148,163,184,.12);
    padding:11px 15px;
}
.stTabs [aria-selected="true"] {
    background:linear-gradient(90deg,rgba(124,58,237,.38),rgba(14,165,233,.24));
    border-bottom:2px solid #38bdf8;
}
div[data-testid="stExpander"] {
    border:0 !important;
    border-radius:18px !important;
    overflow:hidden !important;
    margin-bottom:12px !important;
}
div[data-testid="stExpander"] details summary {
    list-style-type:none !important;
    min-height:auto !important;
    padding:15px 18px !important;
    border-radius:18px !important;
    background:rgba(15,23,42,.70) !important;
    border:1px solid rgba(148,163,184,.14) !important;
}
div[data-testid="stExpander"] details summary::-webkit-details-marker {
    display:none !important;
}


.upload-label {
    color:#f8fafc;
    font-weight:850;
    margin:12px 0 8px;
}

/* ==========================================================
   UPLOAD STARTUP — dropzone isolado e sem duplicação de texto
   ========================================================== */
.upload-label {
    color:#f8fafc;
    font-weight:850;
    margin:12px 0 8px;
}

/* Container geral do uploader */
[data-testid="stFileUploader"] {
    width:100% !important;
    margin-top:8px !important;
    margin-bottom:22px !important;
}

/* Evita que labels internos do Streamlit entrem no layout visual */
[data-testid="stFileUploader"] > label {
    color:#e5edf8 !important;
    font-weight:800 !important;
    margin-bottom:10px !important;
}

/* Área principal do upload */
[data-testid="stFileUploaderDropzone"] {
    position:relative !important;
    min-height:150px !important;
    padding:30px 30px !important;
    border-radius:22px !important;
    border:1px dashed rgba(96,165,250,.50) !important;
    background:
      radial-gradient(circle at 92% 50%, rgba(14,165,233,.16), transparent 26%),
      linear-gradient(145deg,rgba(13,22,39,.94),rgba(17,25,42,.72)) !important;
    display:flex !important;
    align-items:center !important;
    justify-content:space-between !important;
    overflow:hidden !important;
    box-shadow:inset 0 0 0 1px rgba(255,255,255,.025), 0 18px 48px rgba(2,6,23,.18) !important;
}

/* Texto premium do dropzone — gerenciado pelo nosso CSS */
[data-testid="stFileUploaderDropzone"]::before {
    content:"Arraste arquivos aqui ou clique para selecionar";
    position:absolute;
    left:30px;
    top:34px;
    max-width:calc(100% - 230px);
    color:#e5edf8;
    font-size:1.04rem;
    font-weight:850;
    line-height:1.2;
    pointer-events:none;
}
[data-testid="stFileUploaderDropzone"]::after {
    content:"PDF, Word, Excel, imagens e documentos clínicos";
    position:absolute;
    left:30px;
    top:68px;
    max-width:calc(100% - 230px);
    color:#94a3b8;
    font-size:.9rem;
    font-weight:500;
    line-height:1.2;
    pointer-events:none;
}

/* Esconde instruções nativas que costumam duplicar o conteúdo */
[data-testid="stFileUploaderDropzoneInstructions"],
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] div:has(small) {
    display:none !important;
    visibility:hidden !important;
    width:0 !important;
    height:0 !important;
    overflow:hidden !important;
}

/* Botão interno do Streamlit: isolado dos botões globais do app */
[data-testid="stFileUploaderDropzone"] button,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"] {
    position:absolute !important;
    right:30px !important;
    top:50% !important;
    transform:translateY(-50%) !important;
    width:150px !important;
    min-width:150px !important;
    max-width:150px !important;
    height:50px !important;
    min-height:50px !important;
    padding:0 !important;
    border-radius:15px !important;
    border:1px solid rgba(255,255,255,.14) !important;
    background:linear-gradient(90deg,#7c3aed,#0ea5e9) !important;
    box-shadow:0 16px 38px rgba(59,130,246,.28) !important;
    color:transparent !important;
    overflow:hidden !important;
    z-index:4 !important;
}

/* Esconde completamente o texto nativo do botão para impedir uploadupload */
[data-testid="stFileUploaderDropzone"] button *,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"] * {
    font-size:0 !important;
    line-height:0 !important;
    color:transparent !important;
    opacity:0 !important;
    width:0 !important;
    max-width:0 !important;
    overflow:hidden !important;
}

/* Texto único e controlado do botão */
[data-testid="stFileUploaderDropzone"] button::after,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"]::after {
    content:"Selecionar";
    position:absolute;
    inset:0;
    display:flex;
    align-items:center;
    justify-content:center;
    color:#fff;
    font-size:.94rem;
    font-weight:900;
    letter-spacing:-.01em;
    line-height:1;
}

[data-testid="stFileUploaderDropzone"] button:hover,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"]:hover {
    transform:translateY(-51%) scale(1.015) !important;
    filter:brightness(1.04) !important;
}

/* Arquivos selecionados */
[data-testid="stFileUploaderFile"] {
    border-radius:14px !important;
    border:1px solid rgba(148,163,184,.16) !important;
    background:rgba(15,23,42,.78) !important;
    color:#e5edf8 !important;
}

@media(max-width:900px){
    [data-testid="stFileUploaderDropzone"] {
        min-height:190px !important;
        padding:26px !important;
    }
    [data-testid="stFileUploaderDropzone"]::before,
    [data-testid="stFileUploaderDropzone"]::after {
        max-width:calc(100% - 52px);
        left:26px;
    }
    [data-testid="stFileUploaderDropzone"] button,
    [data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"] {
        left:26px !important;
        right:26px !important;
        top:auto !important;
        bottom:24px !important;
        transform:none !important;
        width:calc(100% - 52px) !important;
        min-width:0 !important;
        max-width:none !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover,
    [data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"]:hover {
        transform:scale(1.01) !important;
    }
}


[data-testid="stMetric"] {
    background:linear-gradient(145deg,rgba(13,22,39,.90),rgba(17,25,42,.60));
    border:1px solid rgba(148,163,184,.14);
    border-radius:20px;
    padding:1rem;
}
.logo-wrap {
    display:flex;
    align-items:center;
    justify-content:center;
    border-radius:16px;
    background:linear-gradient(135deg,rgba(124,58,237,.12),rgba(14,165,233,.10));
}
.sidebar-brand { padding:10px 2px 26px; }
.sidebar-brand-row { display:flex; align-items:center; gap:12px; }
.sidebar-brand h1 { margin:0 !important; font-size:1.8rem !important; font-weight:950 !important; letter-spacing:-.08em; }
.sidebar-brand p { margin:0; color:#aab7cf; font-size:.66rem; letter-spacing:.06em; text-transform:uppercase; }
.sidebar-section { color:#7587a6; font-size:.72rem; letter-spacing:.08em; text-transform:uppercase; margin:20px 0 8px; padding-left:14px; }
.sidebar-bottom {
    position:fixed;
    bottom:22px;
    width:260px;
    border:1px solid rgba(148,163,184,.14);
    border-radius:16px;
    padding:15px;
    background:linear-gradient(145deg,rgba(15,23,42,.82),rgba(30,41,59,.38));
}
.user-line { display:flex; gap:10px; align-items:center; }
.avatar {
    width:46px;
    height:46px;
    border-radius:999px;
    display:flex;
    align-items:center;
    justify-content:center;
    background:linear-gradient(135deg,#7c3aed,#4f46e5);
    color:#fff;
    font-weight:900;
}
.topbar {
    display:grid;
    grid-template-columns:minmax(280px, 440px) 1fr auto;
    gap:18px;
    align-items:center;
    margin-bottom:22px;
}
.searchbox {
    height:48px;
    border-radius:13px;
    border:1px solid rgba(148,163,184,.16);
    background:rgba(15,23,42,.62);
    display:flex;
    align-items:center;
    gap:10px;
    padding:0 16px;
    color:#8192b0;
}
.top-icons { display:flex; justify-content:flex-end; gap:18px; align-items:center; }
.profile-chip {
    display:flex;
    align-items:center;
    gap:12px;
    border-left:1px solid rgba(148,163,184,.18);
    padding-left:18px;
}
.layout-grid { display:grid; grid-template-columns:1fr 310px; gap:18px; }
.hero {
    position:relative;
    overflow:hidden;
    min-height:315px;
    border-radius:18px;
    border:1px solid rgba(59,130,246,.22);
    background:radial-gradient(circle at 74% 44%, rgba(124,58,237,.34), transparent 22%),
               radial-gradient(circle at 88% 34%, rgba(14,165,233,.22), transparent 26%),
               linear-gradient(135deg, rgba(17,24,53,.94), rgba(13,20,38,.88));
    box-shadow:0 24px 68px rgba(2,6,23,.35);
    padding:32px;
}
.hero h2 { font-size:2.25rem !important; line-height:1.08; margin:16px 0 12px !important; max-width:650px; }
.gradient-text { background:linear-gradient(90deg,#60a5fa,#22d3ee,#a78bfa); -webkit-background-clip:text; -webkit-text-fill-color:transparent; }
.hero p { color:#c6d1e3; max-width:620px; font-size:1.02rem; }
.hero-actions { display:flex; gap:14px; margin-top:28px; max-width:430px; }
.fake-button {
    border:1px solid rgba(255,255,255,.14);
    border-radius:13px;
    padding:13px 20px;
    display:flex;
    align-items:center;
    gap:10px;
    font-weight:800;
    background:rgba(15,23,42,.52);
}
.fake-button.primary { background:linear-gradient(90deg,#7c3aed,#0ea5e9); border:0; }
.hero-brain {
    position:absolute;
    right:70px;
    top:62px;
    width:245px;
    height:245px;
    display:flex;
    align-items:center;
    justify-content:center;
    border-radius:999px;
    background:radial-gradient(circle,#172554 0%,rgba(79,70,229,.22) 42%,transparent 72%);
    opacity:.96;
}
.card {
    background:linear-gradient(145deg,rgba(13,22,39,.90),rgba(17,25,42,.62));
    border:1px solid rgba(148,163,184,.14);
    border-radius:18px;
    padding:22px;
    box-shadow:0 18px 52px rgba(2,6,23,.26);
    margin-bottom:18px;
}
.flow-step { display:grid; grid-template-columns:34px 1fr; gap:12px; margin:18px 0; align-items:start; }
.flow-num {
    width:26px; height:26px; border-radius:999px; display:flex; align-items:center; justify-content:center;
    background:linear-gradient(135deg,#7c3aed,#0ea5e9); font-weight:900; color:#fff;
}
.flow-step strong { color:#f8fafc; display:block; }
.flow-step span { color:#a7b3c8; font-size:.88rem; }
.stats-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:14px; margin:18px 0; }
.stat-card {
    background:linear-gradient(145deg,rgba(13,22,39,.90),rgba(17,25,42,.62));
    border:1px solid rgba(148,163,184,.14);
    border-radius:18px;
    padding:20px;
    display:flex;
    gap:16px;
    align-items:center;
}
.stat-icon {
    width:58px; height:58px; border-radius:18px; display:flex; align-items:center; justify-content:center;
    background:rgba(79,70,229,.18);
}
.stat-label { color:#cbd5e1; }
.stat-number { color:#fff; font-weight:950; font-size:2rem; letter-spacing:-.06em; }
.stat-sub { color:#34d399; font-size:.82rem; }
.quick-grid { display:grid; grid-template-columns:repeat(5,1fr); gap:11px; }
.quick-card {
    min-height:62px;
    border:1px solid rgba(148,163,184,.14);
    border-radius:13px;
    background:rgba(15,23,42,.48);
    display:flex;
    align-items:center;
    justify-content:center;
    gap:10px;
    font-weight:760;
    text-align:center;
}
.activity-row {
    display:grid;
    grid-template-columns:1fr auto auto;
    gap:14px;
    align-items:center;
    padding:14px 8px;
    border-bottom:1px solid rgba(148,163,184,.10);
}
.badge {
    display:inline-flex;
    align-items:center;
    justify-content:center;
    border-radius:9px;
    padding:5px 10px;
    background:rgba(124,58,237,.18);
    color:#a78bfa;
    font-size:.82rem;
    font-weight:800;
}
.plan-progress { height:7px; background:rgba(148,163,184,.14); border-radius:999px; overflow:hidden; margin:10px 0 15px; }
.plan-progress div { height:100%; background:linear-gradient(90deg,#7c3aed,#0ea5e9); border-radius:999px; }
.small-muted { color:#a7b3c8; font-size:.93rem; }
.price-grid { display:grid; grid-template-columns:repeat(5,1fr); gap:14px; }
.price-card { min-height:260px; display:flex; flex-direction:column; justify-content:space-between; }
.price { font-size:2rem; font-weight:950; color:#fff; letter-spacing:-.06em; }
@media(max-width:1200px){
    .layout-grid{grid-template-columns:1fr;}
    .stats-grid{grid-template-columns:repeat(2,1fr);}
    .quick-grid{grid-template-columns:repeat(2,1fr);}
    .price-grid{grid-template-columns:repeat(2,1fr);}
    .hero-brain{display:none;}
    .topbar{grid-template-columns:1fr;}
    .sidebar-bottom{position:static;width:auto;margin-top:20px;}
}

/* Correção específica para textos de ícone do expander aparecendo como:
   keyboard_arrow_down / keyboard_double_arrow_down */
div[data-testid="stExpander"] summary span[class*="material-symbols"],
div[data-testid="stExpander"] summary span[class*="material-icons"] {
    font-family:'Material Symbols Outlined','Material Icons' !important;
    overflow:hidden !important;
    max-width:24px !important;
}

</style>
""", unsafe_allow_html=True)



# ==========================================================
# NEXO UI 2.0 — DESIGN PREMIUM HEALTH TECH
# ==========================================================
st.markdown("""
<style>
:root{
  --nexo-bg:#07111f;
  --nexo-panel:#0f1b2d;
  --nexo-panel-2:#132238;
  --nexo-border:rgba(148,163,184,.15);
  --nexo-muted:#94a3b8;
  --nexo-text:#f8fafc;
  --nexo-blue:#60a5fa;
  --nexo-purple:#8b5cf6;
  --nexo-green:#34d399;
  --nexo-orange:#fb923c;
  --nexo-red:#f87171;
}
html, body, [class^="st"], [data-testid="stAppViewContainer"]{font-family:Inter,Segoe UI,system-ui,sans-serif!important;}
[data-testid="stAppViewContainer"]{
  background:
    radial-gradient(circle at 18% 0%, rgba(99,102,241,.20), transparent 30%),
    radial-gradient(circle at 95% 3%, rgba(14,165,233,.16), transparent 24%),
    linear-gradient(135deg,#020617 0%,#07111f 48%,#0b1220 100%)!important;
}
.block-container{max-width:1700px!important;padding:10px 18px 30px!important;}
#MainMenu, footer, header, [data-testid="stDecoration"]{visibility:hidden;height:0;}

/* Layout principal */
[data-testid="column"]:has(.nexo-sidebar-sentinel){
  background:linear-gradient(180deg,rgba(5,12,27,.98),rgba(8,18,33,.96));
  border:1px solid var(--nexo-border);
  border-radius:22px;
  padding:18px 14px!important;
  min-height:calc(100vh - 26px);
  box-shadow:0 24px 70px rgba(0,0,0,.28);
}
.sidebar-brand{padding:6px 6px 16px;border-bottom:1px solid rgba(148,163,184,.10);margin-bottom:14px;}
.sidebar-brand-row{display:flex;align-items:center;gap:12px;}
.sidebar-brand h1{font-size:2rem!important;line-height:.9!important;margin:0!important;letter-spacing:-.08em!important;}
.sidebar-brand p{font-size:.74rem!important;color:#9fb0cf!important;margin:2px 0 0!important;}
.logo-wrap{border-radius:14px;background:linear-gradient(135deg,#7c3aed,#0ea5e9);display:flex;align-items:center;justify-content:center;box-shadow:0 12px 30px rgba(99,102,241,.3);}
.sidebar-section{font-size:.72rem;text-transform:uppercase;letter-spacing:.16em;color:#64748b;margin:18px 8px 8px;font-weight:900;}
.sidebar-user-card,.nexo-plan-card{border:1px solid var(--nexo-border);background:rgba(15,23,42,.72);border-radius:16px;padding:14px;margin:14px 0;}
.user-line{display:flex;align-items:center;gap:10px;}
.avatar{width:42px;height:42px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-weight:900;color:white;background:linear-gradient(135deg,#7c3aed,#2563eb);box-shadow:0 12px 30px rgba(79,70,229,.25);}
.sidebar-actions-note{font-size:.75rem;text-transform:uppercase;letter-spacing:.13em;color:#64748b;margin:18px 8px 8px;font-weight:900;}

/* Navegação como app real */
[data-testid="column"]:has(.nexo-sidebar-sentinel) .stButton>button{
  background:transparent!important;border:1px solid transparent!important;box-shadow:none!important;color:#dbeafe!important;
  justify-content:flex-start!important;text-align:left!important;min-height:46px!important;border-radius:14px!important;padding:0 14px!important;font-weight:750!important;
}
[data-testid="column"]:has(.nexo-sidebar-sentinel) .stButton>button:hover{
  background:rgba(99,102,241,.16)!important;border-color:rgba(99,102,241,.25)!important;transform:none!important;
}
.nexo-nav-active-marker + div .stButton>button,
[data-testid="column"]:has(.nexo-sidebar-sentinel) div:has(>.nexo-nav-active-marker) + div .stButton>button{
  background:linear-gradient(90deg,#6366f1,#2563eb)!important;color:#fff!important;box-shadow:0 12px 35px rgba(37,99,235,.28)!important;
}

/* Topbar */
.nexo-topbar{display:grid;grid-template-columns:1fr 420px 128px;gap:18px;align-items:center;margin-bottom:18px;}
.nexo-greeting h2{font-size:1.95rem!important;letter-spacing:-.055em!important;margin:0!important;color:white!important;}
.nexo-greeting p{margin:4px 0 0;color:#9fb0cf;font-size:.95rem;}
.nexo-top-actions{display:flex;gap:12px;justify-content:flex-end;align-items:center;}
.nexo-icon-btn{width:44px;height:44px;border-radius:14px;border:1px solid var(--nexo-border);background:rgba(15,23,42,.72);display:flex;align-items:center;justify-content:center;position:relative;}
.nexo-notify{position:absolute;right:-4px;top:-5px;background:#ef4444;color:white;border-radius:50%;font-size:.70rem;font-weight:900;width:19px;height:19px;display:flex;align-items:center;justify-content:center;}
[data-testid="stTextInput"] input{min-height:48px!important;border-radius:16px!important;background:rgba(15,23,42,.76)!important;border:1px solid rgba(148,163,184,.16)!important;box-shadow:inset 0 0 0 1px rgba(255,255,255,.02)!important;}

/* Cards */
.nexo-grid-4{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:16px;margin:12px 0 18px;}
.nexo-grid-3{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:16px;}
.nexo-grid-2{display:grid;grid-template-columns:1fr 300px;gap:16px;}
.nexo-card{border:1px solid var(--nexo-border);border-radius:18px;background:linear-gradient(180deg,rgba(20,32,50,.86),rgba(15,23,42,.70));box-shadow:0 18px 55px rgba(0,0,0,.22);padding:20px;position:relative;overflow:hidden;}
.nexo-card:before{content:"";position:absolute;inset:0;background:radial-gradient(circle at 80% 0%,rgba(96,165,250,.10),transparent 36%);pointer-events:none;}
.nexo-card h3{font-size:1.08rem!important;margin:0 0 12px!important;letter-spacing:-.03em!important;color:#f8fafc!important;}
.nexo-muted{color:#94a3b8;font-size:.92rem;}
.nexo-stat{min-height:138px;}
.nexo-stat-head{display:flex;align-items:center;justify-content:space-between;margin-bottom:18px;}
.nexo-stat-icon{width:44px;height:44px;border-radius:13px;display:flex;align-items:center;justify-content:center;background:rgba(99,102,241,.22);border:1px solid rgba(99,102,241,.25);}
.nexo-stat-label{font-size:.92rem;font-weight:850;color:#f8fafc;}
.nexo-stat-num{font-size:2.15rem;font-weight:950;color:#fff;letter-spacing:-.065em;line-height:1;}
.nexo-positive{color:#34d399;font-weight:800;font-size:.86rem;margin-top:10px;}
.nexo-spark{height:34px;width:92px;}

/* Acesso rápido */
.nexo-quick-grid{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:14px;margin-top:12px;}
.nexo-quick-card{height:112px;border-radius:16px;border:1px solid rgba(148,163,184,.14);background:linear-gradient(135deg,rgba(99,102,241,.20),rgba(15,23,42,.70));display:flex;flex-direction:column;align-items:center;justify-content:center;text-align:center;gap:9px;font-weight:850;color:#f8fafc;}
.nexo-quick-card:hover{border-color:rgba(96,165,250,.42);transform:translateY(-1px);}

/* Agenda */
.nexo-agenda-row{display:grid;grid-template-columns:58px 16px 1fr auto;gap:12px;align-items:start;padding:11px 0;border-bottom:1px solid rgba(148,163,184,.10);}
.nexo-dot{width:11px;height:11px;border-radius:50%;background:#34d399;margin-top:5px;box-shadow:0 0 0 4px rgba(52,211,153,.10);}
.nexo-dot.warn{background:#fb923c;box-shadow:0 0 0 4px rgba(251,146,60,.10);}
.nexo-pill{font-size:.72rem;font-weight:900;padding:5px 9px;border-radius:999px;border:1px solid rgba(52,211,153,.22);color:#34d399;background:rgba(52,211,153,.08);}
.nexo-pill.warn{border-color:rgba(251,146,60,.24);color:#fb923c;background:rgba(251,146,60,.08);}

/* Barras e gráficos fake premium */
.nexo-bar-row{display:grid;grid-template-columns:150px 1fr 34px;gap:12px;align-items:center;margin:12px 0;}
.nexo-bar{height:7px;background:rgba(148,163,184,.15);border-radius:999px;overflow:hidden;}
.nexo-bar span{display:block;height:100%;border-radius:999px;background:linear-gradient(90deg,#f87171,#fbbf24,#34d399);}
.nexo-donut{width:170px;height:170px;border-radius:50%;background:conic-gradient(#f87171 0 30%,#fb923c 30% 52%,#34d399 52% 100%);display:flex;align-items:center;justify-content:center;margin:auto;}
.nexo-donut-inner{width:104px;height:104px;border-radius:50%;background:#0f1b2d;display:flex;align-items:center;justify-content:center;flex-direction:column;color:white;font-weight:900;}
.nexo-radar{height:190px;border-radius:16px;background:radial-gradient(circle at center,rgba(99,102,241,.25) 0 10%,transparent 11% 22%,rgba(99,102,241,.18) 23% 24%,transparent 25% 48%,rgba(148,163,184,.16) 49% 50%,transparent 51%),linear-gradient(135deg,rgba(15,23,42,.2),rgba(59,130,246,.08));position:relative;display:flex;align-items:center;justify-content:center;}
.nexo-radar:after{content:"";width:112px;height:112px;background:linear-gradient(135deg,rgba(52,211,153,.55),rgba(139,92,246,.45));clip-path:polygon(50% 0%,80% 30%,72% 74%,38% 86%,14% 48%);filter:drop-shadow(0 0 22px rgba(99,102,241,.45));}
.nexo-action-row{display:grid;grid-template-columns:42px 1fr 18px;gap:12px;align-items:center;padding:14px;border-radius:14px;background:rgba(15,23,42,.46);border:1px solid rgba(148,163,184,.10);margin:10px 0;}
.nexo-alert-row{display:grid;grid-template-columns:30px 1fr auto;gap:10px;padding:11px 0;border-bottom:1px solid rgba(148,163,184,.10);align-items:center;}

/* Widgets e forms premium */
.stButton>button,.stDownloadButton>button{border-radius:15px!important;min-height:46px!important;font-weight:850!important;background:linear-gradient(90deg,#6366f1,#0ea5e9)!important;border:1px solid rgba(255,255,255,.10)!important;box-shadow:0 15px 38px rgba(37,99,235,.18)!important;}
.stButton>button:hover,.stDownloadButton>button:hover{transform:translateY(-1px)!important;}
.stTextArea textarea,.stNumberInput input,.stSelectbox div[data-baseweb="select"],.stMultiSelect div[data-baseweb="select"],.stDateInput input{border-radius:15px!important;background:rgba(15,23,42,.70)!important;border:1px solid rgba(148,163,184,.16)!important;color:#f8fafc!important;}
.stTabs [data-baseweb="tab-list"]{gap:8px;border-bottom:1px solid rgba(148,163,184,.12);}
.stTabs [data-baseweb="tab"]{background:rgba(15,23,42,.58);border-radius:14px 14px 0 0;padding:12px 18px;}
.stTabs [aria-selected="true"]{background:linear-gradient(90deg,rgba(99,102,241,.36),rgba(14,165,233,.22))!important;color:white!important;}
.card,.nexo-legacy-card{border:1px solid var(--nexo-border)!important;border-radius:18px!important;background:linear-gradient(180deg,rgba(20,32,50,.86),rgba(15,23,42,.70))!important;box-shadow:0 18px 55px rgba(0,0,0,.22)!important;}

@media(max-width:1250px){.nexo-grid-4{grid-template-columns:repeat(2,1fr)}.nexo-grid-2{grid-template-columns:1fr}.nexo-quick-grid{grid-template-columns:repeat(3,1fr)}.nexo-topbar{grid-template-columns:1fr}.nexo-grid-3{grid-template-columns:1fr}}
</style>
""", unsafe_allow_html=True)

# ==========================================================
# VISUAL HELPERS
# ==========================================================
def card_inicio(titulo, subtitulo=None):
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.markdown(f"### {titulo}")
    if subtitulo:
        st.markdown(f"<p class='small-muted'>{subtitulo}</p>", unsafe_allow_html=True)

def card_fim():
    st.markdown("</div>", unsafe_allow_html=True)

def initials(nome):
    return "".join([p[0] for p in (nome or "NEXO").split()[:2]]).upper()[:2]

def buscar_global(user_id, termo):
    termo = (termo or "").strip().lower()
    if not termo:
        return []
    resultados = []
    conn = db()
    c = conn.cursor()
    like = f"%{termo}%"
    try:
        c.execute("""
            SELECT id, nome, diagnostico FROM pacientes
            WHERE user_id=? AND (LOWER(nome) LIKE ? OR LOWER(COALESCE(diagnostico,'')) LIKE ? OR LOWER(COALESCE(responsaveis,'')) LIKE ?)
            ORDER BY nome LIMIT 8
        """, (user_id, like, like, like))
        for pid, nome, diag in c.fetchall():
            resultados.append({"tipo":"Paciente", "titulo":nome, "detalhe":diag or "Cadastro do paciente", "pagina":"Pacientes e evolução"})
        c.execute("""
            SELECT nome_arquivo, categoria, avaliacao_detectada FROM uploads
            WHERE user_id=? AND (LOWER(nome_arquivo) LIKE ? OR LOWER(COALESCE(categoria,'')) LIKE ? OR LOWER(COALESCE(avaliacao_detectada,'')) LIKE ?)
            ORDER BY id DESC LIMIT 8
        """, (user_id, like, like, like))
        for arq, cat, av in c.fetchall():
            resultados.append({"tipo":"Upload", "titulo":arq, "detalhe":av or cat or "Arquivo enviado", "pagina":"Upload inteligente"})
    finally:
        conn.close()
    paginas = {
        "dashboard":"Dashboard", "perfil":"Perfil profissional", "paciente":"Pacientes e evolução",
        "evolução":"Pacientes e evolução", "evolucao":"Pacientes e evolução", "upload":"Upload inteligente",
        "avaliação":"Avaliação e relatório", "avaliacao":"Avaliação e relatório", "relatório":"Avaliação e relatório",
        "relatorio":"Avaliação e relatório", "plano":"Planos e acesso", "admin":"Admin — Controle de Acessos"
    }
    for chave, pagina in paginas.items():
        if termo in chave or chave in termo:
            resultados.insert(0, {"tipo":"Tela", "titulo":pagina, "detalhe":"Abrir área do sistema", "pagina":pagina})
            break
    return resultados[:10]


def topbar(user):
    termo_key = "global_search_input"
    st.markdown(f"""
    <div class="nexo-topbar">
      <div class="nexo-greeting">
        <h2>Olá, {user.get('nome') or 'Profissional'} 👋</h2>
        <p>Organize sua prática, transforme dados clínicos em cuidado.</p>
      </div>
      <div></div>
      <div class="nexo-top-actions">
        <div class="nexo-icon-btn">{svg_icon('bell',22,'#9fb7d8')}<span class="nexo-notify">3</span></div>
        <div class="nexo-icon-btn">{svg_icon('help',22,'#9fb7d8')}</div>
      </div>
    </div>
    """, unsafe_allow_html=True)
    termo = st.text_input(
        "Busca global",
        placeholder="🔎  Buscar paciente, avaliação, relatório...",
        label_visibility="collapsed",
        key=termo_key
    )
    if termo:
        resultados = buscar_global(user["id"], termo)
        if resultados:
            st.markdown("<div class='nexo-card' style='margin:10px 0 16px;'>", unsafe_allow_html=True)
            st.caption(f"Resultados para: {termo}")
            for i, r in enumerate(resultados):
                cols = st.columns([0.18, 0.56, 0.26])
                with cols[0]:
                    st.markdown(f"**{r['tipo']}**")
                with cols[1]:
                    st.markdown(f"**{r['titulo']}**")
                    st.caption(r.get("detalhe") or "")
                with cols[2]:
                    if st.button("Abrir", key=f"search_open_{i}_{r['pagina']}", use_container_width=True):
                        go_to(r["pagina"])
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.info("Nenhum resultado encontrado para essa busca.")


def startup_header(user):
    # O dashboard 2.0 já possui saudação, cards e ações rápidas no corpo.
    return


def stat_cards(pacientes, avaliacoes=0, relatorios=0, uploads=0):
    horas = uso_mes(USER_ID, "atendimento") * 1.0
    st.markdown(f"""
    <div class="nexo-grid-4">
      <div class="nexo-card nexo-stat" style="background:linear-gradient(135deg,rgba(124,58,237,.28),rgba(15,23,42,.76));">
        <div class="nexo-stat-head"><div class="nexo-stat-icon">{svg_icon('users',24,'#a78bfa')}</div><div class="nexo-stat-label">Pacientes ativos</div></div>
        <div class="nexo-stat-num">{pacientes}</div><div class="nexo-positive">▲ 12% este mês</div>
      </div>
      <div class="nexo-card nexo-stat" style="background:linear-gradient(135deg,rgba(37,99,235,.24),rgba(15,23,42,.76));">
        <div class="nexo-stat-head"><div class="nexo-stat-icon">{svg_icon('clipboard',24,'#60a5fa')}</div><div class="nexo-stat-label">Atendimentos / avaliações</div></div>
        <div class="nexo-stat-num">{avaliacoes}</div><div class="nexo-positive">▲ registros do mês</div>
      </div>
      <div class="nexo-card nexo-stat" style="background:linear-gradient(135deg,rgba(16,185,129,.22),rgba(15,23,42,.76));">
        <div class="nexo-stat-head"><div class="nexo-stat-icon">{svg_icon('chart',24,'#34d399')}</div><div class="nexo-stat-label">Horas clínicas</div></div>
        <div class="nexo-stat-num">{int(horas)}h</div><div class="nexo-positive">▲ calculado por logs</div>
      </div>
      <div class="nexo-card nexo-stat" style="background:linear-gradient(135deg,rgba(251,146,60,.22),rgba(15,23,42,.76));">
        <div class="nexo-stat-head"><div class="nexo-stat-icon">{svg_icon('file',24,'#fb923c')}</div><div class="nexo-stat-label">Relatórios gerados</div></div>
        <div class="nexo-stat-num">{relatorios}</div><div class="nexo-positive">▲ produção mensal</div>
      </div>
    </div>
    """, unsafe_allow_html=True)


# ==========================================================
# LOGIN / PLANOS
# ==========================================================
def tela_login():
    col1, col2, col3 = st.columns([0.75, 1.1, 0.75])
    with col2:
        st.markdown(f"""
        <div class="card" style="margin-top:6vh;text-align:center;">
            <div style="display:flex;justify-content:center;margin-bottom:16px;">{brain_logo(72)}</div>
            <h1 style="font-size:3rem;margin:0;">NEXO</h1>
            <p class="small-muted">{APP_SLOGAN}</p>
        </div>
        """, unsafe_allow_html=True)
        tab_login, tab_cadastro = st.tabs(["Entrar", "Criar conta"])
        with tab_login:
            email = st.text_input("E-mail", key="login_email")
            senha = st.text_input("Senha", type="password", key="login_senha")
            if st.button("Entrar na plataforma"):
                user = autenticar(email, senha)
                if user:
                    st.session_state.user_id = user["id"]
                    st.rerun()
                else:
                    st.error("E-mail ou senha inválidos.")
        with tab_cadastro:
            nome = st.text_input("Nome completo", key="cad_nome")
            email_cad = st.text_input("E-mail profissional", key="cad_email")
            senha_cad = st.text_input("Criar senha", type="password", key="cad_senha")
            profissao = st.selectbox("Profissão", PROFISSOES, key="cad_profissao")
            registro = st.text_input("Registro / Conselho", key="cad_registro")
            formacao = st.text_area("Formação", key="cad_formacao", height=80)
            especializacoes = st.text_area("Especializações / abordagens", key="cad_espec", height=80)
            if st.button("Criar minha conta"):
                if not nome or not email_cad or not senha_cad:
                    st.warning("Preencha nome, e-mail e senha.")
                else:
                    ok, msg = criar_usuario(nome, email_cad, senha_cad, profissao, registro, formacao, especializacoes)
                    st.success(msg) if ok else st.error(msg)

def tela_escolher_plano(user):
    st.markdown(f"""
    <div class="card" style="text-align:center;">
      <div style="display:flex;justify-content:center;margin-bottom:10px;">{brain_logo(64)}</div>
      <h1>Escolha seu plano para acessar o NEXO</h1>
      <p class="small-muted">Sua conta foi criada. Para liberar o modo de produção, selecione um plano e conclua o pagamento.</p>
    </div>
    """, unsafe_allow_html=True)
    cols = st.columns(5)
    for i, (nome, p) in enumerate(PLANOS.items()):
        with cols[i]:
            st.markdown(f"""
            <div class="card price-card">
              <div>
                <span class="badge">{nome}</span>
                <div class="price">R$ {p['valor']:.0f}</div>
                <p class="small-muted">por mês</p>
                <p>{p['descricao']}</p>
                <p class="small-muted">Pacientes: {'Ilimitado' if p['pacientes'] > 9999 else p['pacientes']}<br>
                Relatórios/mês: {'Ilimitado' if p['relatorios_mes'] > 9999 else p['relatorios_mes']}<br>
                Uploads/mês: {'Ilimitado' if p['uploads_mes'] > 9999 else p['uploads_mes']}</p>
              </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Assinar {nome}", key=f"plano_{nome}"):
                ok, msg, url = criar_preferencia_mercado_pago(user["id"], nome)
                if ok and url:
                    st.success("Assinatura gerado. Abra o link abaixo para pagar.")
                    st.link_button("Ir para o Mercado Pago", url)
                else:
                    st.warning(msg)
                    st.info("Enquanto o token não estiver configurado, o admin pode liberar o plano manualmente em Usuários.")

# ==========================================================
# AVALIAÇÕES
# ==========================================================
DEMUCA_DOMINIOS = {
    "Comportamentos Restritivos": {"tipo": "restritivo", "maximo": 14, "itens": [("Estereotipias", 1), ("Agressividade", 1), ("Desinteresse", 1), ("Passividade", 1), ("Reclusão / Isolamento", 1), ("Resistência", 1), ("Pirraça", 1)]},
    "Interação Social e Cognição": {"tipo": "positivo", "maximo": 18, "itens": [("Contato visual", 1), ("Comunicação verbal", 1), ("Interação com objetos", 1), ("Interação com instrumentos", 1), ("Interação com terapeuta", 1), ("Interação com responsáveis", 1), ("Interação com pares", 1), ("Atenção", 1), ("Imitação", 1)]},
    "Percepção e Exploração Rítmica": {"tipo": "positivo", "maximo": 16, "itens": [("Pulso interno", 1), ("Regulação temporal", 1), ("Ritmo real", 2), ("Apoio", 2), ("Contrastes de andamento", 2)]},
    "Percepção e Exploração Sonora": {"tipo": "positivo", "maximo": 14, "itens": [("Som / silêncio", 1), ("Timbre", 1), ("Planos de altura", 1), ("Movimento sonoro", 1), ("Contrastes de intensidade", 1), ("Repetição de ideias", 1), ("Senso de conclusão", 1)]},
    "Exploração Vocal": {"tipo": "positivo", "maximo": 14, "itens": [("Vocalizações", 1), ("Balbucios", 1), ("Sílabas canônicas", 1), ("Imitação de canções", 2), ("Criação vocal", 2)]},
    "Movimentação Corporal com a Música": {"tipo": "positivo", "maximo": 14, "itens": [("Andar", 1), ("Correr", 1), ("Parar", 1), ("Gesticular", 1), ("Dançar", 1), ("Movimentar-se no lugar", 1), ("Pular", 1)]},
}

def classificar(valor, maximo):
    pc = (valor / maximo) * 100 if maximo else 0
    if pc < 40:
        return "baixo"
    if pc < 70:
        return "moderado"
    return "adequado"

def percentual(valor, maximo):
    return round((valor / maximo) * 100, 1) if maximo else 0

def campo(label, key):
    return st.number_input(label, min_value=0, max_value=5, value=0, step=1, key=key)

def calcular_demuca(respostas):
    totais = {}
    for dominio, itens in respostas.items():
        tipo = DEMUCA_DOMINIOS[dominio]["tipo"]
        total = 0
        for _, dados in itens.items():
            mapa = {"N": 2, "P": 1, "M": 0} if tipo == "restritivo" else {"N": 0, "P": 1, "M": 2}
            total += mapa[dados["resposta"]] * dados["peso"]
        totais[dominio] = total
    return totais

def totals_to_order(totais):
    ordem = ["Improvisação", "Recriação", "Composição", "Escuta Musical"]
    return {k: totais[k] for k in ordem if k in totais}

def identificar_prejuizos(nordoff_total, totais_iaps, totais_demuca):
    prejuizos = []
    if nordoff_total is not None and classificar(nordoff_total, 50) in ["baixo", "moderado"]:
        prejuizos.append({"area": "Comunicabilidade musical", "origem": "Nordoff-Robbins", "nivel": classificar(nordoff_total, 50), "objetivo": "Ampliar a comunicação musical funcional, favorecendo iniciativa, responsividade, reciprocidade e sustentação da interação sonoro-musical.", "habilidade": "comunicação musical, responsividade e vínculo terapêutico"})
    mapa_iaps = {
        "Improvisação": ("Estimular iniciativa sonora, espontaneidade musical e respostas intencionais em contexto improvisacional.", "iniciativa sonora e expressão espontânea"),
        "Recriação": ("Fortalecer memória musical, coordenação motora, imitação, seguimento de modelos e participação estruturada.", "memória musical, coordenação e participação estruturada"),
        "Composição": ("Favorecer criatividade, autoria, simbolização e organização de ideias musicais.", "criatividade, autoria e elaboração simbólica"),
        "Escuta Musical": ("Promover atenção auditiva, escuta ativa, resposta emocional e integração subjetiva da experiência sonora.", "escuta ativa, atenção e integração sonora"),
    }
    for area, valor in totals_to_order(totais_iaps).items():
        nivel = classificar(valor, 20)
        if nivel in ["baixo", "moderado"]:
            prejuizos.append({"area": area, "origem": "IAPS", "nivel": nivel, "objetivo": mapa_iaps[area][0], "habilidade": mapa_iaps[area][1]})
    mapa_demuca = {
        "Comportamentos Restritivos": ("Reduzir interferência de comportamentos restritivos, favorecendo co-regulação, previsibilidade e engajamento funcional.", "regulação e redução de interferências comportamentais"),
        "Interação Social e Cognição": ("Ampliar contato visual, atenção compartilhada, imitação, comunicação e interação social mediada pela música.", "interação social, atenção compartilhada e cognição"),
        "Percepção e Exploração Rítmica": ("Desenvolver pulso interno, regulação temporal, apoio rítmico e percepção de contrastes de andamento.", "percepção rítmica e organização temporal"),
        "Percepção e Exploração Sonora": ("Estimular discriminação sonora, timbre, altura, intensidade e organização musical.", "percepção sonora e discriminação auditiva"),
        "Exploração Vocal": ("Favorecer vocalizações, balbucios, imitação de canções e criação vocal.", "expressão vocal e comunicação vocal"),
        "Movimentação Corporal com a Música": ("Ampliar organização motora, expressão corporal e integração música-movimento.", "expressão corporal e coordenação motora"),
    }
    for dominio, valor in totais_demuca.items():
        nivel = classificar(valor, DEMUCA_DOMINIOS[dominio]["maximo"])
        if nivel in ["baixo", "moderado"]:
            prejuizos.append({"area": dominio, "origem": "DEMUCA", "nivel": nivel, "objetivo": mapa_demuca[dominio][0], "habilidade": mapa_demuca[dominio][1]})
    if not prejuizos:
        prejuizos.append({"area": "Manutenção e ampliação de repertório terapêutico", "origem": "Síntese clínica", "nivel": "adequado", "objetivo": "Aprofundar recursos já estabelecidos, ampliando autonomia, flexibilidade e expressão funcional.", "habilidade": "autonomia, flexibilidade e elaboração expressiva"})
    return prejuizos

def gerar_grafico_barras(titulo, dados, maximos=None):
    categorias, valores = list(dados.keys()), list(dados.values())
    limite = max(maximos.values()) if maximos else (max(valores) if valores else 5)
    fig, ax = plt.subplots(figsize=(11, 5.2))
    fig.patch.set_facecolor("#020617")
    ax.set_facecolor("#0f172a")
    ax.bar(categorias, valores)
    ax.set_title(titulo, color="white", fontsize=14, fontweight="bold")
    ax.set_ylim(0, limite)
    ax.tick_params(axis="x", rotation=25, colors="white")
    ax.tick_params(axis="y", colors="white")
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    ax.spines["bottom"].set_color("#475569")
    ax.spines["left"].set_color("#475569")
    plt.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    buffer.seek(0)
    return buffer

def gerar_grafico_radar(totais_iaps, nordoff_total=None, totais_demuca=None):
    cats, vals = [], []
    if nordoff_total is not None:
        cats.append("Nordoff")
        vals.append(percentual(nordoff_total, 50))
    if totais_iaps:
        cats += ["Improvisação", "Recriação", "Composição", "Escuta"]
        vals += [percentual(totais_iaps.get("Improvisação", 0),20), percentual(totais_iaps.get("Recriação", 0),20), percentual(totais_iaps.get("Composição", 0),20), percentual(totais_iaps.get("Escuta Musical", 0),20)]
    if totais_demuca:
        for d, v in totais_demuca.items():
            cats.append(d[:16])
            vals.append(percentual(v, DEMUCA_DOMINIOS[d]["maximo"]))
    while len(cats) < 3:
        cats.append("")
        vals.append(0)
    angles = [n / float(len(cats)) * 2 * math.pi for n in range(len(cats))]
    values = vals + vals[:1]
    angles += angles[:1]
    fig = plt.figure(figsize=(8, 8))
    fig.patch.set_facecolor("#020617")
    ax = plt.subplot(111, polar=True)
    ax.set_facecolor("#0f172a")
    ax.plot(angles, values, linewidth=2)
    ax.fill(angles, values, alpha=.25)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(cats, color="white")
    ax.set_ylim(0, 100)
    ax.grid(color="#334155")
    ax.set_title("Perfil Geral (%)", color="white", pad=20, fontsize=14, fontweight="bold")
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    buffer.seek(0)
    return buffer

def interpretar_iaps(totais):
    out = ""
    for area, valor in totals_to_order(totais).items():
        out += f"{area}: desempenho {classificar(valor, 20)}.\n"
    return out

def interpretar_demuca(totais):
    return "".join([f"{d}: {v}/{DEMUCA_DOMINIOS[d]['maximo']} — {classificar(v, DEMUCA_DOMINIOS[d]['maximo'])}.\n" for d, v in totais.items()])

def gerar_resumo(nome, idade, diagnostico, nordoff_total, nordoff, iaps, demuca, profissao="Musicoterapeuta"):
    perfil = perfil_profissional(profissao)
    txt = (
        f"{nome or 'O paciente'}, {idade} anos, com diagnóstico de {diagnostico or 'não informado'}, "
        f"foi avaliado(a) a partir de parâmetros coerentes com a área de {profissao}. "
        f"A leitura clínica prioriza {perfil['foco']}.\n\n"
    )
    if nordoff_total is not None:
        txt += f"Na Nordoff-Robbins, o desempenho geral foi {classificar(nordoff_total, 50)}. Este dado deve ser compreendido em articulação com os indicadores funcionais observados no contexto terapêutico.\n\n"
    if iaps:
        txt += "Nos IAPS:\n" + interpretar_iaps(iaps) + "\n"
    if demuca:
        txt += "Na DEMUCA:\n" + interpretar_demuca(demuca) + "\n"
    txt += (
        f"A integração dos dados sugere planejamento individualizado, com objetivos mensuráveis, acompanhamento evolutivo e intervenções alinhadas a "
        f"{perfil['prescricao']}"
    )
    return txt

def gerar_gas(prejuizos, profissao="Musicoterapeuta"):
    perfil = perfil_profissional(profissao)
    metas = []
    fonte = prejuizos[:3]
    if not fonte:
        fonte = [{"habilidade": a, "objetivo": o} for a, o in zip(perfil["areas"][:3], perfil["objetivos_base"][:3])]
    for i, p in enumerate(fonte, 1):
        h = p.get("habilidade") or p.get("area") or perfil["areas"][0]
        obj = p.get("objetivo") or perfil["objetivos_base"][0]
        metas.append({
            "meta": f"META {i:02d}: {h.capitalize()}",
            "-2": f"Não apresenta {h} de forma funcional no contexto observado.",
            "-1": f"Apresenta {h} de forma inicial, inconsistente ou dependente de alto suporte.",
            "0": obj,
            "+1": f"Apresenta {h} com maior consistência, menor suporte e melhor adaptação ao contexto.",
            "+2": f"Apresenta {h} de forma funcional, consistente e com sinais de generalização."
        })
    return metas

def gerar_prescricao(prejuizos, profissao="Musicoterapeuta"):
    perfil = perfil_profissional(profissao)
    principais = ", ".join([p["area"] for p in prejuizos[:4]]) or ", ".join(perfil["areas"][:4])
    return (
        f"Recomenda-se acompanhamento em {profissao}, com plano individualizado e objetivos mensuráveis, priorizando: {principais}. "
        f"A prescrição deve contemplar {perfil['prescricao']}"
    )

def gerar_conduta(prejuizos, profissao="Musicoterapeuta"):
    perfil = perfil_profissional(profissao)
    txt = (
        f"A partir da análise integrada, recomenda-se continuidade do acompanhamento com foco em {perfil['foco']}. "
        f"O plano deve ser revisado periodicamente por indicadores funcionais e pela evolução observada em sessão.\n\nFocos identificados:\n"
    )
    for p in prejuizos[:5]:
        txt += f"- {p['area']} ({p['origem']}): {p['objetivo']}\n"
    txt += "\nSugestões técnicas da área: " + ", ".join(perfil["termos"][:5]) + "."
    return txt

# ==========================================================
# WORD / PDF
# ==========================================================
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)

def add_section(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)

def add_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text or "Não informado.")
    r.font.size = Pt(11)

def add_bullet(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"• {text}")
    r.font.size = Pt(11)

def inserir_logo_doc(doc, logo_path, posicao):
    if not logo_path or not os.path.exists(logo_path):
        return
    try:
        if posicao == "Topo centralizado":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.25))
        elif posicao == "Topo direito":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.25))
        elif posicao == "Rodapé":
            section = doc.sections[0]
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.0))
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.25))
    except Exception:
        pass

def criar_word(ctx, dados, user):
    doc = Document()
    inserir_logo_doc(doc, user.get("logo_path"), user.get("logo_posicao") or "Topo esquerdo")
    add_title(doc, perfil_profissional(ctx.get("profissao_rel") or user.get("profissao") or "Outro")["rotulo_relatorio"].upper() + " — NEXO")
    add_section(doc, "IDENTIFICAÇÃO")
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    campos = [
        ("Paciente", ctx.get("nome")), ("Idade", str(ctx.get("idade"))), ("Diagnóstico", ctx.get("diagnostico")),
        ("Profissional", ctx.get("terapeuta")), ("Profissão", ctx.get("profissao_rel")), ("Registro", ctx.get("registro"))
    ]
    for k, v in campos:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v or ""
    justificativa_texto = "\n\n".join(dados.get("justificativas", []))
    for sec, val in [
        ("HISTÓRIA CLÍNICA", ctx.get("historia")),
        ("QUEIXA", ctx.get("queixa")),
        ("OBSERVAÇÕES", ctx.get("observacoes")),
        ("JUSTIFICATIVA TÉCNICA DAS AVALIAÇÕES", justificativa_texto or "Não informado."),
        ("RESUMO GERAL", dados["resumo"]),
        ("PRESCRIÇÃO", dados["prescricao"]),
        ("CONDUTA", dados["conduta"])
    ]:
        add_section(doc, sec)
        add_text(doc, val)
    add_section(doc, "OBJETIVOS")
    for p in dados["prejuizos"]:
        add_bullet(doc, p["objetivo"])
    if dados.get("gas"):
        add_section(doc, dados.get("nome_escala_objetivos", "Escala de Objetivos Mensuráveis"))
        for m in dados["gas"]:
            add_text(doc, str(m))
    for chave, titulo in [("grafico_nordoff", "Gráfico - Nordoff-Robbins"), ("grafico_iaps", "Gráfico - IAPS"), ("grafico_demuca", "Gráfico - DEMUCA")]:
        if dados.get(chave):
            add_section(doc, titulo)
            doc.add_picture(dados[chave], width=Inches(5.7))
    if dados.get("avaliacoes_personalizadas"):
        add_section(doc, "AVALIAÇÕES PERSONALIZADAS E GRÁFICOS")
        for av in dados.get("avaliacoes_personalizadas", []):
            add_text(doc, f"{av.get('modelo_nome')} — {av.get('avaliacao_nome')}\n{av.get('interpretacao') or ''}")
            fig_av = gerar_grafico_avaliacao_modelo(av.get('modelo_nome') or av.get('avaliacao_nome'), av.get('grafico_spec'), av.get('grafico_valores'))
            buf_av = fig_to_buffer(fig_av)
            if fig_av:
                plt.close(fig_av)
            if buf_av:
                doc.add_picture(buf_av, width=Inches(5.7))
    if dados.get("grafico_radar"):
        add_section(doc, "Gráfico - Perfil Geral")
        doc.add_picture(dados["grafico_radar"], width=Inches(5.7))
    doc.add_paragraph("\n\n__________________________________")
    doc.add_paragraph(ctx.get("terapeuta") or "")
    doc.add_paragraph(ctx.get("profissao_rel") or "")
    doc.add_paragraph(ctx.get("registro") or "")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def criar_pdf_simples(ctx, dados, user):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    logo_path = user.get("logo_path")
    if logo_path and os.path.exists(logo_path) and user.get("logo_posicao") != "Rodapé":
        story.append(Image(logo_path, width=1.2*inch, height=1.2*inch))
        story.append(Spacer(1, 10))

    story.append(Paragraph(f"<b>{perfil_profissional(ctx.get('profissao_rel') or user.get('profissao') or 'Outro')['rotulo_relatorio'].upper()} — NEXO</b>", styles["Title"]))
    story.append(Spacer(1, 12))
    for titulo, texto in [
        ("Identificação", f"Paciente: {ctx.get('nome') or ''}<br/>Idade: {ctx.get('idade') or ''}<br/>Diagnóstico: {ctx.get('diagnostico') or ''}<br/>Profissional: {ctx.get('terapeuta') or ''}"),
        ("História clínica", ctx.get("historia") or "Não informado."),
        ("Queixa", ctx.get("queixa") or "Não informado."),
        ("Justificativa técnica das avaliações", "<br/><br/>".join(dados.get("justificativas", [])) or "Não informado."),
        ("Resumo geral", dados["resumo"]),
        ("Prescrição", dados["prescricao"]),
        ("Conduta", dados["conduta"]),
    ]:
        story.append(Paragraph(f"<b>{titulo}</b>", styles["Heading2"]))
        story.append(Paragraph(str(texto).replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 10))

    # Gráficos das avaliações estruturadas e personalizadas
    for chave, titulo in [("grafico_nordoff", "Gráfico - Nordoff-Robbins"), ("grafico_iaps", "Gráfico - IAPS"), ("grafico_demuca", "Gráfico - DEMUCA"), ("grafico_radar", "Gráfico - Perfil Geral")]:
        if dados.get(chave):
            story.append(Paragraph(f"<b>{titulo}</b>", styles["Heading2"]))
            story.append(Image(dados[chave], width=5.8*inch, height=3.2*inch))
            story.append(Spacer(1, 10))
    if dados.get("avaliacoes_personalizadas"):
        story.append(Paragraph("<b>Avaliações personalizadas e gráficos</b>", styles["Heading2"]))
        for av in dados.get("avaliacoes_personalizadas", []):
            story.append(Paragraph(f"<b>{av.get('modelo_nome')} — {av.get('avaliacao_nome')}</b>", styles["Heading3"]))
            story.append(Paragraph(str(av.get('interpretacao') or '').replace("\n", "<br/>"), styles["BodyText"]))
            fig_av = gerar_grafico_avaliacao_modelo(av.get('modelo_nome') or av.get('avaliacao_nome'), av.get('grafico_spec'), av.get('grafico_valores'))
            buf_av = fig_to_buffer(fig_av)
            if fig_av:
                plt.close(fig_av)
            if buf_av:
                story.append(Image(buf_av, width=5.8*inch, height=3.2*inch))
                story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer



def nexo_count_table(user_id, table, where_extra=""):
    conn = db(); c = conn.cursor()
    try:
        sql = f"SELECT COUNT(*) FROM {table} WHERE user_id=? {where_extra}"
        c.execute(sql, (user_id,))
        return int(c.fetchone()[0] or 0)
    except Exception:
        return 0
    finally:
        conn.close()

def nexo_agenda_hoje(user_id, limite=5):
    hoje = date.today().strftime("%Y-%m-%d")
    conn = db(); c = conn.cursor()
    try:
        c.execute("""
            SELECT e.hora_inicio, COALESCE(p.nome,e.titulo), e.titulo, e.status
            FROM agenda_eventos e
            LEFT JOIN pacientes p ON p.id=e.paciente_id
            WHERE e.user_id=? AND e.data=?
            ORDER BY e.hora_inicio LIMIT ?
        """, (user_id, hoje, limite))
        return c.fetchall()
    except Exception:
        return []
    finally:
        conn.close()

def nexo_ultimas_atividades(user_id, limite=4):
    atividades = []
    conn = db(); c = conn.cursor()
    try:
        c.execute("SELECT avaliacao_nome, criado_em FROM avaliacao_respostas WHERE user_id=? ORDER BY id DESC LIMIT ?", (user_id, limite))
        for nome, criado in c.fetchall():
            atividades.append(("Avaliação", f"Avaliação {nome or ''} preenchida", criado or ""))
        c.execute("SELECT nome_arquivo, criado_em FROM uploads WHERE user_id=? ORDER BY id DESC LIMIT ?", (user_id, limite))
        for nome, criado in c.fetchall():
            atividades.append(("Upload", f"Upload {nome or ''}", criado or ""))
    except Exception:
        pass
    finally:
        conn.close()
    atividades = sorted(atividades, key=lambda x: x[2] or "", reverse=True)[:limite]
    return atividades

def nexo_quick_action_card(icon, label, key, page):
    if st.button(f"{icon}\n\n{label}", key=key, use_container_width=True):
        go_to(page)

def nexo_render_quick_actions():
    st.markdown("<div class='nexo-card'><h3>Acesso rápido</h3>", unsafe_allow_html=True)
    cols = st.columns(6)
    actions = [
        ("📋", "Nova avaliação", "quick_avaliacao", "Avaliação e relatório"),
        ("📅", "Novo atendimento", "quick_atendimento", "Agenda e calendário"),
        ("📝", "Nova evolução", "quick_evolucao", "Pacientes e evolução"),
        ("📄", "Gerar relatório", "quick_relatorio", "Avaliação e relatório"),
        ("🎯", "Objetivos", "quick_objetivos", "Central clínica inteligente"),
        ("🧩", "Intervenções", "quick_intervencoes", "Central clínica inteligente"),
    ]
    for col, action in zip(cols, actions):
        with col:
            nexo_quick_action_card(*action)
    st.markdown("</div>", unsafe_allow_html=True)


# ==========================================================
# AGENDA DO TERAPEUTA — DIA, MÊS E ANO
# ==========================================================
def agenda_status_opcoes():
    return ["previsto", "confirmado", "realizado", "evolucao_ok", "faltou", "cancelado", "reagendado"]

def listar_eventos_agenda(user_id, inicio, fim):
    conn = db(); c = conn.cursor()
    try:
        c.execute("""
            SELECT e.id, e.data, e.hora_inicio, e.hora_fim, COALESCE(p.nome, e.titulo),
                   e.titulo, e.modalidade, e.status, e.observacao, e.paciente_id
            FROM agenda_eventos e
            LEFT JOIN pacientes p ON p.id=e.paciente_id
            WHERE e.user_id=? AND e.data BETWEEN ? AND ?
            ORDER BY e.data, e.hora_inicio, e.id
        """, (user_id, inicio, fim))
        return c.fetchall()
    except Exception:
        return []
    finally:
        conn.close()

def salvar_evento_agenda(user_id, paciente_id, data_ev, hora_inicio, hora_fim, titulo, modalidade, status, observacao):
    conn = db(); c = conn.cursor()
    try:
        c.execute("""
            INSERT INTO agenda_eventos
            (user_id, paciente_id, data, hora_inicio, hora_fim, titulo, modalidade, status, origem, observacao, criado_em, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            user_id, paciente_id, data_ev, hora_inicio, hora_fim, titulo, modalidade, status,
            "agenda_terapeuta", observacao, datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")
        ))
        conn.commit()
        return True, "Atendimento criado na agenda."
    except Exception as e:
        return False, f"Erro ao criar atendimento: {e}"
    finally:
        conn.close()

def atualizar_status_evento(event_id, user_id, status):
    conn = db(); c = conn.cursor()
    try:
        c.execute("UPDATE agenda_eventos SET status=?, atualizado_em=? WHERE id=? AND user_id=?", (status, datetime.now().isoformat(timespec="seconds"), event_id, user_id))
        conn.commit(); return True
    except Exception:
        return False
    finally:
        conn.close()

def excluir_evento_agenda(event_id, user_id):
    conn = db(); c = conn.cursor()
    try:
        c.execute("DELETE FROM agenda_eventos WHERE id=? AND user_id=?", (event_id, user_id))
        conn.commit(); return True
    except Exception:
        return False
    finally:
        conn.close()

def _agenda_badge(status):
    s = str(status or "previsto").lower()
    if s in ["realizado", "evolucao_ok", "confirmado"]:
        return "🟢"
    if s in ["faltou", "cancelado"]:
        return "🔴"
    if s in ["reagendado", "pendente"]:
        return "🟠"
    return "🔵"

def render_calendario_mes(user_id, ano, mes):
    primeiro = date(ano, mes, 1)
    ultimo = date(ano, mes, calendar.monthrange(ano, mes)[1])
    eventos = listar_eventos_agenda(user_id, primeiro.strftime("%Y-%m-%d"), ultimo.strftime("%Y-%m-%d"))
    por_dia = {}
    for ev in eventos:
        por_dia.setdefault(ev[1], []).append(ev)

    st.markdown(f"### {calendar.month_name[mes].capitalize()} / {ano}")
    semanas = calendar.Calendar(firstweekday=0).monthdatescalendar(ano, mes)
    nomes = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]
    cols = st.columns(7)
    for col, nome in zip(cols, nomes):
        col.markdown(f"<div class='nexo-muted' style='text-align:center;font-weight:900'>{nome}</div>", unsafe_allow_html=True)
    for semana in semanas:
        cols = st.columns(7)
        for col, dia in zip(cols, semana):
            data_key = dia.strftime("%Y-%m-%d")
            eventos_dia = por_dia.get(data_key, [])
            opacity = "0.42" if dia.month != mes else "1"
            border = "rgba(96,165,250,.30)" if data_key == date.today().strftime("%Y-%m-%d") else "rgba(148,163,184,.12)"
            resumo = "".join([f"<div style='font-size:.74rem;margin-top:5px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;'>{_agenda_badge(ev[7])} {ev[2] or ''} {ev[4] or ev[5] or ''}</div>" for ev in eventos_dia[:3]])
            extra = f"<div class='nexo-muted' style='font-size:.72rem;margin-top:4px;'>+{len(eventos_dia)-3} evento(s)</div>" if len(eventos_dia) > 3 else ""
            col.markdown(f"""
            <div style='min-height:116px;border:1px solid {border};border-radius:16px;background:rgba(15,23,42,.52);padding:10px;opacity:{opacity};margin-bottom:8px;'>
                <div style='font-weight:950;color:#f8fafc;'>{dia.day}</div>
                {resumo}{extra}
            </div>
            """, unsafe_allow_html=True)

def render_tabela_eventos(eventos, prefixo):
    if not eventos:
        st.info("Nenhum atendimento encontrado neste período.")
        return
    for ev in eventos:
        ev_id, data_ev, hi, hf, paciente, titulo, modalidade, status, obs, paciente_id = ev
        with st.expander(f"{_agenda_badge(status)} {data_ev} • {hi or '--:--'}–{hf or '--:--'} • {paciente or titulo or 'Atendimento'}", expanded=False):
            c1, c2, c3 = st.columns([0.33, 0.34, 0.33])
            c1.markdown(f"**Paciente:** {paciente or 'Não vinculado'}")
            c2.markdown(f"**Modalidade:** {modalidade or 'Presencial'}")
            c3.markdown(f"**Status atual:** `{status or 'previsto'}`")
            if obs:
                st.markdown(f"**Observação:** {obs}")
            novo_status = st.selectbox("Atualizar status", agenda_status_opcoes(), index=agenda_status_opcoes().index(status) if status in agenda_status_opcoes() else 0, key=f"{prefixo}_status_{ev_id}")
            b1, b2, b3 = st.columns(3)
            if b1.button("Salvar status", key=f"{prefixo}_save_status_{ev_id}"):
                if atualizar_status_evento(ev_id, USER_ID, novo_status):
                    st.success("Status atualizado."); st.rerun()
                else:
                    st.error("Não foi possível atualizar.")
            if b2.button("Criar evolução", key=f"{prefixo}_evolucao_{ev_id}"):
                st.session_state["agenda_evento_para_evolucao"] = ev_id
                go_to("Pacientes e evolução")
            if b3.button("Remover", key=f"{prefixo}_delete_{ev_id}"):
                if excluir_evento_agenda(ev_id, USER_ID):
                    st.warning("Atendimento removido."); st.rerun()
                else:
                    st.error("Não foi possível remover.")

def tela_agenda_terapeuta():
    card_inicio("Agenda e calendário", "Visualize e organize seus atendimentos por dia, mês e ano.")
    tab_dia, tab_mes, tab_ano, tab_novo = st.tabs(["Hoje / Dia", "Mês", "Ano", "Novo atendimento"])

    with tab_dia:
        data_sel = st.date_input("Escolha o dia", value=date.today(), key="agenda_dia_data")
        eventos = listar_eventos_agenda(USER_ID, data_sel.strftime("%Y-%m-%d"), data_sel.strftime("%Y-%m-%d"))
        st.markdown(f"### Atendimentos de {data_sel.strftime('%d/%m/%Y')}")
        render_tabela_eventos(eventos, "dia")

    with tab_mes:
        c1, c2 = st.columns([0.25, 0.25])
        mes = c1.selectbox("Mês", list(range(1, 13)), index=date.today().month-1, format_func=lambda m: calendar.month_name[m].capitalize(), key="agenda_mes")
        ano = c2.number_input("Ano", min_value=2020, max_value=2100, value=date.today().year, step=1, key="agenda_ano_mes")
        render_calendario_mes(USER_ID, int(ano), int(mes))
        inicio = date(int(ano), int(mes), 1)
        fim = date(int(ano), int(mes), calendar.monthrange(int(ano), int(mes))[1])
        st.markdown("### Lista do mês")
        render_tabela_eventos(listar_eventos_agenda(USER_ID, inicio.strftime("%Y-%m-%d"), fim.strftime("%Y-%m-%d")), "mes")

    with tab_ano:
        ano_visao = st.number_input("Ano da visão anual", min_value=2020, max_value=2100, value=date.today().year, step=1, key="agenda_ano_visao")
        eventos = listar_eventos_agenda(USER_ID, f"{int(ano_visao)}-01-01", f"{int(ano_visao)}-12-31")
        por_mes = {m: 0 for m in range(1, 13)}
        realizados = {m: 0 for m in range(1, 13)}
        for ev in eventos:
            try:
                m = int(str(ev[1]).split("-")[1]); por_mes[m] += 1
                if str(ev[7]).lower() in ["realizado", "evolucao_ok"]:
                    realizados[m] += 1
            except Exception:
                pass
        cols = st.columns(4)
        for idx, m in enumerate(range(1, 13)):
            with cols[idx % 4]:
                st.markdown(f"""
                <div class='nexo-card' style='padding:14px;margin-bottom:12px;'>
                  <h3>{calendar.month_name[m].capitalize()}</h3>
                  <div style='font-size:1.7rem;font-weight:950;color:#fff;'>{por_mes[m]}</div>
                  <div class='nexo-muted'>{realizados[m]} realizados</div>
                </div>
                """, unsafe_allow_html=True)
        st.markdown("### Todos os atendimentos do ano")
        render_tabela_eventos(eventos, "ano")

    with tab_novo:
        pacientes = listar_pacientes(USER_ID)
        mapa_pac = {"Sem paciente vinculado": None}
        for p in pacientes:
            mapa_pac[f"#{p[0]} — {p[1]}"] = p[0]
        c1, c2 = st.columns(2)
        with c1:
            paciente_label = st.selectbox("Paciente", list(mapa_pac.keys()), key="agenda_novo_paciente")
            data_ev = st.date_input("Data", value=date.today(), key="agenda_novo_data")
            hora_inicio = st.text_input("Hora inicial", value="08:00", key="agenda_novo_hi")
            hora_fim = st.text_input("Hora final", value="08:50", key="agenda_novo_hf")
        with c2:
            titulo = st.text_input("Título", value="Atendimento", key="agenda_novo_titulo")
            modalidade = st.selectbox("Modalidade", ["Presencial", "Online", "Domiciliar", "Híbrido"], key="agenda_novo_modalidade")
            status = st.selectbox("Status", agenda_status_opcoes(), key="agenda_novo_status")
            repetir = st.selectbox("Repetição", ["Não repetir", "Semanal até o fim do mês", "Semanal até o fim do ano"], key="agenda_novo_repetir")
        obs = st.text_area("Observação", height=90, key="agenda_novo_obs")
        if st.button("Salvar atendimento na agenda", key="agenda_novo_salvar", use_container_width=True):
            datas = [data_ev]
            if repetir != "Não repetir":
                limite = date(data_ev.year, data_ev.month, calendar.monthrange(data_ev.year, data_ev.month)[1]) if repetir == "Semanal até o fim do mês" else date(data_ev.year, 12, 31)
                datas = []
                d = data_ev
                while d <= limite:
                    datas.append(d)
                    d = d + timedelta(days=7)
            ok_count = 0
            last_msg = ""
            for d in datas[:60]:
                ok, msg = salvar_evento_agenda(USER_ID, mapa_pac[paciente_label], d.strftime("%Y-%m-%d"), hora_inicio, hora_fim, titulo, modalidade, status, obs)
                ok_count += 1 if ok else 0
                last_msg = msg
            if ok_count:
                st.success(f"{ok_count} atendimento(s) salvo(s) na agenda.")
                st.rerun()
            else:
                st.error(last_msg or "Não foi possível salvar.")

# ==========================================================
# TELAS
# ==========================================================
def tela_dashboard():
    pacientes = contar_pacientes(USER_ID)
    rel = uso_mes(USER_ID, "relatorio")
    up = uso_mes(USER_ID, "upload")
    avaliacoes = uso_mes(USER_ID, "avaliacao") or nexo_count_table(USER_ID, "avaliacao_respostas")

    stat_cards(pacientes, avaliacoes=avaliacoes, relatorios=rel, uploads=up)
    nexo_render_quick_actions()

    agenda = nexo_agenda_hoje(USER_ID)
    atividades = nexo_ultimas_atividades(USER_ID)

    col_main, col_side = st.columns([0.73, 0.27], gap="medium")
    with col_main:
        st.markdown("""
        <div class="nexo-card"><h3>Visão clínica inteligente <span class="nexo-muted">ⓘ</span></h3>
          <div class="nexo-grid-3">
            <div>
              <div class="nexo-muted" style="font-weight:850;margin-bottom:12px;">Domínios mais comprometidos</div>
              <div class="nexo-bar-row"><span>Regulação emocional</span><div class="nexo-bar"><span style="width:42%"></span></div><b>2.1</b></div>
              <div class="nexo-bar-row"><span>Interação social</span><div class="nexo-bar"><span style="width:54%"></span></div><b>2.7</b></div>
              <div class="nexo-bar-row"><span>Comunicação</span><div class="nexo-bar"><span style="width:64%"></span></div><b>3.2</b></div>
              <div class="nexo-bar-row"><span>Flexibilidade</span><div class="nexo-bar"><span style="width:72%"></span></div><b>3.6</b></div>
              <div class="nexo-muted" style="font-size:.78rem;margin-top:12px;">Escala: 1 baixo a 5 alto</div>
            </div>
            <div>
              <div class="nexo-muted" style="font-weight:850;margin-bottom:12px;">Evolução média por domínio</div>
              <div class="nexo-radar"></div>
              <div class="nexo-muted" style="font-size:.78rem;text-align:center;margin-top:10px;">Últimos 3 meses</div>
            </div>
            <div>
              <div class="nexo-muted" style="font-weight:850;margin-bottom:12px;">Distribuição de níveis</div>
              <div class="nexo-donut"><div class="nexo-donut-inner"><span style="font-size:2rem;">{}</span><span class="nexo-muted">pacientes</span></div></div>
              <div class="nexo-muted" style="font-size:.82rem;text-align:center;margin-top:12px;">Crítico • Moderado • Adequado</div>
            </div>
          </div>
        </div>
        """.format(pacientes), unsafe_allow_html=True)

        c1, c2 = st.columns([0.42,0.58], gap="medium")
        with c1:
            st.markdown("<div class='nexo-card'><h3>Pacientes com atenção</h3>", unsafe_allow_html=True)
            alertas = [
                ("🟠", "Reavaliação pendente", "sem reavaliação há 120 dias"),
                ("🟡", "Faltas recorrentes", "acompanhar adesão"),
                ("🔴", "Evolução abaixo do esperado", "revisar objetivos"),
            ]
            for icon, nome, motivo in alertas:
                st.markdown(f"<div class='nexo-alert-row'><div>{icon}</div><div><b>{nome}</b><br><span class='nexo-muted'>{motivo}</span></div><span>›</span></div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
        with c2:
            st.markdown("<div class='nexo-card'><h3>Últimas atividades</h3>", unsafe_allow_html=True)
            if atividades:
                for tipo, titulo, criado in atividades:
                    st.markdown(f"<div class='nexo-action-row'><div>📌</div><div><b>{titulo}</b><br><span class='nexo-muted'>{tipo} • {str(criado)[:16]}</span></div><span>›</span></div>", unsafe_allow_html=True)
            else:
                st.markdown("<p class='nexo-muted'>Nenhuma atividade recente registrada ainda.</p>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

    with col_side:
        st.markdown("<div class='nexo-card'><h3>Agenda de hoje</h3>", unsafe_allow_html=True)
        if agenda:
            for hora, paciente, titulo, status in agenda:
                warn = "warn" if str(status).lower() in ["pendente","cancelado","faltou"] else ""
                pill = "Pendente" if warn else "Confirmado"
                st.markdown(f"<div class='nexo-agenda-row'><b>{hora or '--:--'}</b><span class='nexo-dot {warn}'></span><div><b>{paciente}</b><br><span class='nexo-muted'>{titulo or 'Atendimento'}</span></div><span class='nexo-pill {warn}'>{pill}</span></div>", unsafe_allow_html=True)
        else:
            st.markdown("<p class='nexo-muted'>Nenhum atendimento agendado para hoje.</p>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='nexo-card'><h3>Ações rápidas</h3>", unsafe_allow_html=True)
        if st.button("📤 Upload de avaliação", key="dash_upload_avaliacao", use_container_width=True):
            go_to("Upload inteligente")
        if st.button("📅 Abrir agenda", key="dash_abrir_agenda", use_container_width=True):
            go_to("Agenda e calendário")
        if st.button("📚 Biblioteca de intervenções", key="dash_biblioteca_intervencoes", use_container_width=True):
            go_to("Central clínica inteligente")
        if st.button("📊 Comparar evoluções", key="dash_comparar_evolucoes", use_container_width=True):
            go_to("Central clínica inteligente")
        if st.button("🎯 Gerar objetivos", key="dash_gerar_objetivos", use_container_width=True):
            go_to("Central clínica inteligente")
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("""
        <div class="nexo-card" style="background:linear-gradient(135deg,rgba(124,58,237,.34),rgba(15,23,42,.72));">
          <h3>✦ Insights automáticos</h3>
          <p class="nexo-muted">A área de Regulação emocional aparece como a mais desafiadora entre seus pacientes.</p>
        </div>
        """, unsafe_allow_html=True)


def tela_perfil():
    card_inicio("Perfil profissional e logo", "Dados usados na assinatura, Word e PDF.")
    c1, c2 = st.columns(2)
    with c1:
        prof = st.selectbox("Profissão", PROFISSOES, index=PROFISSOES.index(USER["profissao"]) if USER.get("profissao") in PROFISSOES else 0)
        registro = st.text_input("Registro profissional", value=USER.get("registro") or "")
        assinatura = st.text_input("Assinatura", value=USER.get("assinatura") or USER["nome"])
        logo_pos = st.selectbox("Posição da logo no relatório", ["Topo esquerdo", "Topo centralizado", "Topo direito", "Rodapé"],
                                index=["Topo esquerdo", "Topo centralizado", "Topo direito", "Rodapé"].index(USER.get("logo_posicao") or "Topo esquerdo") if USER.get("logo_posicao") in ["Topo esquerdo", "Topo centralizado", "Topo direito", "Rodapé"] else 0)
    with c2:
        formacao = st.text_area("Formação", value=USER.get("formacao") or "", height=100)
        espec = st.text_area("Especializações", value=USER.get("especializacoes") or "", height=100)
        st.markdown("<div class='upload-label'>Upload da sua logo</div>", unsafe_allow_html=True)
        logo = st.file_uploader("Upload da sua logo", type=["png", "jpg", "jpeg"], key="logo_upload", label_visibility="collapsed")
        logo_path = None
        if logo:
            ext = logo.name.split(".")[-1].lower()
            logo_path = os.path.join(UPLOAD_DIR, f"logo_user_{USER_ID}.{ext}")
            with open(logo_path, "wb") as f:
                f.write(logo.getbuffer())
            st.success("Logo carregada.")
    if st.button("Salvar perfil e logo"):
        atualizar_perfil(USER_ID, prof, registro, formacao, espec, assinatura, logo_path, logo_pos)
        st.success("Perfil atualizado.")
        st.rerun()
    card_fim()

def tela_pacientes():
    pacientes = listar_pacientes(USER_ID)
    card_inicio("Banco de pacientes", "Cada profissional acessa apenas seus próprios pacientes.")
    st.metric("Pacientes cadastrados", len(pacientes))

    # Accordion premium próprio: substitui st.expander para eliminar definitivamente
    # o bug visual keyboard_arrow_down / _arrow_down do Streamlit.
    if "mostrar_form_paciente" not in st.session_state:
        st.session_state.mostrar_form_paciente = False

    aberto = st.session_state.mostrar_form_paciente
    seta = "⌄" if aberto else "›"
    st.markdown(f'''
    <div class="nexo-accordion-head">
        <div class="nexo-accordion-icon">{seta}</div>
        <div>
            <strong>Cadastrar novo paciente</strong><br>
            <span>Abra o formulário para registrar dados clínicos essenciais.</span>
        </div>
    </div>
    ''', unsafe_allow_html=True)

    toggle_label = "Fechar cadastro" if aberto else "Abrir cadastro"
    if st.button(toggle_label, key="toggle_cadastrar_paciente", use_container_width=True):
        st.session_state.mostrar_form_paciente = not st.session_state.mostrar_form_paciente
        st.rerun()

    if st.session_state.mostrar_form_paciente:
        st.markdown('<div class="nexo-accordion-body">', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            nome = st.text_input("Nome do paciente", key="p_nome")
            idade = st.number_input("Idade", 0, 120, 0, key="p_idade")
        with c2:
            nasc = st.text_input("Data de nascimento", key="p_nasc")
            diag = st.text_input("Diagnóstico", key="p_diag")
        with c3:
            esc = st.text_input("Escolaridade", key="p_esc")
            resp = st.text_input("Responsáveis", key="p_resp")
        agenda_inicial = False
        dia_idx = None
        hora_agenda = ""
        dur_agenda = 50
        freq_agenda = "Semanal"
        mod_agenda = "Presencial"
        if modo_clinica_liberado(USER):
            st.markdown("#### Agenda recorrente inicial")
            agenda_inicial = st.checkbox("Adicionar este paciente à agenda recorrente", value=True, key="p_agenda_inicial")
            if agenda_inicial:
                a1, a2, a3, a4, a5 = st.columns(5)
                with a1:
                    dia_nome = st.selectbox("Dia fixo", list(DIAS_SEMANA.values()), key="p_dia_agenda")
                    dia_idx = [k for k,v in DIAS_SEMANA.items() if v == dia_nome][0]
                with a2:
                    hora_agenda = time_to_str(st.time_input("Horário", value=parse_time_str("08:00"), key="p_hora_agenda"))
                with a3:
                    dur_agenda = st.number_input("Duração", 20, 180, 50, step=5, key="p_dur_agenda")
                with a4:
                    freq_agenda = st.selectbox("Frequência", ["Semanal", "Quinzenal", "Mensal"], key="p_freq_agenda")
                with a5:
                    mod_agenda = st.selectbox("Modalidade", ["Presencial", "Online", "Domiciliar"], key="p_mod_agenda")
        if st.button("Salvar paciente", key="salvar_paciente_form"):
            if len(pacientes) >= LIM_PAC:
                st.error("Limite de pacientes atingido.")
            elif not nome:
                st.warning("Informe o nome.")
            else:
                salvar_paciente(USER_ID, nome, idade, nasc, diag, esc, resp, dia_idx, hora_agenda, dur_agenda, freq_agenda, mod_agenda)
                st.session_state.mostrar_form_paciente = False
                st.success("Paciente salvo. Se houver agenda recorrente, ele aparecerá automaticamente no calendário.")
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    card_fim()

    card_inicio("Pacientes cadastrados")
    if not pacientes:
        st.info("Nenhum paciente cadastrado ainda.")
    else:
        for p in pacientes:
            st.markdown(f"**{p[1]}** — {p[4] or 'sem diagnóstico'}")
    card_fim()

def tela_evolucao():
    pacientes = listar_pacientes(USER_ID)
    if not pacientes:
        st.info("Cadastre um paciente primeiro.")
        return
    nomes = {f"{p[1]} — {p[4] or 'sem diagnóstico'}": p for p in pacientes}
    p = nomes[st.selectbox("Selecionar paciente", list(nomes.keys()))]
    card_inicio(f"Evolução de {p[1]}")
    c1, c2, c3 = st.columns(3)
    with c1:
        data_ev = st.text_input("Data", value=date.today().strftime("%d/%m/%Y"))
        titulo = st.text_input("Título")
    with c2:
        humor = st.selectbox("Estado observado", ["Regulado", "Oscilante", "Desorganizado", "Engajado", "Evitativo", "Outro"])
        eng = st.slider("Engajamento", 0, 10, 5)
    with c3:
        prog = st.slider("Progresso", 0, 10, 5)
    desc = st.text_area("Descrição clínica", height=150)
    if st.button("Salvar evolução"):
        salvar_evolucao(USER_ID, p[0], data_ev, titulo, desc, humor, eng, prog)
        st.success("Evolução salva.")
        st.rerun()
    card_fim()

    card_inicio("Linha do tempo")
    for ev in listar_evolucoes(USER_ID, p[0]):
        st.markdown(f"**{ev[0]} — {ev[1] or 'Evolução clínica'}**")
        st.caption(f"Estado: {ev[3]} • Engajamento: {ev[4]}/10 • Progresso: {ev[5]}/10")
        st.write(ev[2])
        st.markdown("---")
    card_fim()

def reconhecer_avaliacao_por_nome(nome_arquivo: str, profissao: str):
    nome = (nome_arquivo or "").lower()
    mapa = {
        "nordoff": "Nordoff-Robbins",
        "iaps": "IAPS",
        "demuca": "DEMUCA",
        "mel": "MEL",
        "musicmed": "MUSICMED",
        "gas": "Escala GAS",
        "vbmapp": "VB-MAPP",
        "vb-mapp": "VB-MAPP",
        "ablls": "ABLLS-R",
        "afls": "AFLS",
        "sensorial": "Perfil Sensorial",
        "linguagem": "Avaliação de Linguagem",
        "fono": "Avaliação Fonoaudiológica",
        "psicoped": "Avaliação Psicopedagógica",
        "psicomotor": "Avaliação Psicomotora",
        "aba": "Registro/Protocolo ABA",
        "abc": "Registro ABC",
    }
    for chave, avaliacao in mapa.items():
        if chave in nome:
            return avaliacao
    perfil = perfil_profissional(profissao)
    return perfil.get("avaliacoes", ["Avaliação clínica"])[0]

def tela_upload():
    """Upload inteligente blindado para modelos de avaliação.

    Objetivo desta versão:
    - não quebrar com campo repetido;
    - não quebrar com arquivo vazio, corrompido ou não legível;
    - salvar bytes do arquivo com segurança;
    - permitir configurar manualmente quando o texto não puder ser extraído;
    - impedir DataFrame/Streamlit keys duplicadas;
    - permitir qualquer avaliação como modelo personalizado do profissional.
    """
    pacientes = listar_pacientes(USER_ID)
    paciente_id = None

    def _uid_arquivo(uploaded_file, idx):
        nome = nome_arquivo_seguro(getattr(uploaded_file, "name", f"arquivo_{idx+1}"))
        try:
            tamanho = len(uploaded_file.getvalue())
        except Exception:
            tamanho = 0
        base = f"{idx}_{nome}_{tamanho}"
        return slug_streamlit_key(hashlib.sha1(base.encode("utf-8", errors="ignore")).hexdigest()[:14])

    def _bytes_arquivo(uploaded_file):
        try:
            return uploaded_file.getvalue() or b""
        except Exception:
            try:
                uploaded_file.seek(0)
                return uploaded_file.read() or b""
            except Exception:
                return b""

    def _campos_base_seguros(secoes_auto, texto, avaliacao, profissao):
        campos = []
        for sec in secoes_auto or []:
            titulo = sec.get("titulo", "")
            for campo in sec.get("campos", []) or []:
                campo = normalizar_nome_campo(campo)
                if not campo:
                    continue
                # Mantém coerência por domínio quando a escala foi lida em seções.
                if titulo and titulo not in ["Identificação da avaliação", "Síntese clínica", "Resultados e indicadores"] and not campo.lower().startswith("item"):
                    campos.append(f"{titulo} — {campo}")
                else:
                    campos.append(campo)
        campos.extend(inferir_campos_por_texto(texto or "", avaliacao, profissao) or [])
        campos = campos_unicos(campos, limite=135)
        if len(campos) < 4:
            campos.extend(campos_fallback_por_profissao(profissao))
        # Agora preserva o sentido dos enunciados resumidos em vez de trocar tudo por rótulos genéricos.
        return campos_unicos(campos, limite=135)[:36]

    def _secoes_finais(campos_editados):
        campos_editados = campos_unicos(campos_editados) or [
            "Observações clínicas",
            "Pontuação/resultado",
            "Indicadores observados",
            "Conduta sugerida",
        ]
        return [
            {"titulo": "Identificação da avaliação", "campos": ["Data da avaliação", "Contexto de aplicação", "Responsável pelo preenchimento"]},
            {"titulo": "Resultados e indicadores", "campos": campos_editados},
            {"titulo": "Síntese clínica", "campos": ["Interpretação clínica", "Impacto funcional", "Prioridades terapêuticas"]},
        ]

    def _salvar_bytes_em_upload_dir(nome_original, conteudo_bytes):
        nome_limpo = nome_arquivo_seguro(nome_original)
        stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
        safe_name = f"{USER_ID}_{stamp}_{nome_limpo}".replace("/", "_").replace("\\", "_")
        caminho = os.path.join(UPLOAD_DIR, safe_name)
        with open(caminho, "wb") as f:
            f.write(conteudo_bytes or b"")
        return caminho

    card_inicio(
        "Upload inteligente + avaliações personalizadas",
        "Envie qualquer avaliação, protocolo, planilha, Word, PDF, imagem ou documento clínico. Se o NEXO não conseguir ler o conteúdo, ele cria um modelo configurável manualmente sem travar o app."
    )

    c1, c2 = st.columns([0.58, 0.42])
    with c1:
        profissao_upload = st.selectbox(
            "Área profissional para interpretação",
            PROFISSOES,
            index=PROFISSOES.index(USER.get("profissao")) if USER.get("profissao") in PROFISSOES else 0,
            key="upload_profissao_area"
        )
        perfil = perfil_profissional(profissao_upload)
        st.info(f"Perfil ativo: {perfil['rotulo_relatorio']} • foco: {perfil['foco']}")
    with c2:
        categoria = st.selectbox(
            "Categoria do upload",
            ["Avaliação", "Modelo de avaliação", "Documento clínico", "Planilha", "Imagem", "PDF", "Outro"],
            key="upload_categoria"
        )
        if pacientes:
            opts = {f"{p[1]} — {p[4] or 'sem diagnóstico'}": p[0] for p in pacientes}
            paciente_id = opts[st.selectbox("Vincular ao paciente", list(opts.keys()), key="upload_paciente_vinculado")]
        else:
            st.caption("Nenhum paciente cadastrado. O upload será salvo sem vínculo com paciente.")

    st.markdown("""
    <div class="stats-grid" style="grid-template-columns:repeat(4,1fr);">
      <div class="stat-card"><div class="flow-num">1</div><div><strong>Ler</strong><br><span class="small-muted">Quando possível</span></div></div>
      <div class="stat-card"><div class="flow-num">2</div><div><strong>Estruturar</strong><br><span class="small-muted">Campos únicos</span></div></div>
      <div class="stat-card"><div class="flow-num">3</div><div><strong>Validar</strong><br><span class="small-muted">Domínios e gráfico</span></div></div>
      <div class="stat-card"><div class="flow-num">4</div><div><strong>Relatório</strong><br><span class="small-muted">Pronto para uso</span></div></div>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["Enviar e automatizar", "Minhas avaliações", "Preencher avaliação"])

    with tab1:
        st.markdown("**Motor blindado de upload:** o arquivo é salvo mesmo quando não é possível extrair texto. O NEXO cria uma avaliação clínica automatizada derivada, fiel à finalidade do arquivo, mas sem reproduzir integralmente itens protegidos por direitos autorais.")
        st.info("Proteção autoral: a avaliação criada é uma versão clínica estruturada para uso interno do profissional. Revise nomes, domínios e critérios antes de usar em relatório.")
        st.markdown("**Sugestões da área:** " + ", ".join(perfil.get("avaliacoes", []) or ["Avaliação clínica"]))
        if usuario_deve_usar_modelos_proprios(USER):
            st.warning("Seu perfil usa avaliações próprias. Envie seus protocolos/modelos nesta tela para criar sua biblioteca personalizada.")
        else:
            st.success("Seu perfil está autorizado pelo administrador a usar também as avaliações padrão internas do sistema.")

        arquivos = st.file_uploader(
            "Selecionar avaliações ou modelos",
            type=None,
            accept_multiple_files=True,
            help="Você pode enviar PDF, Word, Excel, CSV, TXT, imagens ou qualquer arquivo usado como protocolo clínico.",
            key="upload_inteligente_modelos"
        )

        analises_preview = []
        if arquivos:
            total_uploads_mes = uso_mes(USER_ID, "upload")
            if total_uploads_mes + len(arquivos) > LIM_UP:
                st.error("Limite mensal de uploads atingido. Exclua arquivos da seleção ou atualize o plano.")
            else:
                st.markdown("### Revisão antes de salvar")
                for idx, arquivo in enumerate(arquivos):
                    uid = _uid_arquivo(arquivo, idx)
                    nome_original = nome_arquivo_seguro(getattr(arquivo, "name", f"arquivo_{idx+1}"))
                    conteudo_bytes = _bytes_arquivo(arquivo)
                    tamanho_kb = len(conteudo_bytes) / 1024 if conteudo_bytes else 0
                    ext = extensao_arquivo(nome_original) or "sem extensão"

                    texto = ""
                    try:
                        texto = texto_arquivo_upload(arquivo) or ""
                    except Exception:
                        texto = ""
                    try:
                        arquivo.seek(0)
                    except Exception:
                        pass

                    try:
                        avaliacao_auto, secoes_auto = montar_modelo_automatizado(nome_original, texto, profissao_upload)
                    except Exception:
                        avaliacao_auto, secoes_auto = reconhecer_avaliacao_por_nome(nome_original, profissao_upload), []

                    campos_base = _campos_base_seguros(secoes_auto, texto, avaliacao_auto, profissao_upload)

                    card_inicio(f"Arquivo {idx+1}: {nome_original}", "Revise o reconhecimento. Nada é salvo sem clicar no botão final.")
                    if not conteudo_bytes:
                        st.error("Arquivo vazio ou não pôde ser lido pelo navegador. Remova este item da seleção e tente novamente.")
                    else:
                        if texto.strip():
                            st.success(f"Texto extraído com sucesso • {ext.upper()} • {tamanho_kb:.1f} KB")
                            with st.expander("Prévia do texto extraído", expanded=False):
                                st.text((texto[:2500] + "...") if len(texto) > 2500 else texto)
                        else:
                            st.warning(f"Não foi possível extrair texto automaticamente • {ext.upper()} • {tamanho_kb:.1f} KB. O modelo será configurado manualmente sem erro.")

                    c1, c2 = st.columns([0.42, 0.58])
                    with c1:
                        opcoes_base = [avaliacao_auto] + (perfil.get("avaliacoes", []) or []) + [
                            "Escala de Objetivos Mensuráveis", "Avaliação clínica personalizada", "Outro"
                        ]
                        if usuario_pode_usar_avaliacoes_padrao(USER):
                            opcoes_base += ["Nordoff-Robbins", "IAPS", "DEMUCA", "MEL"]
                        opcoes = list(dict.fromkeys([x for x in opcoes_base if x]))
                        avaliacao_confirmada = st.selectbox("Avaliação reconhecida", opcoes, key=f"modelo_avaliacao_confirmada_{uid}")
                        if avaliacao_confirmada == "Outro":
                            avaliacao_confirmada = st.text_input("Digite o nome da avaliação", value="Avaliação clínica personalizada", key=f"avaliacao_outro_{uid}")
                        nome_modelo = st.text_input("Nome salvo no perfil", value=avaliacao_confirmada, key=f"nome_modelo_{uid}")
                        salvar_como_modelo = st.checkbox("Salvar como avaliação personalizada do meu perfil", value=True, key=f"salvar_modelo_{uid}")
                        usar_no_relatorio = st.checkbox("Usar como base do relatório", value=True, key=f"usar_relatorio_modelo_{uid}")
                    with c2:
                        campos_txt = st.text_area(
                            "Campos da avaliação automatizada — um por linha",
                            value="\n".join(campos_base),
                            height=210,
                            key=f"campos_modelo_{uid}",
                            help="Campos repetidos, vazios ou enormes são normalizados automaticamente."
                        )
                        justificativa = st.text_area(
                            "Justificativa clínica",
                            value=texto_justificativa_avaliacao(avaliacao_confirmada, profissao_upload),
                            height=120,
                            key=f"just_modelo_{uid}"
                        )

                    campos_editados = campos_unicos([x for x in campos_txt.splitlines()])
                    if not campos_editados:
                        campos_editados = ["Observações clínicas", "Pontuação/resultado", "Indicadores observados", "Conduta sugerida"]

                    st.markdown("#### Gráfico da avaliação")
                    try:
                        spec_auto = especificacao_grafico_avaliacao(avaliacao_confirmada, texto, profissao_upload)
                    except Exception:
                        spec_auto = {"tipo": "barras_percentual", "fonte": "Modelo configurado manualmente.", "dominios": []}
                    tipos_validos = ["barras_percentual", "radar_percentual", "linha_objetivos"]
                    tipo_sugerido = spec_auto.get("tipo", "barras_percentual")
                    if tipo_sugerido not in tipos_validos:
                        tipo_sugerido = "barras_percentual"
                    tipo_grafico = st.selectbox("Tipo de gráfico", tipos_validos, index=tipos_validos.index(tipo_sugerido), key=f"tipo_grafico_modelo_{uid}")

                    df_spec = pd.DataFrame([
                        {"Domínio": normalizar_nome_campo(d.get("dominio", "")), "Mínimo": d.get("minimo", 0), "Máximo": d.get("maximo", 10)}
                        for d in (spec_auto.get("dominios", []) or [])
                    ])
                    if df_spec.empty or "Domínio" not in df_spec.columns:
                        df_spec = pd.DataFrame([
                            {"Domínio": "Resultado global", "Mínimo": 0, "Máximo": 10},
                            {"Domínio": "Funcionalidade observada", "Mínimo": 0, "Máximo": 10},
                            {"Domínio": "Necessidade de suporte", "Mínimo": 0, "Máximo": 10},
                        ])
                    df_spec = df_spec.dropna(how="all")

                    df_edit = st.data_editor(
                        df_spec,
                        use_container_width=True,
                        num_rows="dynamic",
                        key=f"grafico_spec_editor_{uid}",
                        column_config={
                            "Domínio": st.column_config.TextColumn("Domínio/escala", required=True),
                            "Mínimo": st.column_config.NumberColumn("Mínimo", step=1.0),
                            "Máximo": st.column_config.NumberColumn("Máximo", step=1.0),
                        }
                    )
                    fonte_ref = st.text_area(
                        "Fonte/referência do modelo de gráfico",
                        value=spec_auto.get("fonte", "Modelo validado pelo profissional."),
                        height=80,
                        key=f"fonte_grafico_{uid}"
                    )
                    try:
                        grafico_spec = normalizar_grafico_spec(df_edit, tipo_grafico, fonte_ref)
                    except Exception:
                        grafico_spec = {"tipo": tipo_grafico, "fonte": fonte_ref or "Modelo validado pelo profissional.", "dominios": [{"dominio": "Resultado global", "minimo": 0, "maximo": 10}]}
                    st.caption(texto_grafico_fiel(avaliacao_confirmada, grafico_spec))

                    secoes_final = _secoes_finais(campos_editados)
                    analises_preview.append({
                        "bytes": conteudo_bytes,
                        "nome_original": nome_original,
                        "avaliacao_confirmada": avaliacao_confirmada,
                        "nome_modelo": nome_modelo or avaliacao_confirmada,
                        "justificativa": (justificativa or texto_justificativa_avaliacao(avaliacao_confirmada, profissao_upload)) + "\n\n" + aviso_direitos_autorais_modelo(avaliacao_confirmada),
                        "secoes_final": secoes_final,
                        "salvar_como_modelo": salvar_como_modelo,
                        "usar_no_relatorio": usar_no_relatorio,
                        "grafico_spec": grafico_spec,
                        "fonte_ref": fonte_ref,
                        "tipo_arq": getattr(arquivo, "type", "") or ext,
                    })
                    card_fim()

                if st.button("Salvar e criar avaliações automatizadas", use_container_width=True, key="salvar_uploads_automatizados"):
                    erros, salvos, modelos_criados = [], 0, 0
                    for item in analises_preview:
                        try:
                            if not item["bytes"]:
                                erros.append(f"{item['nome_original']}: arquivo vazio ou ilegível.")
                                continue
                            caminho = _salvar_bytes_em_upload_dir(item["nome_original"], item["bytes"])
                            status_rel = "Base de relatório: sim" if item["usar_no_relatorio"] else "Base de relatório: não"
                            registrar_upload(
                                USER_ID, paciente_id, item["nome_original"], caminho, item["tipo_arq"], categoria,
                                item["avaliacao_confirmada"], item["justificativa"] + "\n" + status_rel
                            )
                            salvos += 1
                            if item["salvar_como_modelo"]:
                                salvar_modelo_avaliacao(
                                    USER_ID,
                                    item["nome_modelo"],
                                    profissao_upload,
                                    item["avaliacao_confirmada"],
                                    item["justificativa"],
                                    item["secoes_final"],
                                    item["nome_original"],
                                    caminho,
                                    item["grafico_spec"],
                                    item["fonte_ref"],
                                )
                                modelos_criados += 1
                            log_uso(USER_ID, "upload")
                        except Exception as e:
                            erros.append(f"{item.get('nome_original', 'arquivo')}: {e}")
                    if salvos:
                        st.success(f"{salvos} arquivo(s) salvo(s). {modelos_criados} modelo(s) automatizado(s) criado(s).")
                    if erros:
                        st.error("Alguns itens não puderam ser processados:\n" + "\n".join(erros))
                    if salvos and not erros:
                        st.rerun()

    with tab2:
        st.markdown("### Biblioteca personalizada do profissional")
        try:
            modelos = listar_modelos_avaliacao(USER_ID)
        except Exception as e:
            modelos = []
            st.error(f"Não foi possível carregar os modelos: {e}")
        if modelos:
            df = dataframe_modelos_seguro(modelos)
            st.dataframe(df, use_container_width=True, hide_index=True)
        else:
            st.info("Nenhuma avaliação personalizada salva ainda. Envie um modelo na primeira aba.")

    with tab3:
        try:
            modelos = listar_modelos_avaliacao(USER_ID, profissao_upload)
        except Exception as e:
            modelos = []
            st.error(f"Não foi possível carregar avaliações para preenchimento: {e}")
        if not modelos:
            st.info("Você ainda não tem modelos personalizados para esta área. Envie uma avaliação/modelo primeiro.")
        else:
            opcoes = {f"#{m[0]} — {m[1]} — {m[3]}": m[0] for m in modelos}
            modelo_id = opcoes[st.selectbox("Escolha a avaliação automatizada", list(opcoes.keys()), key="select_modelo_preencher")]
            modelo = buscar_modelo_avaliacao(USER_ID, modelo_id, gestor_id=USER_ID)
            if modelo:
                card_inicio(modelo["nome"], modelo["descricao"])
                respostas = {}
                pontuacoes_campos = {}
                contador_campos = {}
                grafico_spec = modelo.get("grafico") or especificacao_grafico_avaliacao(modelo["avaliacao_detectada"], "", profissao_upload)
                dominios = grafico_spec.get("dominios", []) or []
                minimo_auto = min([float(d.get("minimo", 0) or 0) for d in dominios], default=0.0)
                maximo_auto = max([float(d.get("maximo", 10) or 10) for d in dominios], default=10.0)
                if maximo_auto <= minimo_auto:
                    maximo_auto = minimo_auto + 10

                for sec_idx, sec in enumerate(modelo.get("secoes", []) or []):
                    sec_titulo = sec.get('titulo','Seção')
                    st.markdown(f"#### {sec_titulo}")
                    for campo_idx, campo in enumerate(sec.get("campos", []) or []):
                        campo_limpo = normalizar_nome_campo(campo) or f"Campo {campo_idx+1}"
                        chave_base = campo_limpo.lower()
                        contador_campos[chave_base] = contador_campos.get(chave_base, 0) + 1
                        nome_resposta = campo_limpo if contador_campos[chave_base] == 1 else f"{campo_limpo} ({contador_campos[chave_base]})"
                        widget_key = f"resp_modelo_{modelo_id}_{sec_idx}_{campo_idx}_{contador_campos[chave_base]}_{slug_streamlit_key(campo_limpo)}"
                        if eh_campo_pontuavel(campo_limpo, sec_titulo):
                            col_txt, col_score = st.columns([0.68, 0.32])
                            with col_txt:
                                respostas[nome_resposta] = st.text_area(campo_limpo, key=widget_key, height=76)
                            with col_score:
                                score = widget_pontuacao_automatica(
                                    campo_limpo,
                                    sec_titulo,
                                    minimo_auto,
                                    maximo_auto,
                                    key=f"score_modelo_{modelo_id}_{sec_idx}_{campo_idx}_{contador_campos[chave_base]}_{slug_streamlit_key(campo_limpo)}"
                                )
                                pontuacoes_campos[campo_limpo] = score
                        else:
                            respostas[nome_resposta] = st.text_area(campo_limpo, key=widget_key, height=76)

                st.markdown("### Gráfico automático do relatório")
                st.caption("O gráfico é calculado automaticamente a partir das pontuações dos itens preenchidos acima. Não é necessário lançar pontuação manual separada.")
                st.caption(texto_grafico_fiel(modelo["avaliacao_detectada"], grafico_spec))
                grafico_valores = calcular_grafico_automatico(modelo.get("secoes", []) or [], pontuacoes_campos, grafico_spec)
                if not grafico_valores:
                    st.info("Este modelo ainda não possui domínios gráficos ou itens pontuáveis. A avaliação textual será salva normalmente.")
                else:
                    try:
                        fig_modelo = gerar_grafico_avaliacao_modelo(modelo["nome"], grafico_spec, grafico_valores)
                        if fig_modelo:
                            st.pyplot(fig_modelo, use_container_width=True)
                            plt.close(fig_modelo)
                    except Exception as e:
                        st.warning(f"O gráfico não pôde ser renderizado, mas a avaliação poderá ser salva. Detalhe: {e}")
                    with st.expander("Ajuste fino opcional dos valores do gráfico"):
                        cols_graf = st.columns(2)
                        for i, dominio in enumerate(dominios):
                            label = normalizar_nome_campo(dominio.get("dominio", f"Domínio {i+1}")) or f"Domínio {i+1}"
                            try:
                                minimo = float(dominio.get("minimo", 0) or 0)
                                maximo = float(dominio.get("maximo", 10) or 10)
                            except Exception:
                                minimo, maximo = 0.0, 10.0
                            if maximo <= minimo:
                                maximo = minimo + 1
                            with cols_graf[i % 2]:
                                grafico_valores[label] = st.number_input(
                                    f"{label} ({minimo:g} a {maximo:g})",
                                    min_value=minimo,
                                    max_value=maximo,
                                    value=float(grafico_valores.get(label, minimo)),
                                    step=1.0,
                                    key=f"graf_auto_ajuste_{modelo_id}_{i}_{slug_streamlit_key(label)}"
                                )

                usar_rel = st.checkbox("Usar esta avaliação preenchida no relatório", value=True, key=f"usar_resp_rel_{modelo_id}")
                interpretacao = gerar_interpretacao_modelo(profissao_upload, modelo["avaliacao_detectada"], respostas)
                if grafico_valores:
                    interpretacao += "\n\nResultados quantitativos para gráfico: " + "; ".join([f"{k}: {v}" for k, v in grafico_valores.items()]) + "."
                st.markdown("### Interpretação automática")
                interpretacao_editada = st.text_area("Texto gerado pelo sistema", value=interpretacao, height=180, key=f"interp_modelo_{modelo_id}")
                if st.button("Salvar avaliação preenchida", use_container_width=True, key=f"salvar_resp_{modelo_id}"):
                    try:
                        salvar_resposta_avaliacao(USER_ID, paciente_id, modelo_id, modelo["avaliacao_detectada"], respostas, interpretacao_editada, modelo["descricao"], usar_rel, grafico_valores)
                        st.success("Avaliação preenchida salva com segurança.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Não foi possível salvar a avaliação: {e}")
                card_fim()

    st.markdown("### Últimos uploads")
    rows = listar_uploads(USER_ID)
    if rows:
        try:
            st.dataframe(pd.DataFrame(rows, columns=["Arquivo", "Categoria", "Avaliação detectada", "Criado em"]), use_container_width=True)
        except Exception:
            st.dataframe(pd.DataFrame(rows), use_container_width=True)
    else:
        st.info("Nenhum upload registrado ainda.")
    card_fim()

def tela_avaliacoes_relatorios():
    nordoff = {}
    i_imp = {}
    i_rec = {}
    i_comp = {}
    i_esc = {}
    demuca = {}

    tabs = st.tabs(["Paciente", "Escalas", "Análise", "Relatório"])
    with tabs[0]:
        card_inicio("Dados do profissional")
        c1, c2, c3 = st.columns(3)
        with c1:
            terapeuta = st.text_input("Nome", value=USER.get("assinatura") or USER["nome"])
        with c2:
            profissao_rel = st.selectbox("Profissão", PROFISSOES, index=PROFISSOES.index(USER["profissao"]) if USER.get("profissao") in PROFISSOES else 0)
        with c3:
            registro = st.text_input("Registro", value=USER.get("registro") or "")
        perfil_area = perfil_profissional(profissao_rel)
        st.info(f"Perfil ativo: {perfil_area['rotulo_relatorio']} • foco técnico: {perfil_area['foco']}")
        st.caption("O relatório, a prescrição, os objetivos mensuráveis e as justificativas serão adaptados automaticamente para esta área profissional.")
        card_fim()

        card_inicio("Identificação do paciente")
        c1, c2, c3 = st.columns(3)
        with c1:
            nome = st.text_input("Nome do paciente")
            idade = st.number_input("Idade", 0, 120, 0)
        with c2:
            diagnostico = st.text_input("Diagnóstico")
            data_nascimento = st.text_input("Data de nascimento")
        with c3:
            escolaridade = st.text_input("Escolaridade")
            responsaveis = st.text_input("Responsáveis")
        metodo = st.selectbox("Método / abordagem", ["", "MIG", "TREINI", "ABA", "Particular", "Outro"])
        metodo_forca_gas = metodo in ["MIG", "TREINI"]
        usar_objetivos_mensuraveis = True if metodo_forca_gas else st.checkbox(
            "Usar Escala de Objetivos Mensuráveis no relatório",
            value=True,
            help="Para MIG e TREINI o sistema mantém a escala GAS automaticamente. Para outros métodos, o terapeuta decide se quer incluir a escala."
        )
        historia = st.text_area("História clínica", height=100)
        queixa = st.text_area("Queixa", height=90)
        observacoes = st.text_area("Observações", height=90)
        card_fim()

    with tabs[1]:
        escalas_implementadas, sugestoes_area = sugestoes_avaliacoes_por_profissao(profissao_rel)
        escalas_implementadas = [e for e in escalas_implementadas if e in avaliacoes_padrao_disponiveis(USER, profissao_rel)]
        card_inicio("Avaliações sugeridas para a área", "O sistema adapta a linguagem do relatório ao perfil profissional selecionado.")
        st.markdown("**Sugestões clínicas da área:** " + ", ".join(sugestoes_area))
        if escalas_implementadas:
            st.caption("Avaliações padrão liberadas para este perfil pelo administrador. As demais podem ser anexadas pelo Upload Inteligente.")
        else:
            st.warning("Avaliações padrão internas não estão liberadas para este perfil. Use o Upload Inteligente para cadastrar suas próprias avaliações automatizadas.")
        card_fim()
        escalas = st.multiselect("Selecione as avaliações estruturadas", escalas_implementadas, default=escalas_implementadas[:2]) if escalas_implementadas else []
        
        # Avaliações criadas pelo Upload Inteligente aparecem aqui como avaliações estruturadas do profissional.
        respostas_modelos_upload = []
        try:
            modelos_upload = listar_modelos_avaliacao(USER_ID, profissao_rel)
        except Exception as e:
            modelos_upload = []
            st.warning(f"Não foi possível carregar avaliações automatizadas do upload: {e}")

        if modelos_upload:
            st.markdown("### Avaliações automatizadas do Upload")
            st.caption("Estas avaliações foram criadas a partir dos arquivos enviados pelo profissional. Elas entram no relatório após preenchimento e salvamento.")
            opcoes_modelos_upload = {f"#{m[0]} — {m[1]} — {m[3]}": m[0] for m in modelos_upload}
            selecionados_upload = st.multiselect(
                "Selecione avaliações criadas por upload",
                list(opcoes_modelos_upload.keys()),
                key="rel_modelos_upload_selecionados"
            )
            for selecao in selecionados_upload:
                mid = opcoes_modelos_upload[selecao]
                modelo = buscar_modelo_avaliacao(USER_ID, mid, gestor_id=USER_ID)
                if not modelo:
                    continue
                card_inicio(f"Preencher: {modelo['nome']}", modelo.get("descricao") or aviso_direitos_autorais_modelo(modelo.get("avaliacao_detectada", "avaliação")))
                st.info(aviso_direitos_autorais_modelo(modelo.get("avaliacao_detectada", modelo.get("nome", "avaliação"))))
                respostas_upload = {}
                pontuacoes_upload = {}
                contador_upload = {}
                grafico_spec_upload = modelo.get("grafico") or especificacao_grafico_avaliacao(modelo.get("avaliacao_detectada", modelo.get("nome", "")), "", profissao_rel)
                dominios_upload = grafico_spec_upload.get("dominios", []) or []
                minimo_auto_upload = min([float(d.get("minimo", 0) or 0) for d in dominios_upload], default=0.0)
                maximo_auto_upload = max([float(d.get("maximo", 10) or 10) for d in dominios_upload], default=10.0)
                if maximo_auto_upload <= minimo_auto_upload:
                    maximo_auto_upload = minimo_auto_upload + 10

                for sec_idx, sec in enumerate(modelo.get("secoes", []) or []):
                    sec_titulo = sec.get('titulo','Seção')
                    st.markdown(f"#### {sec_titulo}")
                    for campo_idx, campo_nome in enumerate(sec.get("campos", []) or []):
                        campo_limpo = normalizar_nome_campo(campo_nome) or f"Campo {campo_idx+1}"
                        chave_base = campo_limpo.lower()
                        contador_upload[chave_base] = contador_upload.get(chave_base, 0) + 1
                        nome_resposta = campo_limpo if contador_upload[chave_base] == 1 else f"{campo_limpo} ({contador_upload[chave_base]})"
                        if eh_campo_pontuavel(campo_limpo, sec_titulo):
                            col_txt, col_score = st.columns([0.68, 0.32])
                            with col_txt:
                                respostas_upload[nome_resposta] = st.text_area(
                                    campo_limpo,
                                    height=70,
                                    key=f"rel_resp_upload_{mid}_{sec_idx}_{campo_idx}_{contador_upload[chave_base]}_{slug_streamlit_key(campo_limpo)}"
                                )
                            with col_score:
                                score = widget_pontuacao_automatica(
                                    campo_limpo,
                                    sec_titulo,
                                    minimo_auto_upload,
                                    maximo_auto_upload,
                                    key=f"rel_score_upload_{mid}_{sec_idx}_{campo_idx}_{contador_upload[chave_base]}_{slug_streamlit_key(campo_limpo)}"
                                )
                                pontuacoes_upload[campo_limpo] = score
                        else:
                            respostas_upload[nome_resposta] = st.text_area(
                                campo_limpo,
                                height=70,
                                key=f"rel_resp_upload_{mid}_{sec_idx}_{campo_idx}_{contador_upload[chave_base]}_{slug_streamlit_key(campo_limpo)}"
                            )

                grafico_valores_upload = calcular_grafico_automatico(modelo.get("secoes", []) or [], pontuacoes_upload, grafico_spec_upload)
                if dominios_upload:
                    st.markdown("#### Gráfico automático do relatório")
                    st.caption("O gráfico é gerado automaticamente a partir das pontuações dos itens. Ajuste manual é opcional.")
                    try:
                        fig_upload = gerar_grafico_avaliacao_modelo(modelo["nome"], grafico_spec_upload, grafico_valores_upload)
                        if fig_upload:
                            st.pyplot(fig_upload, use_container_width=True)
                            plt.close(fig_upload)
                    except Exception as e:
                        st.warning(f"O gráfico não pôde ser exibido, mas a avaliação ainda pode ser salva. Detalhe: {e}")
                    with st.expander("Ajuste fino opcional dos valores do gráfico"):
                        cg = st.columns(2)
                        for i, dominio in enumerate(dominios_upload):
                            label = normalizar_nome_campo(dominio.get("dominio", f"Domínio {i+1}")) or f"Domínio {i+1}"
                            try:
                                minimo = float(dominio.get("minimo", 0) or 0)
                                maximo = float(dominio.get("maximo", 10) or 10)
                            except Exception:
                                minimo, maximo = 0.0, 10.0
                            if maximo <= minimo:
                                maximo = minimo + 1
                            with cg[i % 2]:
                                grafico_valores_upload[label] = st.number_input(
                                    f"{label} ({minimo:g} a {maximo:g})",
                                    min_value=minimo,
                                    max_value=maximo,
                                    value=float(grafico_valores_upload.get(label, minimo)),
                                    step=1.0,
                                    key=f"rel_graf_auto_ajuste_{mid}_{i}_{slug_streamlit_key(label)}"
                                )
                else:
                    st.info("Este modelo não possui domínios gráficos; o relatório usará a interpretação textual.")
                interpretacao_upload = gerar_interpretacao_modelo(profissao_rel, modelo.get("avaliacao_detectada", modelo.get("nome", "Avaliação")), respostas_upload)
                if grafico_valores_upload:
                    interpretacao_upload += "\n\nResultados quantitativos para gráfico: " + "; ".join([f"{k}: {v}" for k, v in grafico_valores_upload.items()]) + "."
                interpretacao_upload = st.text_area(
                    "Interpretação automática para o relatório",
                    value=interpretacao_upload,
                    height=150,
                    key=f"rel_interp_upload_{mid}"
                )
                if st.button("Salvar esta avaliação para o relatório", use_container_width=True, key=f"rel_salvar_upload_{mid}"):
                    try:
                        salvar_resposta_avaliacao(USER_ID, None, mid, modelo.get("avaliacao_detectada", modelo.get("nome", "Avaliação")), respostas_upload, interpretacao_upload, modelo.get("descricao", ""), True, grafico_valores_upload)
                        log_uso(USER_ID, "avaliacao")
                        st.success("Avaliação estruturada salva e disponível na confecção do relatório automático.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Não foi possível salvar esta avaliação: {e}")
                card_fim()
        else:
            st.info("Nenhuma avaliação automatizada criada por upload para esta área ainda.")
        if "Nordoff-Robbins" in escalas:
            card_inicio("Escala Nordoff-Robbins", "Pontuação de 0 a 5 por domínio.")
            c1, c2 = st.columns(2)
            with c1:
                for k in ["Expressão emocional", "Exploração sonora", "Interação musical", "Engajamento", "Responsividade musical"]:
                    nordoff[k] = campo(k, "n_" + k)
            with c2:
                for k in ["Iniciativa musical", "Sustentação da atividade musical", "Comunicação não verbal", "Reciprocidade musical", "Organização musical"]:
                    nordoff[k] = campo(k, "n_" + k)
            card_fim()

        if "IAPS" in escalas:
            card_inicio("IAPS")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("#### IAPS — Improvisação")
                for k in ["Iniciativa sonora", "Resposta musical", "Organização sonora", "Interação musical"]:
                    i_imp[k] = campo(k, "i_" + k)
                st.markdown("#### IAPS — Recriação")
                for k in ["Memória musical", "Coordenação motora", "Seguimento musical", "Participação"]:
                    i_rec[k] = campo(k, "r_" + k)
            with c2:
                st.markdown("#### IAPS — Composição")
                for k in ["Criatividade", "Organização de ideias", "Expressão simbólica", "Autoria"]:
                    i_comp[k] = campo(k, "c_" + k)
                st.markdown("#### IAPS — Escuta Musical")
                for k in ["Atenção auditiva", "Resposta emocional", "Reflexão", "Integração da experiência sonora"]:
                    i_esc[k] = campo(k, "e_" + k)
            card_fim()

        if "DEMUCA" in escalas:
            card_inicio("DEMUCA", "N = Não | P = Pouco | M = Muito")
            for dominio, info in DEMUCA_DOMINIOS.items():
                st.markdown(f"#### DEMUCA — {dominio}")
                demuca[dominio] = {}
                for item, peso in info["itens"]:
                    resp = st.radio(f"{item}" + (" • peso 2" if peso == 2 else ""), ["N", "P", "M"], horizontal=True, key=f"d_{dominio}_{item}")
                    demuca[dominio][item] = {"resposta": resp, "peso": peso}
                st.markdown("<hr style='border-color:rgba(148,163,184,.12);'>", unsafe_allow_html=True)
            card_fim()

    def calc():
        nt = sum(nordoff.values()) if nordoff else None
        ti = {"Improvisação": sum(i_imp.values()), "Recriação": sum(i_rec.values()), "Composição": sum(i_comp.values()), "Escuta Musical": sum(i_esc.values())} if i_imp else {}
        td = calcular_demuca(demuca) if demuca else {}
        prej = identificar_prejuizos(nt, ti, td)
        gas = gerar_gas(prej, profissao_rel) if usar_objetivos_mensuraveis else []
        nome_escala_objetivos = "Escala GAS" if metodo in ["MIG", "TREINI"] else "Escala de Objetivos Mensuráveis"
        resumo = gerar_resumo(nome, idade, diagnostico, nt, nordoff, ti, td, profissao_rel)
        justificativas = [texto_justificativa_avaliacao(e, profissao_rel) for e in escalas]
        return {
            "nordoff_total": nt, "totais_iaps": ti, "totais_demuca": td, "prejuizos": prej, "gas": gas,
            "nome_escala_objetivos": nome_escala_objetivos, "justificativas": justificativas,
            "resumo": resumo, "prescricao": gerar_prescricao(prej, profissao_rel), "conduta": gerar_conduta(prej, profissao_rel),
            "grafico_nordoff": gerar_grafico_barras("Nordoff-Robbins", nordoff, {k: 5 for k in nordoff}) if nordoff else None,
            "grafico_iaps": gerar_grafico_barras("IAPS", ti, {k: 20 for k in ti}) if ti else None,
            "grafico_demuca": gerar_grafico_barras("DEMUCA", td, {d: DEMUCA_DOMINIOS[d]["maximo"] for d in td}) if td else None,
            "avaliacoes_personalizadas": listar_respostas_relatorio(USER_ID, None),
            "grafico_radar": gerar_grafico_radar(ti, nt, td)
        }

    with tabs[2]:
        dados = calc()
        stat_cards(contar_pacientes(USER_ID), len(escalas), len(dados.get("gas", [])), uso_mes(USER_ID, "upload"))
        card_inicio("Áreas prioritárias")
        for p in dados["prejuizos"]:
            st.markdown(f"**{p['area']}** — {p['origem']} | `{p['nivel']}`")
            st.caption(p["objetivo"])
        card_fim()
        card_inicio("Gráficos")
        for k, t in [("grafico_nordoff", "Nordoff-Robbins"), ("grafico_iaps", "IAPS"), ("grafico_demuca", "DEMUCA"), ("grafico_radar", "Perfil Geral")]:
            if dados[k]:
                st.subheader(t)
                st.image(dados[k], use_container_width=True)
        if dados.get("avaliacoes_personalizadas"):
            st.subheader("Avaliações personalizadas salvas no perfil")
            for av in dados["avaliacoes_personalizadas"]:
                st.markdown(f"**{av['modelo_nome']}** — {av['avaliacao_nome']}")
                fig_av = gerar_grafico_avaliacao_modelo(av['modelo_nome'], av['grafico_spec'], av['grafico_valores'])
                if fig_av:
                    st.pyplot(fig_av, use_container_width=True)
                    plt.close(fig_av)
                st.caption((av.get('interpretacao') or '')[:500])
        card_fim()

    with tabs[3]:
        dados = calc()
        card_inicio("Relatório clínico automatizado")
        formato = st.radio("Formato do arquivo", ["Word .docx", "PDF .pdf"], horizontal=True)
        if dados.get("justificativas"):
            st.markdown("### Justificativa técnica das avaliações")
            for j in dados["justificativas"]:
                st.info(j)
        st.text_area("Resumo", dados["resumo"], height=180)
        st.text_area("Plano e prescrição", dados["prescricao"], height=140)
        st.text_area("Parecer / Conduta", dados["conduta"], height=180)
        if dados.get("gas"):
            st.caption(f"{dados.get('nome_escala_objetivos', 'Escala de Objetivos Mensuráveis')} será incluída apenas no documento gerado, conforme a regra selecionada.")
        else:
            st.caption("Escala de Objetivos Mensuráveis desativada para este relatório.")

        if st.button("Preparar relatório"):
            if uso_mes(USER_ID, "relatorio") >= LIM_REL:
                st.error("Limite mensal de relatórios atingido.")
            else:
                ctx = locals()
                if formato == "Word .docx":
                    st.session_state.relatorio_buffer = criar_word(ctx, dados, USER)
                    st.session_state.relatorio_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.session_state.relatorio_ext = "docx"
                else:
                    st.session_state.relatorio_buffer = criar_pdf_simples(ctx, dados, USER)
                    st.session_state.relatorio_mime = "application/pdf"
                    st.session_state.relatorio_ext = "pdf"
                st.session_state.relatorio_gerado = True
                log_uso(USER_ID, "relatorio")
                st.success("Relatório preparado.")

        if st.session_state.get("relatorio_gerado"):
            base = nome.replace(" ", "_") if nome else "paciente"
            st.download_button(
                "Baixar relatório",
                st.session_state.relatorio_buffer,
                file_name=f"relatorio_{base}.{st.session_state.relatorio_ext}",
                mime=st.session_state.relatorio_mime
            )
        card_fim()

def tela_planos():
    card_inicio("Planos e acesso", "Escolha ou altere o plano conectado ao Mercado Pago.")
    st.markdown('<div class="price-grid">', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    cols = st.columns(5)
    for i, (nome, p) in enumerate(PLANOS.items()):
        with cols[i]:
            st.markdown(f"""
            <div class="card price-card">
              <div>
                <span class="badge">{nome}</span>
                <div class="price">R$ {p['valor']:.0f}</div>
                <p class="small-muted">mensal</p>
                <p>{p['descricao']}</p>
                <p class="small-muted">Pacientes: {'Ilimitado' if p['pacientes'] > 9999 else p['pacientes']}<br>Relatórios: {'Ilimitado' if p['relatorios_mes'] > 9999 else p['relatorios_mes']}<br>Uploads: {'Ilimitado' if p['uploads_mes'] > 9999 else p['uploads_mes']}</p>
              </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Assinar {nome}", key=f"assinar_{nome}"):
                ok, msg, url = criar_preferencia_mercado_pago(USER_ID, nome)
                if ok and url:
                    st.success("Assinatura criada.")
                    st.link_button("Pagar com Mercado Pago", url)
                else:
                    st.warning(msg)
    card_fim()

def tela_admin():
    if USER["role"] != "admin":
        st.error("Acesso restrito.")
        return

    tabs = st.tabs(["Usuários", "Criar acesso", "Pagamentos", "Logs", "Configurações"])
    with tabs[0]:
        rows = listar_usuarios()
        card_inicio("Gerenciamento de usuários")
        df = pd.DataFrame([{
            "ID": u["id"], "Nome": u.get("nome"), "E-mail": u.get("email"), "Role": u.get("role"),
            "Status": u.get("status"), "Plano": u.get("plano"), "Tipo": u.get("tipo_acesso"),
            "Expiração": u.get("data_expiracao"), "Pacientes": u.get("limite_pacientes"),
            "Relatórios": u.get("limite_relatorios_mes"), "Uploads": u.get("limite_uploads_mes"), "Avaliações padrão": "Sim" if int(u.get("avaliacoes_padrao_autorizado") or 0) else "Não"
        } for u in rows])
        st.dataframe(df, use_container_width=True)
        opts = {f"{u['id']} — {u['nome']} — {u['email']}": u for u in rows}
        u = opts[st.selectbox("Editar usuário", list(opts.keys()))]
        c1, c2 = st.columns(2)
        with c1:
            nome = st.text_input("Nome", value=u.get("nome") or "", key="adm_nome")
            email = st.text_input("E-mail", value=u.get("email") or "", key="adm_email")
            senha = st.text_input("Nova senha opcional", type="password", key="adm_senha")
            prof = st.selectbox("Profissão", PROFISSOES, index=PROFISSOES.index(u.get("profissao")) if u.get("profissao") in PROFISSOES else 0, key="adm_prof")
        with c2:
            tipo = st.selectbox("Tipo de acesso", TIPOS_ACESSO, index=TIPOS_ACESSO.index(u.get("tipo_acesso")) if u.get("tipo_acesso") in TIPOS_ACESSO else 1, key="adm_tipo")
            status = st.selectbox("Status", STATUS_ACESSO, index=STATUS_ACESSO.index(u.get("status")) if u.get("status") in STATUS_ACESSO else 0, key="adm_status")
            plano = st.selectbox("Plano", list(PLANOS.keys()), index=list(PLANOS.keys()).index(normalizar_plano_nome(u.get("plano"))) if normalizar_plano_nome(u.get("plano")) in PLANOS else 0, key="adm_plano")
            exp = st.date_input("Expiração", value=datetime.strptime(u.get("data_expiracao") or date.today().strftime("%Y-%m-%d"), "%Y-%m-%d").date(), key="adm_exp")
        p = PLANOS[plano]
        c3, c4, c5 = st.columns(3)
        with c3:
            lp = st.number_input("Limite pacientes", 0, 999999, int(u.get("limite_pacientes") or p["pacientes"]), key="adm_lp")
        with c4:
            lr = st.number_input("Relatórios/mês", 0, 999999, int(u.get("limite_relatorios_mes") or p["relatorios_mes"]), key="adm_lr")
        with c5:
            lu = st.number_input("Uploads/mês", 0, 999999, int(u.get("limite_uploads_mes") or p["uploads_mes"]), key="adm_lu")
        padrao_aut = st.checkbox("Autorizar avaliações padrão do sistema para este usuário", value=bool(int(u.get("avaliacoes_padrao_autorizado") or 0)), key="adm_avaliacoes_padrao")
        obs = st.text_area("Observação interna", value=u.get("observacao_admin") or "", key="adm_obs")

        if st.button("Salvar alterações do usuário"):
            conn = db()
            c = conn.cursor()
            role = "admin" if tipo == "admin" else "user"
            if senha:
                c.execute("""UPDATE users SET nome=?, email=?, senha_hash=?, profissao=?, assinatura=?, plano=?, status=?,
                role=?, tipo_acesso=?, data_expiracao=?, limite_pacientes=?, limite_relatorios_mes=?, limite_uploads_mes=?,
                observacao_admin=?, avaliacoes_padrao_autorizado=? WHERE id=?""",
                (nome, email.lower().strip(), hash_senha(senha), prof, nome, plano, status, role, tipo, exp.strftime("%Y-%m-%d"),
                 int(lp), int(lr), int(lu), obs, 1 if padrao_aut else 0, u["id"]))
            else:
                c.execute("""UPDATE users SET nome=?, email=?, profissao=?, assinatura=?, plano=?, status=?,
                role=?, tipo_acesso=?, data_expiracao=?, limite_pacientes=?, limite_relatorios_mes=?, limite_uploads_mes=?,
                observacao_admin=?, avaliacoes_padrao_autorizado=? WHERE id=?""",
                (nome, email.lower().strip(), prof, nome, plano, status, role, tipo, exp.strftime("%Y-%m-%d"),
                 int(lp), int(lr), int(lu), obs, 1 if padrao_aut else 0, u["id"]))
            conn.commit()
            conn.close()
            st.success("Usuário atualizado.")
            st.rerun()

        if st.button("Liberar plano selecionado por 30 dias"):
            atualizar_plano_usuario(u["id"], plano, "ativo")
            st.success("Plano liberado.")
            st.rerun()
        card_fim()

    with tabs[1]:
        card_inicio("Criar novo acesso")
        c1, c2 = st.columns(2)
        with c1:
            nome = st.text_input("Nome completo", key="cr_nome")
            email = st.text_input("E-mail", key="cr_email")
            senha = st.text_input("Senha temporária", value="Nexo2026!", key="cr_senha")
            prof = st.selectbox("Profissão", PROFISSOES, key="cr_prof")
            registro = st.text_input("Registro", key="cr_reg")
        with c2:
            tipo = st.selectbox("Tipo", TIPOS_ACESSO, index=1, key="cr_tipo")
            status = st.selectbox("Status", STATUS_ACESSO, index=0, key="cr_status")
            plano = st.selectbox("Plano", list(PLANOS.keys()), key="cr_plano")
            exp = st.date_input("Expiração", value=date.today() + timedelta(days=30), key="cr_exp")
        p = PLANOS[plano]
        if st.button("Criar acesso personalizado"):
            conn = db()
            c = conn.cursor()
            try:
                role = "admin" if tipo == "admin" else "user"
                c.execute("""INSERT INTO users (nome,email,senha_hash,profissao,registro,assinatura,plano,status,role,tipo_acesso,
                data_inicio,data_expiracao,limite_pacientes,limite_relatorios_mes,limite_uploads_mes,observacao_admin,criado_em)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (nome, email.lower().strip(), hash_senha(senha), prof, registro, nome, plano, status, role, tipo,
                 date.today().strftime("%Y-%m-%d"), exp.strftime("%Y-%m-%d"), p["pacientes"], p["relatorios_mes"],
                 p["uploads_mes"], "Criado pelo admin.", datetime.now().isoformat(timespec="seconds")))
                conn.commit()
                st.success("Acesso criado.")
            except sqlite3.IntegrityError:
                st.error("E-mail já cadastrado.")
            finally:
                conn.close()
        card_fim()

    with tabs[2]:
        card_inicio("Pagamentos Mercado Pago")
        conn = db()
        df = pd.read_sql_query("SELECT * FROM pagamentos ORDER BY id DESC", conn)
        conn.close()
        st.dataframe(df, use_container_width=True)
        st.info("Em produção, configure um webhook do Mercado Pago para atualizar o status automaticamente.")
        card_fim()

    with tabs[3]:
        card_inicio("Logs de uso")
        conn = db()
        df = pd.read_sql_query("SELECT * FROM usage_logs ORDER BY id DESC LIMIT 500", conn)
        conn.close()
        st.dataframe(df, use_container_width=True)
        card_fim()

    with tabs[4]:
        card_inicio("Configurações")
        st.write("Token Mercado Pago configurado:", "Sim" if os.getenv("MERCADO_PAGO_ACCESS_TOKEN") else "Não")
        st.write("URL pública:", os.getenv("APP_PUBLIC_URL", "http://localhost:8501"))
        st.code("set MERCADO_PAGO_ACCESS_TOKEN=SEU_ACCESS_TOKEN\nset APP_PUBLIC_URL=https://nexoclinic.streamlit.app/", language="bash")
        card_fim()

# ==========================================================
# SESSION / ROUTER
# ==========================================================
if "relatorio_gerado" not in st.session_state:
    st.session_state.relatorio_gerado = False
if "user_id" not in st.session_state:
    st.session_state.user_id = None

if not st.session_state.user_id:
    tela_login()
    st.stop()

USER = buscar_usuario(st.session_state.user_id)
if not USER:
    st.session_state.clear()
    st.rerun()

if USER["role"] != "admin" and USER["status"] in ["aguardando_pagamento", "expirado", "bloqueado"]:
    if USER["status"] == "bloqueado":
        st.error("Seu acesso está bloqueado. Fale com o administrador.")
        st.stop()
    tela_escolher_plano(USER)
    if st.button("Sair"):
        st.session_state.clear()
        st.rerun()
    st.stop()

if USER["role"] != "admin" and USER.get("data_expiracao"):
    try:
        if datetime.strptime(USER["data_expiracao"], "%Y-%m-%d").date() < date.today():
            conn = db()
            c = conn.cursor()
            c.execute("UPDATE users SET status='expirado' WHERE id=?", (USER["id"],))
            conn.commit()
            conn.close()
            st.rerun()
    except ValueError:
        pass

USER_ID = USER["id"]
LIM_PAC = int(USER.get("limite_pacientes") or plano_config(USER.get("plano"))["pacientes"])
LIM_REL = int(USER.get("limite_relatorios_mes") or plano_config(USER.get("plano"))["relatorios_mes"])
LIM_UP = int(USER.get("limite_uploads_mes") or plano_config(USER.get("plano"))["uploads_mes"])

if "page" not in st.session_state:
    st.session_state.page = "Dashboard"

if "page_history" not in st.session_state:
    st.session_state.page_history = []



def go_to(page_name):
    if "page" not in st.session_state:
        st.session_state.page = "Dashboard"
    if "page_history" not in st.session_state:
        st.session_state.page_history = []
    if st.session_state.page != page_name:
        st.session_state.page_history.append(st.session_state.page)
        st.session_state.page = page_name
        st.rerun()

def render_back_button():
    if "page_history" not in st.session_state:
        st.session_state.page_history = []
    c_back, _ = st.columns([0.16, 0.84])
    with c_back:
        if st.session_state.page_history:
            if st.button("← Voltar", use_container_width=True, key="global_back_button"):
                st.session_state.page = st.session_state.page_history.pop()
                st.rerun()



# ==========================================================
# MODO CLÍNICA — AGENDA + EVOLUÇÕES OBRIGATÓRIAS
# ==========================================================
DIAS_SEMANA = {
    0: "Segunda", 1: "Terça", 2: "Quarta", 3: "Quinta", 4: "Sexta", 5: "Sábado", 6: "Domingo"
}

def modo_clinica_liberado(user):
    """Libera o Modo Clínica para o admin master e para planos Clinic/Enterprise.

    A checagem é propositalmente robusta para cobrir variações do cadastro:
    - role admin;
    - tipo_acesso admin;
    - e-mail master;
    - plano Clinic, Enterprise ou qualquer plano contendo clínica/clinic/enterprise/ilimitado/empresarial.
    """
    if not user:
        return False
    plano = str(user.get("plano") or "").strip().lower()
    role = str(user.get("role") or "").strip().lower()
    tipo = str(user.get("tipo_acesso") or "").strip().lower()
    email = str(user.get("email") or "").strip().lower()
    return (
        role == "admin"
        or tipo == "admin"
        or email == ADMIN_EMAIL.lower()
        or plano in {"clinic", "enterprise"}
        or "clinic" in plano
        or "clínica" in plano
        or "clinica" in plano
        or "ilimitado" in plano
        or "empresarial" in plano
    )

def parse_time_str(s, default="08:00"):
    try:
        h, m = (s or default).split(":")[:2]
        return datetime.strptime(f"{int(h):02d}:{int(m):02d}", "%H:%M").time()
    except Exception:
        return datetime.strptime(default, "%H:%M").time()

def time_to_str(t):
    return t.strftime("%H:%M") if hasattr(t, "strftime") else str(t or "")[:5]

def get_agenda_config(user_id):
    conn = db(); c = conn.cursor()
    c.execute("SELECT dias_semana, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao FROM agenda_config WHERE user_id=?", (user_id,))
    row = c.fetchone()
    if not row:
        c.execute("""INSERT INTO agenda_config (user_id, dias_semana, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao, atualizado_em)
                  VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""", (user_id, "0,1,2,3,4", "08:00", "18:00", "12:00", "13:00", 50, 10, datetime.now().isoformat(timespec="seconds")))
        conn.commit()
        row = ("0,1,2,3,4", "08:00", "18:00", "12:00", "13:00", 50, 10)
    conn.close()
    dias = [int(x) for x in (row[0] or "0,1,2,3,4").split(",") if str(x).strip().isdigit()]
    return {"dias_semana": dias, "hora_inicio": row[1], "hora_fim": row[2], "almoco_inicio": row[3], "almoco_fim": row[4], "duracao_padrao": row[5], "intervalo_padrao": row[6]}

def salvar_agenda_config(user_id, dias_semana, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao):
    conn = db(); c = conn.cursor()
    dias_txt = ",".join(str(d) for d in dias_semana)
    c.execute("""INSERT INTO agenda_config (user_id, dias_semana, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao, atualizado_em)
              VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
              ON CONFLICT(user_id) DO UPDATE SET dias_semana=excluded.dias_semana, hora_inicio=excluded.hora_inicio, hora_fim=excluded.hora_fim,
              almoco_inicio=excluded.almoco_inicio, almoco_fim=excluded.almoco_fim, duracao_padrao=excluded.duracao_padrao,
              intervalo_padrao=excluded.intervalo_padrao, atualizado_em=excluded.atualizado_em""",
              (user_id, dias_txt, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao, datetime.now().isoformat(timespec="seconds")))
    conn.commit(); conn.close()

def pacientes_com_agenda(user_id):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT id, nome, diagnostico, dia_semana, hora_atendimento, duracao_sessao, frequencia, modalidade, agenda_status
              FROM pacientes WHERE user_id=? AND agenda_status='ativo' AND dia_semana IS NOT NULL AND hora_atendimento IS NOT NULL AND hora_atendimento!=''
              ORDER BY dia_semana, hora_atendimento""", (user_id,))
    rows = c.fetchall(); conn.close(); return rows

def ensure_evento_recorrente(user_id, paciente, data_obj):
    pid, nome, diag, dia_semana, hora, dur, freq, modalidade, status = paciente
    if dia_semana is None or int(dia_semana) != data_obj.weekday():
        return None
    if freq == "Quinzenal" and (data_obj.isocalendar().week % 2 != 0):
        return None
    if freq == "Mensal" and data_obj.day > 7:
        return None
    inicio_dt = datetime.combine(data_obj, parse_time_str(hora))
    fim_dt = inicio_dt + timedelta(minutes=int(dur or 50))
    conn = db(); c = conn.cursor()
    c.execute("""SELECT id FROM agenda_eventos WHERE user_id=? AND paciente_id=? AND data=? AND hora_inicio=? AND origem='recorrente'""",
              (user_id, pid, data_obj.strftime("%Y-%m-%d"), hora))
    existing = c.fetchone()
    if existing:
        conn.close(); return existing[0]
    c.execute("""INSERT INTO agenda_eventos (user_id, paciente_id, data, hora_inicio, hora_fim, titulo, modalidade, status, origem, criado_em, atualizado_em)
              VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (user_id, pid, data_obj.strftime("%Y-%m-%d"), hora, fim_dt.strftime("%H:%M"), nome, modalidade or "Presencial", "previsto", "recorrente", datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")))
    event_id = c.lastrowid
    conn.commit(); conn.close(); return event_id

def gerar_eventos_semana(user_id, inicio_semana):
    gerar_eventos_periodo_recorrentes(user_id, inicio_semana, inicio_semana + timedelta(days=6))

def gerar_eventos_periodo_recorrentes(user_id, data_inicio, data_fim):
    """Gera compromissos recorrentes dentro de qualquer intervalo: semana, mês ou ano."""
    atual = data_inicio
    pacientes = pacientes_com_agenda(user_id)
    while atual <= data_fim:
        for pac in pacientes:
            ensure_evento_recorrente(user_id, pac, atual)
        atual += timedelta(days=1)

def inicio_fim_mes(data_ref):
    inicio = data_ref.replace(day=1)
    ultimo = calendar.monthrange(data_ref.year, data_ref.month)[1]
    fim = data_ref.replace(day=ultimo)
    return inicio, fim

def nome_mes_pt(mes):
    nomes = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    return nomes[mes - 1]

def render_evento_card(e, compact=False):
    cor, label = cor_status(e[6], e[1])
    min_h = "76px" if compact else "120px"
    fonte = ".78rem" if compact else ".86rem"
    st.markdown(f"""
    <div style="border-left:5px solid {cor}; background:rgba(15,23,42,.68); border:1px solid rgba(148,163,184,.14); border-radius:14px; padding:10px; margin:8px 0; min-height:{min_h};">
        <strong>{e[2]} - {e[3]}</strong><br>
        <span style="color:#f8fafc;font-size:{fonte};">{e[4]}</span><br>
        <span style="color:#9fb0cf;font-size:.76rem;">{e[5]} • {label}</span>
    </div>
    """, unsafe_allow_html=True)

def render_acoes_evento(e, prefix):
    if st.button("Evoluir", key=f"{prefix}_evoluir_evento_{e[0]}", use_container_width=True):
        st.session_state.evento_evolucao_id = e[0]
        st.session_state.evento_evolucao_paciente_id = e[7]
        st.session_state.evento_evolucao_nome = e[4]
        st.session_state.evento_evolucao_data = e[1]
        st.rerun()
    cc1, cc2 = st.columns(2)
    with cc1:
        if st.button("Cancelar", key=f"{prefix}_cancelar_evento_{e[0]}", use_container_width=True):
            atualizar_evento_status(USER_ID, e[0], "cancelado"); st.rerun()
    with cc2:
        if st.button("Falta", key=f"{prefix}_falta_evento_{e[0]}", use_container_width=True):
            atualizar_evento_status(USER_ID, e[0], "falta_justificada"); st.rerun()

def render_form_evolucao_agenda():
    if st.session_state.get("evento_evolucao_id"):
        eid = st.session_state.evento_evolucao_id
        pid = st.session_state.evento_evolucao_paciente_id
        st.markdown("---")
        card_inicio(f"Registrar evolução — {st.session_state.get('evento_evolucao_nome')}")
        c1,c2,c3 = st.columns(3)
        with c1:
            titulo = st.text_input("Título", value="Evolução clínica", key="agenda_ev_titulo")
            humor = st.selectbox("Estado observado", ["Regulado", "Oscilante", "Desorganizado", "Engajado", "Evitativo", "Outro"], key="agenda_ev_humor")
        with c2:
            eng = st.slider("Engajamento", 0, 10, 5, key="agenda_ev_eng")
        with c3:
            prog = st.slider("Progresso", 0, 10, 5, key="agenda_ev_prog")
        desc = st.text_area("Descrição clínica da sessão", height=150, key="agenda_ev_desc")
        if st.button("Salvar evolução e marcar como OK", key="agenda_salvar_evolucao"):
            salvar_evolucao(USER_ID, pid, datetime.strptime(st.session_state.evento_evolucao_data, "%Y-%m-%d").strftime("%d/%m/%Y"), titulo, desc, humor, eng, prog, event_id=eid)
            for k in ["evento_evolucao_id", "evento_evolucao_paciente_id", "evento_evolucao_nome", "evento_evolucao_data"]:
                st.session_state.pop(k, None)
            st.success("Evolução registrada e atendimento marcado como OK.")
            st.rerun()
        card_fim()

def listar_eventos_periodo(user_id, data_inicio, data_fim):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT e.id, e.data, e.hora_inicio, e.hora_fim, e.titulo, e.modalidade, e.status, e.paciente_id, p.diagnostico
              FROM agenda_eventos e LEFT JOIN pacientes p ON p.id=e.paciente_id
              WHERE e.user_id=? AND e.data BETWEEN ? AND ? ORDER BY e.data, e.hora_inicio""",
              (user_id, data_inicio.strftime("%Y-%m-%d"), data_fim.strftime("%Y-%m-%d")))
    rows = c.fetchall(); conn.close(); return rows

def eventos_contadores(user_id, inicio, fim):
    eventos = listar_eventos_periodo(user_id, inicio, fim)
    realizados = sum(1 for e in eventos if e[6] == "evolucao_ok")
    pendentes = sum(1 for e in eventos if e[6] in ["previsto", "realizado_sem_evolucao"] and datetime.strptime(e[1], "%Y-%m-%d").date() <= date.today())
    cancelados = sum(1 for e in eventos if e[6] in ["cancelado", "falta_justificada", "remarcado"])
    horas = 0
    for e in eventos:
        if e[6] == "evolucao_ok":
            try:
                ini = datetime.strptime(e[2], "%H:%M"); fim_h = datetime.strptime(e[3], "%H:%M")
                horas += (fim_h - ini).seconds / 3600
            except Exception:
                pass
    return {"total": len(eventos), "realizados": realizados, "pendentes": pendentes, "cancelados": cancelados, "horas": round(horas, 1)}

def atualizar_evento_status(user_id, event_id, status):
    conn = db(); c = conn.cursor()
    c.execute("UPDATE agenda_eventos SET status=?, atualizado_em=? WHERE id=? AND user_id=?", (status, datetime.now().isoformat(timespec="seconds"), event_id, user_id))
    conn.commit(); conn.close()

def cor_status(status, data_txt):
    try:
        passado = datetime.strptime(data_txt, "%Y-%m-%d").date() <= date.today()
    except Exception:
        passado = False
    if status == "evolucao_ok": return "#22c55e", "Evolução OK"
    if status == "cancelado": return "#64748b", "Cancelado"
    if status == "falta_justificada": return "#f59e0b", "Falta justificada"
    if status == "remarcado": return "#94a3b8", "Remarcado"
    if passado: return "#ef4444", "Evolução pendente"
    return "#38bdf8", "Previsto"


# ==========================================================
# MODO CLÍNICA — GESTÃO DE PROFISSIONAIS, SETORES E AVALIAÇÕES
# ==========================================================
def listar_setores_clinica(gestor_id, apenas_ativos=True):
    conn = db(); c = conn.cursor()
    sql = "SELECT id,nome,area,descricao,ativo,criado_em FROM clinica_setores WHERE gestor_id=?"
    params=[gestor_id]
    if apenas_ativos:
        sql += " AND ativo=1"
    sql += " ORDER BY area,nome"
    c.execute(sql, params)
    rows=c.fetchall(); conn.close(); return rows

def salvar_setor_clinica(gestor_id, nome, area, descricao):
    conn=db(); c=conn.cursor(); now=datetime.now().isoformat(timespec="seconds")
    c.execute("""INSERT INTO clinica_setores (gestor_id,nome,area,descricao,ativo,criado_em,atualizado_em)
              VALUES (?,?,?,?,1,?,?)""", (gestor_id,nome,area,descricao,now,now))
    conn.commit(); conn.close()

def atualizar_status_setor(gestor_id, setor_id, ativo):
    conn=db(); c=conn.cursor()
    c.execute("UPDATE clinica_setores SET ativo=?, atualizado_em=? WHERE id=? AND gestor_id=?", (1 if ativo else 0, datetime.now().isoformat(timespec="seconds"), setor_id, gestor_id))
    conn.commit(); conn.close()

def listar_profissionais_clinica(gestor_id):
    conn=db(); c=conn.cursor()
    c.execute("""SELECT p.id,p.nome,p.email,p.profissao,p.registro,p.status,s.nome,s.area,p.observacoes
              FROM clinica_profissionais p LEFT JOIN clinica_setores s ON s.id=p.setor_id
              WHERE p.gestor_id=? ORDER BY p.status,p.nome""", (gestor_id,))
    rows=c.fetchall(); conn.close(); return rows

def salvar_profissional_clinica(gestor_id, nome, email, profissao, registro, setor_id, observacoes):
    conn=db(); c=conn.cursor(); now=datetime.now().isoformat(timespec="seconds")
    c.execute("""INSERT INTO clinica_profissionais (gestor_id,nome,email,profissao,registro,setor_id,status,observacoes,criado_em,atualizado_em)
              VALUES (?,?,?,?,?,?, 'ativo', ?, ?, ?)""", (gestor_id,nome,email,profissao,registro,setor_id,observacoes,now,now))
    conn.commit(); conn.close()

def atualizar_status_profissional(gestor_id, prof_id, status):
    conn=db(); c=conn.cursor()
    c.execute("UPDATE clinica_profissionais SET status=?, atualizado_em=? WHERE id=? AND gestor_id=?", (status, datetime.now().isoformat(timespec="seconds"), prof_id, gestor_id))
    conn.commit(); conn.close()

def listar_avaliacoes_clinica(gestor_id, setor_id=None):
    conn=db(); c=conn.cursor()
    if setor_id:
        c.execute("""SELECT m.id,m.nome,m.profissao,m.avaliacao_detectada,m.descricao,m.criado_em,s.nome,s.area,m.fonte_referencia
                  FROM avaliacao_modelos m LEFT JOIN clinica_setores s ON s.id=m.setor_id
                  WHERE m.gestor_id=? AND m.escopo='clinica' AND m.ativo=1 AND m.setor_id=? ORDER BY m.id DESC""", (gestor_id,setor_id))
    else:
        c.execute("""SELECT m.id,m.nome,m.profissao,m.avaliacao_detectada,m.descricao,m.criado_em,s.nome,s.area,m.fonte_referencia
                  FROM avaliacao_modelos m LEFT JOIN clinica_setores s ON s.id=m.setor_id
                  WHERE m.gestor_id=? AND m.escopo='clinica' AND m.ativo=1 ORDER BY s.area,s.nome,m.nome""", (gestor_id,))
    rows=c.fetchall(); conn.close(); return rows

def desativar_avaliacao_clinica(gestor_id, modelo_id):
    conn=db(); c=conn.cursor()
    c.execute("UPDATE avaliacao_modelos SET ativo=0, atualizado_em=? WHERE id=? AND gestor_id=? AND escopo='clinica'", (datetime.now().isoformat(timespec="seconds"), modelo_id, gestor_id))
    conn.commit(); conn.close()

def tela_gestao_clinica():
    st.markdown("### Gestão da clínica")
    st.caption("Cadastre setores, profissionais e avaliações padrão por área. Essas avaliações alimentam a construção dos relatórios dos profissionais vinculados ao Modo Clínica.")
    if not modo_clinica_liberado(USER):
        st.info("Esse painel é exclusivo do plano Clinic ou Enterprise / Modo Clínica.")
        return

    tab_s, tab_p, tab_av = st.tabs(["Setores por área", "Profissionais", "Avaliações padrão por setor"])
    with tab_s:
        card_inicio("Cadastrar setor", "Ex.: Musicoterapia, Fonoaudiologia, Terapia Ocupacional, Psicologia, ABA, Psicomotricidade.")
        c1,c2 = st.columns(2)
        with c1:
            nome_setor = st.text_input("Nome do setor", key="clin_setor_nome")
            area_setor = st.selectbox("Área profissional", PROFISSOES, key="clin_setor_area")
        with c2:
            desc_setor = st.text_area("Descrição / foco do setor", height=120, key="clin_setor_desc")
        if st.button("Criar setor", key="clin_criar_setor"):
            if not nome_setor.strip():
                st.warning("Informe o nome do setor.")
            else:
                salvar_setor_clinica(USER_ID, nome_setor.strip(), area_setor, desc_setor)
                st.success("Setor cadastrado."); st.rerun()
        card_fim()
        setores = listar_setores_clinica(USER_ID, apenas_ativos=False)
        if setores:
            st.markdown("#### Setores cadastrados")
            for sid,nome,area,desc,ativo,criado in setores:
                with st.container(border=True):
                    c1,c2,c3 = st.columns([0.48,0.26,0.26])
                    c1.markdown(f"**{nome}**  \n{area}  \n<span style='color:#94a3b8'>{desc or ''}</span>", unsafe_allow_html=True)
                    c2.caption("Ativo" if ativo else "Inativo")
                    with c3:
                        novo = st.toggle("Ativar", value=bool(ativo), key=f"clin_setor_ativo_{sid}")
                        if novo != bool(ativo):
                            atualizar_status_setor(USER_ID, sid, novo); st.rerun()
        else:
            st.info("Nenhum setor cadastrado ainda.")

    with tab_p:
        setores = listar_setores_clinica(USER_ID)
        if not setores:
            st.warning("Cadastre pelo menos um setor antes de cadastrar profissionais.")
        else:
            mapa_setores = {f"{r[1]} — {r[2]}": r[0] for r in setores}
            card_inicio("Cadastrar profissional da clínica", "O gestor organiza a equipe por setor e área. Em uma fase futura, esse cadastro pode virar convite de acesso ao app.")
            c1,c2,c3 = st.columns(3)
            with c1:
                nome_prof = st.text_input("Nome do profissional", key="clin_prof_nome")
                email_prof = st.text_input("E-mail", key="clin_prof_email")
            with c2:
                prof_area = st.selectbox("Profissão", PROFISSOES, key="clin_prof_area")
                registro_prof = st.text_input("Registro / conselho", key="clin_prof_reg")
            with c3:
                setor_label = st.selectbox("Setor", list(mapa_setores.keys()), key="clin_prof_setor")
                obs_prof = st.text_area("Observações", height=92, key="clin_prof_obs")
            if st.button("Cadastrar profissional", key="clin_prof_salvar"):
                if not nome_prof.strip():
                    st.warning("Informe o nome do profissional.")
                else:
                    salvar_profissional_clinica(USER_ID, nome_prof.strip(), email_prof.strip(), prof_area, registro_prof.strip(), mapa_setores[setor_label], obs_prof)
                    st.success("Profissional cadastrado."); st.rerun()
            card_fim()
        profissionais = listar_profissionais_clinica(USER_ID)
        if profissionais:
            st.markdown("#### Equipe cadastrada")
            for pid,nome,email,prof,reg,status,setor,area,obs in profissionais:
                with st.container(border=True):
                    c1,c2,c3 = st.columns([0.48,0.28,0.24])
                    c1.markdown(f"**{nome}**  \n{email or 'sem e-mail'}  \n{prof} • {reg or 'sem registro'}")
                    c2.markdown(f"Setor: **{setor or 'não definido'}**  \nÁrea: {area or '-'}")
                    with c3:
                        novo_status = st.selectbox("Status", ["ativo","pausado","desligado"], index=["ativo","pausado","desligado"].index(status if status in ["ativo","pausado","desligado"] else "ativo"), key=f"clin_prof_status_{pid}")
                        if novo_status != status:
                            atualizar_status_profissional(USER_ID, pid, novo_status); st.rerun()

    with tab_av:
        setores = listar_setores_clinica(USER_ID)
        if not setores:
            st.warning("Cadastre setores antes de subir avaliações padrão.")
            return
        mapa_setores = {f"{r[1]} — {r[2]}": r for r in setores}
        card_inicio("Upload de avaliação padrão do setor", "O arquivo será reconhecido, convertido em modelo preenchível e salvo como padrão da área selecionada.")
        setor_label = st.selectbox("Setor que usará esta avaliação", list(mapa_setores.keys()), key="clin_av_setor")
        setor = mapa_setores[setor_label]
        sid, setor_nome, setor_area = setor[0], setor[1], setor[2]
        c1,c2 = st.columns([0.55,0.45])
        with c1:
            arquivos = st.file_uploader("Enviar avaliação padrão", type=["pdf","docx","xlsx","xls","csv","txt","md","png","jpg","jpeg"], accept_multiple_files=True, key="clin_av_upload")
        with c2:
            fonte = st.text_input("Fonte/referência do modelo", placeholder="Ex.: protocolo próprio, autorização interna, referência pública etc.", key="clin_av_fonte")
            nota = st.text_area("Observação para a equipe", height=110, key="clin_av_nota")
        if arquivos:
            previews=[]
            for i, arq in enumerate(arquivos):
                texto = texto_arquivo_upload(arq)
                av_detectada, secoes = montar_modelo_automatizado(arq.name, texto, setor_area)
                grafico = especificacao_grafico_avaliacao(av_detectada, texto, setor_area)
                st.markdown(f"#### {arq.name}")
                c1,c2 = st.columns(2)
                with c1:
                    nome_modelo = st.text_input("Nome no sistema", value=av_detectada, key=f"clin_av_nome_{i}")
                    av_conf = st.text_input("Avaliação reconhecida", value=av_detectada, key=f"clin_av_rec_{i}")
                with c2:
                    desc = st.text_area("Justificativa/descrição", value=justificativa_avaliacao(av_detectada, setor_area) + ("\n\n" + nota if nota else ""), height=130, key=f"clin_av_desc_{i}")
                st.caption("Domínios do gráfico: " + ", ".join([d.get("nome", "") for d in grafico.get("dominios", [])]))
                previews.append((arq,nome_modelo,av_conf,desc,secoes,grafico))
            if st.button("Salvar avaliações como padrão deste setor", key="clin_av_salvar"):
                for arq,nome_modelo,av_conf,desc,secoes,grafico in previews:
                    safe_name = f"clinica_{USER_ID}_{sid}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{arq.name}".replace("/","_").replace("\\","_")
                    caminho = os.path.join(UPLOAD_DIR, safe_name)
                    with open(caminho,"wb") as f:
                        f.write(arq.getbuffer())
                    salvar_modelo_avaliacao(USER_ID, nome_modelo or av_conf, setor_area, av_conf, desc, secoes, arq.name, caminho, grafico, fonte, gestor_id=USER_ID, setor_id=sid, escopo="clinica", area_alvo=setor_area)
                st.success("Avaliações padrão do setor salvas."); st.rerun()
        card_fim()
        avals = listar_avaliacoes_clinica(USER_ID)
        if avals:
            st.markdown("#### Biblioteca padrão da clínica")
            for mid,nome,prof,av,desc,criado,setor_nome,area,fonte_ref in avals:
                with st.container(border=True):
                    c1,c2,c3 = st.columns([0.55,0.25,0.20])
                    c1.markdown(f"**{nome}**  \n{av}  \n<span style='color:#94a3b8'>{(desc or '')[:220]}</span>", unsafe_allow_html=True)
                    c2.markdown(f"Setor: **{setor_nome or '-'}**  \nÁrea: {area or prof}  \nFonte: {fonte_ref or '-'}")
                    with c3:
                        if st.button("Desativar", key=f"clin_av_del_{mid}"):
                            desativar_avaliacao_clinica(USER_ID, mid); st.rerun()
        else:
            st.info("Nenhuma avaliação padrão de setor cadastrada ainda.")

def tela_agenda_clinica():
    st.markdown("## Agenda clínica e evoluções obrigatórias")
    if not modo_clinica_liberado(USER):
        st.info("Modo Clínica disponível no plano Clinic ou Enterprise: agenda recorrente, carga horária personalizável, controle de evolução pendente e indicadores.")
        if st.button("Ver planos", key="agenda_ver_planos"):
            go_to("Planos e acesso")
        return

    cfg = get_agenda_config(USER_ID)
    tabs = st.tabs(["Calendário", "Carga horária", "Agenda dos pacientes", "Indicadores", "Gestão clínica"])

    with tabs[0]:
        hoje = date.today()
        st.markdown("### Calendário clínico")
        cview, cyear, cmonth = st.columns([0.32, 0.22, 0.46])
        with cview:
            visualizacao = st.radio("Visualização", ["Semana", "Mês", "Ano inteiro"], horizontal=True, key="agenda_visualizacao")
        with cyear:
            ano_ref = st.number_input("Ano", min_value=2020, max_value=2100, value=hoje.year, step=1, key="agenda_ano_ref")
        with cmonth:
            mes_nome = st.selectbox("Mês", [nome_mes_pt(i) for i in range(1, 13)], index=hoje.month-1, key="agenda_mes_ref")
        mes_ref = [nome_mes_pt(i) for i in range(1, 13)].index(mes_nome) + 1

        if visualizacao == "Semana":
            data_base = st.date_input("Semana de referência", value=hoje, key="agenda_semana_ref")
            inicio_periodo = data_base - timedelta(days=data_base.weekday())
            fim_periodo = inicio_periodo + timedelta(days=6)
            titulo_periodo = f"Semana de {inicio_periodo.strftime('%d/%m/%Y')} a {fim_periodo.strftime('%d/%m/%Y')}"
        elif visualizacao == "Mês":
            inicio_periodo, fim_periodo = inicio_fim_mes(date(int(ano_ref), int(mes_ref), 1))
            titulo_periodo = f"{nome_mes_pt(int(mes_ref))} de {int(ano_ref)}"
        else:
            inicio_periodo = date(int(ano_ref), 1, 1)
            fim_periodo = date(int(ano_ref), 12, 31)
            titulo_periodo = f"Ano de {int(ano_ref)}"

        gerar_eventos_periodo_recorrentes(USER_ID, inicio_periodo, fim_periodo)
        eventos = listar_eventos_periodo(USER_ID, inicio_periodo, fim_periodo)
        cont = eventos_contadores(USER_ID, inicio_periodo, fim_periodo)
        c1,c2,c3,c4 = st.columns(4)
        c1.metric("Atendimentos no período", cont["total"])
        c2.metric("Evoluções OK", cont["realizados"])
        c3.metric("Pendentes", cont["pendentes"])
        c4.metric("Horas realizadas", cont["horas"])
        st.caption("Verde = evolução realizada. Vermelho = atendimento vencido sem evolução. Azul = futuro. Cinza/amarelo = cancelamento, remarcação ou falta.")
        st.markdown(f"### {titulo_periodo}")

        if visualizacao == "Semana":
            cols = st.columns(7)
            for i, col in enumerate(cols):
                dia = inicio_periodo + timedelta(days=i)
                with col:
                    st.markdown(f"**{DIAS_SEMANA[i]}**")
                    st.caption(dia.strftime("%d/%m"))
                    eventos_dia = [e for e in eventos if e[1] == dia.strftime("%Y-%m-%d")]
                    if not eventos_dia:
                        st.caption("Sem atendimentos")
                    for e in eventos_dia:
                        render_evento_card(e)
                        render_acoes_evento(e, "semana")

        elif visualizacao == "Mês":
            cal = calendar.Calendar(firstweekday=0)
            semanas = cal.monthdatescalendar(int(ano_ref), int(mes_ref))
            header = st.columns(7)
            for i, col in enumerate(header):
                with col:
                    st.markdown(f"<div style='text-align:center;color:#9fb0cf;font-weight:800;'>{DIAS_SEMANA[i][:3]}</div>", unsafe_allow_html=True)
            for sidx, semana in enumerate(semanas):
                cols = st.columns(7)
                for i, dia in enumerate(semana):
                    with cols[i]:
                        fora_mes = dia.month != int(mes_ref)
                        bg = "rgba(15,23,42,.38)" if not fora_mes else "rgba(15,23,42,.12)"
                        borda = "rgba(148,163,184,.14)" if not fora_mes else "rgba(148,163,184,.06)"
                        st.markdown(f"""
                        <div style="background:{bg};border:1px solid {borda};border-radius:16px;padding:10px;min-height:150px;margin-bottom:8px;">
                            <strong style="color:{'#f8fafc' if not fora_mes else '#475569'};">{dia.day}</strong>
                        </div>
                        """, unsafe_allow_html=True)
                        eventos_dia = [e for e in eventos if e[1] == dia.strftime("%Y-%m-%d")]
                        for e in eventos_dia[:3]:
                            render_evento_card(e, compact=True)
                            render_acoes_evento(e, f"mes_{sidx}_{i}")
                        if len(eventos_dia) > 3:
                            st.caption(f"+ {len(eventos_dia)-3} atendimentos")

        else:
            st.markdown("#### Calendário anual")
            meses = list(range(1, 13))
            for linha in range(4):
                cols = st.columns(3)
                for j in range(3):
                    mes = meses[linha*3 + j]
                    with cols[j]:
                        ini_m, fim_m = inicio_fim_mes(date(int(ano_ref), mes, 1))
                        eventos_mes = [e for e in eventos if ini_m.strftime("%Y-%m-%d") <= e[1] <= fim_m.strftime("%Y-%m-%d")]
                        realizados = sum(1 for e in eventos_mes if e[6] == "evolucao_ok")
                        pendentes = sum(1 for e in eventos_mes if e[6] in ["previsto", "realizado_sem_evolucao"] and datetime.strptime(e[1], "%Y-%m-%d").date() <= date.today())
                        st.markdown(f"""
                        <div style="background:linear-gradient(145deg,rgba(13,22,39,.90),rgba(17,25,42,.62));border:1px solid rgba(148,163,184,.14);border-radius:18px;padding:18px;margin-bottom:14px;min-height:180px;">
                            <h4 style="margin:0 0 10px 0;color:#f8fafc;">{nome_mes_pt(mes)}</h4>
                            <div style="color:#9fb0cf;font-size:.9rem;">Atendimentos: <strong style="color:#f8fafc;">{len(eventos_mes)}</strong></div>
                            <div style="color:#22c55e;font-size:.9rem;">Evoluções OK: {realizados}</div>
                            <div style="color:#ef4444;font-size:.9rem;">Pendentes: {pendentes}</div>
                        </div>
                        """, unsafe_allow_html=True)
                        eventos_preview = sorted(eventos_mes, key=lambda x: (x[1], x[2]))[:3]
                        for e in eventos_preview:
                            cor, label = cor_status(e[6], e[1])
                            st.markdown(f"<div style='border-left:4px solid {cor};padding:6px 8px;margin:5px 0;background:rgba(15,23,42,.48);border-radius:10px;font-size:.78rem;'><strong>{datetime.strptime(e[1], '%Y-%m-%d').strftime('%d/%m')}</strong> • {e[2]} • {e[4]}</div>", unsafe_allow_html=True)
                        if len(eventos_mes) > 3:
                            st.caption(f"+ {len(eventos_mes)-3} eventos no mês")

        render_form_evolucao_agenda()

    with tabs[1]:
        card_inicio("Carga horária do terapeuta", "Personalize sua rotina. Essa configuração pode ser alterada a qualquer momento.")
        dias_label = {v:k for k,v in DIAS_SEMANA.items()}
        selecionados = st.multiselect("Dias de atendimento", list(dias_label.keys()), default=[DIAS_SEMANA[d] for d in cfg["dias_semana"]], key="cfg_dias")
        c1,c2,c3,c4 = st.columns(4)
        with c1: hi = st.time_input("Início", value=parse_time_str(cfg["hora_inicio"]), key="cfg_hi")
        with c2: hf = st.time_input("Fim", value=parse_time_str(cfg["hora_fim"]), key="cfg_hf")
        with c3: ai = st.time_input("Início almoço", value=parse_time_str(cfg["almoco_inicio"], "12:00"), key="cfg_ai")
        with c4: af = st.time_input("Fim almoço", value=parse_time_str(cfg["almoco_fim"], "13:00"), key="cfg_af")
        c5,c6 = st.columns(2)
        with c5: dur = st.number_input("Duração padrão da sessão", 20, 180, int(cfg["duracao_padrao"] or 50), step=5, key="cfg_dur")
        with c6: inter = st.number_input("Intervalo padrão", 0, 60, int(cfg["intervalo_padrao"] or 10), step=5, key="cfg_inter")
        if st.button("Salvar carga horária", key="salvar_carga_horaria"):
            salvar_agenda_config(USER_ID, [dias_label[d] for d in selecionados], time_to_str(hi), time_to_str(hf), time_to_str(ai), time_to_str(af), dur, inter)
            st.success("Carga horária atualizada.")
            st.rerun()
        card_fim()

    with tabs[2]:
        card_inicio("Agenda recorrente dos pacientes", "Defina dia, horário e frequência para o paciente aparecer automaticamente no calendário.")
        pacientes = listar_pacientes(USER_ID)
        if not pacientes:
            st.info("Cadastre pacientes primeiro.")
        else:
            conn = db(); c = conn.cursor()
            for pac in pacientes:
                pid, nome, idade, nasc, diag, esc, resp = pac[:7]
                c.execute("SELECT dia_semana, hora_atendimento, duracao_sessao, frequencia, modalidade, agenda_status FROM pacientes WHERE id=? AND user_id=?", (pid, USER_ID))
                row = c.fetchone() or (None, "", 50, "Semanal", "Presencial", "ativo")
                with st.container(border=True):
                    st.markdown(f"**{nome}** — {diag or 'sem diagnóstico'}")
                    c1,c2,c3,c4,c5 = st.columns(5)
                    with c1:
                        dia_nome = st.selectbox("Dia", list(DIAS_SEMANA.values()), index=int(row[0]) if row[0] is not None else 0, key=f"agenda_p_dia_{pid}")
                    with c2:
                        hora = st.time_input("Horário", value=parse_time_str(row[1] or cfg["hora_inicio"]), key=f"agenda_p_hora_{pid}")
                    with c3:
                        durp = st.number_input("Duração", 20, 180, int(row[2] or 50), step=5, key=f"agenda_p_dur_{pid}")
                    with c4:
                        freq = st.selectbox("Frequência", ["Semanal", "Quinzenal", "Mensal"], index=["Semanal","Quinzenal","Mensal"].index(row[3] if row[3] in ["Semanal","Quinzenal","Mensal"] else "Semanal"), key=f"agenda_p_freq_{pid}")
                    with c5:
                        mod = st.selectbox("Modalidade", ["Presencial", "Online", "Domiciliar"], index=["Presencial","Online","Domiciliar"].index(row[4] if row[4] in ["Presencial","Online","Domiciliar"] else "Presencial"), key=f"agenda_p_mod_{pid}")
                    if st.button("Salvar agenda deste paciente", key=f"salvar_agenda_p_{pid}"):
                        dia_idx = [k for k,v in DIAS_SEMANA.items() if v == dia_nome][0]
                        c.execute("""UPDATE pacientes SET dia_semana=?, hora_atendimento=?, duracao_sessao=?, frequencia=?, modalidade=?, agenda_status='ativo' WHERE id=? AND user_id=?""",
                                  (dia_idx, time_to_str(hora), durp, freq, mod, pid, USER_ID))
                        conn.commit(); st.success("Agenda do paciente atualizada."); st.rerun()
            conn.close()
        card_fim()

    with tabs[3]:
        hoje = date.today(); inicio_mes = hoje.replace(day=1); fim_mes = (inicio_mes + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        gerar_eventos_semana(USER_ID, hoje - timedelta(days=hoje.weekday()))
        cont = eventos_contadores(USER_ID, inicio_mes, fim_mes)
        st.markdown("### Indicadores do mês")
        c1,c2,c3,c4,c5 = st.columns(5)
        c1.metric("Atendimentos previstos", cont["total"])
        c2.metric("Atendimentos feitos", cont["realizados"])
        c3.metric("Evoluções pendentes", cont["pendentes"])
        c4.metric("Cancelamentos/faltas", cont["cancelados"])
        c5.metric("Horas clínicas", cont["horas"])
        st.caption("Verde = evolução realizada. Vermelho = atendimento vencido sem evolução. Azul = futuro. Cinza/amarelo = cancelamento, remarcação ou falta.")

    with tabs[4]:
        tela_gestao_clinica()

def tela_pacientes_e_evolucao():
    st.markdown("## Banco de pacientes, agenda e evolução")
    tab_p, tab_a, tab_e = st.tabs(["Pacientes", "Agenda clínica", "Evolução avulsa"])
    with tab_p:
        tela_pacientes()
    with tab_a:
        tela_agenda_clinica()
    with tab_e:
        tela_evolucao()




# ==========================================================
# MODO CLÍNICA — UPGRADE GESTÃO COMPLETA
# ==========================================================
def ensure_clinica_upgrade_schema():
    """Garante as tabelas/colunas extras do Modo Clínica sem quebrar bancos antigos."""
    add_col_if_missing("pacientes", "clinica_profissional_id", "ALTER TABLE pacientes ADD COLUMN clinica_profissional_id INTEGER DEFAULT NULL")
    add_col_if_missing("pacientes", "clinica_setor_id", "ALTER TABLE pacientes ADD COLUMN clinica_setor_id INTEGER DEFAULT NULL")
    add_col_if_missing("agenda_eventos", "clinica_profissional_id", "ALTER TABLE agenda_eventos ADD COLUMN clinica_profissional_id INTEGER DEFAULT NULL")
    conn = db(); c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS clinica_carga_horaria (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        gestor_id INTEGER,
        profissional_id INTEGER UNIQUE,
        dias_semana TEXT DEFAULT '0,1,2,3,4',
        hora_inicio TEXT DEFAULT '08:00',
        hora_fim TEXT DEFAULT '18:00',
        almoco_inicio TEXT DEFAULT '12:00',
        almoco_fim TEXT DEFAULT '13:00',
        duracao_padrao INTEGER DEFAULT 50,
        intervalo_padrao INTEGER DEFAULT 10,
        observacoes TEXT DEFAULT '',
        atualizado_em TEXT
    )
    """)
    conn.commit(); conn.close()

ensure_clinica_upgrade_schema()


def safe_int(v, default=None):
    try:
        if v is None or str(v).strip() == "":
            return default
        return int(v)
    except Exception:
        return default


def get_profissionais_map(gestor_id, ativos=True):
    rows = listar_profissionais_clinica(gestor_id)
    if ativos:
        rows = [r for r in rows if str(r[5] or '').lower() == 'ativo']
    mapa = {}
    for r in rows:
        pid,nome,email,prof,reg,status,setor,area,obs = r
        mapa[f"{nome} — {prof or area or 'Área'}"] = r
    return mapa


def get_carga_profissional(gestor_id, profissional_id):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT dias_semana,hora_inicio,hora_fim,almoco_inicio,almoco_fim,duracao_padrao,intervalo_padrao,observacoes
              FROM clinica_carga_horaria WHERE gestor_id=? AND profissional_id=?""", (gestor_id, profissional_id))
    row = c.fetchone()
    if not row:
        c.execute("""INSERT OR IGNORE INTO clinica_carga_horaria
                  (gestor_id,profissional_id,dias_semana,hora_inicio,hora_fim,almoco_inicio,almoco_fim,duracao_padrao,intervalo_padrao,observacoes,atualizado_em)
                  VALUES (?,?,?,?,?,?,?,?,?,?,?)""", (gestor_id, profissional_id, "0,1,2,3,4", "08:00", "18:00", "12:00", "13:00", 50, 10, "", datetime.now().isoformat(timespec="seconds")))
        conn.commit()
        row = ("0,1,2,3,4", "08:00", "18:00", "12:00", "13:00", 50, 10, "")
    conn.close()
    dias = [int(x) for x in str(row[0] or "0,1,2,3,4").split(",") if str(x).strip().isdigit()]
    return {"dias_semana": dias, "hora_inicio": row[1], "hora_fim": row[2], "almoco_inicio": row[3], "almoco_fim": row[4], "duracao_padrao": row[5], "intervalo_padrao": row[6], "observacoes": row[7] or ""}


def salvar_carga_profissional(gestor_id, profissional_id, dias_semana, hora_inicio, hora_fim, almoco_inicio, almoco_fim, duracao_padrao, intervalo_padrao, observacoes=""):
    conn = db(); c = conn.cursor(); now = datetime.now().isoformat(timespec="seconds")
    dias_txt = ",".join(str(d) for d in dias_semana)
    c.execute("""INSERT INTO clinica_carga_horaria
              (gestor_id,profissional_id,dias_semana,hora_inicio,hora_fim,almoco_inicio,almoco_fim,duracao_padrao,intervalo_padrao,observacoes,atualizado_em)
              VALUES (?,?,?,?,?,?,?,?,?,?,?)
              ON CONFLICT(profissional_id) DO UPDATE SET
              gestor_id=excluded.gestor_id, dias_semana=excluded.dias_semana, hora_inicio=excluded.hora_inicio,
              hora_fim=excluded.hora_fim, almoco_inicio=excluded.almoco_inicio, almoco_fim=excluded.almoco_fim,
              duracao_padrao=excluded.duracao_padrao, intervalo_padrao=excluded.intervalo_padrao,
              observacoes=excluded.observacoes, atualizado_em=excluded.atualizado_em""",
              (gestor_id, profissional_id, dias_txt, hora_inicio, hora_fim, almoco_inicio, almoco_fim, int(duracao_padrao), int(intervalo_padrao), observacoes, now))
    conn.commit(); conn.close()


def listar_pacientes_clinica(gestor_id):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT p.id,p.nome,p.idade,p.nascimento,p.diagnostico,p.escolaridade,p.responsaveis,
                     p.dia_semana,p.hora_atendimento,p.duracao_sessao,p.frequencia,p.modalidade,p.agenda_status,
                     p.clinica_profissional_id,p.clinica_setor_id, cp.nome, cp.profissao, cs.nome, cs.area
              FROM pacientes p
              LEFT JOIN clinica_profissionais cp ON cp.id=p.clinica_profissional_id
              LEFT JOIN clinica_setores cs ON cs.id=p.clinica_setor_id
              WHERE p.user_id=? ORDER BY p.agenda_status DESC, COALESCE(cp.nome,''), p.nome""", (gestor_id,))
    rows = c.fetchall(); conn.close(); return rows


def salvar_paciente_clinica(gestor_id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis, profissional_id=None, setor_id=None):
    salvar_paciente(gestor_id, nome, idade, nascimento, diagnostico, escolaridade, responsaveis)
    conn = db(); c = conn.cursor()
    c.execute("SELECT id FROM pacientes WHERE user_id=? ORDER BY id DESC LIMIT 1", (gestor_id,))
    row = c.fetchone()
    if row:
        c.execute("UPDATE pacientes SET clinica_profissional_id=?, clinica_setor_id=?, agenda_status='ativo' WHERE id=? AND user_id=?", (profissional_id, setor_id, row[0], gestor_id))
    conn.commit(); conn.close()


def atualizar_paciente_clinica(gestor_id, paciente_id, profissional_id=None, setor_id=None, agenda_status="ativo"):
    conn = db(); c = conn.cursor()
    c.execute("""UPDATE pacientes SET clinica_profissional_id=?, clinica_setor_id=?, agenda_status=?
              WHERE id=? AND user_id=?""", (profissional_id, setor_id, agenda_status, paciente_id, gestor_id))
    c.execute("UPDATE agenda_eventos SET clinica_profissional_id=? WHERE user_id=? AND paciente_id=? AND status IN ('previsto','realizado_sem_evolucao')", (profissional_id, gestor_id, paciente_id))
    conn.commit(); conn.close()


def atualizar_agenda_paciente_clinica(gestor_id, paciente_id, dia_semana, hora, duracao, frequencia, modalidade, profissional_id=None, setor_id=None):
    conn = db(); c = conn.cursor()
    c.execute("""UPDATE pacientes SET dia_semana=?, hora_atendimento=?, duracao_sessao=?, frequencia=?, modalidade=?,
              clinica_profissional_id=COALESCE(?, clinica_profissional_id), clinica_setor_id=COALESCE(?, clinica_setor_id), agenda_status='ativo'
              WHERE id=? AND user_id=?""", (dia_semana, hora, duracao, frequencia, modalidade, profissional_id, setor_id, paciente_id, gestor_id))
    conn.commit(); conn.close()


def profissionais_setores_labels(gestor_id):
    profs = get_profissionais_map(gestor_id, ativos=False)
    setores = listar_setores_clinica(gestor_id)
    prof_label_to_id = {label: row[0] for label, row in profs.items()}
    prof_id_to_label = {row[0]: label for label, row in profs.items()}
    prof_id_to_setor = {row[0]: row[0] and None for label, row in profs.items()}
    setor_label_to_id = {f"{r[1]} — {r[2]}": r[0] for r in setores}
    setor_id_to_label = {r[0]: f"{r[1]} — {r[2]}" for r in setores}
    # setor do profissional vem oculto no SELECT: índice 6 é nome do setor e 7 é área; buscamos o ID direto quando necessário.
    return profs, setores, prof_label_to_id, prof_id_to_label, setor_label_to_id, setor_id_to_label


def obter_setor_profissional(gestor_id, profissional_id):
    conn = db(); c = conn.cursor()
    c.execute("SELECT setor_id FROM clinica_profissionais WHERE id=? AND gestor_id=?", (profissional_id, gestor_id))
    row = c.fetchone(); conn.close()
    return row[0] if row else None


def pacientes_com_agenda_clinica(gestor_id):
    conn = db(); c = conn.cursor()
    c.execute("""SELECT p.id,p.nome,p.diagnostico,p.dia_semana,p.hora_atendimento,p.duracao_sessao,p.frequencia,p.modalidade,p.agenda_status,p.clinica_profissional_id
              FROM pacientes p
              WHERE p.user_id=? AND p.agenda_status='ativo' AND p.dia_semana IS NOT NULL AND p.hora_atendimento IS NOT NULL AND p.hora_atendimento!=''
              ORDER BY p.clinica_profissional_id, p.dia_semana, p.hora_atendimento""", (gestor_id,))
    rows = c.fetchall(); conn.close(); return rows


def ensure_evento_recorrente_clinica(gestor_id, paciente, data_obj):
    pid, nome, diag, dia_semana, hora, dur, freq, modalidade, status, prof_id = paciente
    if dia_semana is None or int(dia_semana) != data_obj.weekday():
        return None
    if freq == "Quinzenal" and (data_obj.isocalendar().week % 2 != 0):
        return None
    if freq == "Mensal" and data_obj.day > 7:
        return None
    inicio_dt = datetime.combine(data_obj, parse_time_str(hora))
    fim_dt = inicio_dt + timedelta(minutes=int(dur or 50))
    conn = db(); c = conn.cursor()
    c.execute("""SELECT id FROM agenda_eventos WHERE user_id=? AND paciente_id=? AND data=? AND hora_inicio=? AND origem='recorrente'""",
              (gestor_id, pid, data_obj.strftime("%Y-%m-%d"), hora))
    existing = c.fetchone()
    if existing:
        c.execute("UPDATE agenda_eventos SET clinica_profissional_id=COALESCE(?, clinica_profissional_id), atualizado_em=? WHERE id=?", (prof_id, datetime.now().isoformat(timespec="seconds"), existing[0]))
        conn.commit(); conn.close(); return existing[0]
    c.execute("""INSERT INTO agenda_eventos (user_id, paciente_id, data, hora_inicio, hora_fim, titulo, modalidade, status, origem, clinica_profissional_id, criado_em, atualizado_em)
              VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (gestor_id, pid, data_obj.strftime("%Y-%m-%d"), hora, fim_dt.strftime("%H:%M"), nome, modalidade or "Presencial", "previsto", "recorrente", prof_id, datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")))
    event_id = c.lastrowid
    conn.commit(); conn.close(); return event_id


def gerar_eventos_periodo_clinica(gestor_id, data_inicio, data_fim):
    atual = data_inicio
    pacientes = pacientes_com_agenda_clinica(gestor_id)
    while atual <= data_fim:
        for pac in pacientes:
            ensure_evento_recorrente_clinica(gestor_id, pac, atual)
        atual += timedelta(days=1)


def listar_eventos_periodo_clinica(gestor_id, data_inicio, data_fim, profissional_id=None):
    conn = db(); c = conn.cursor()
    sql = """SELECT e.id, e.data, e.hora_inicio, e.hora_fim, e.titulo, e.modalidade, e.status, e.paciente_id, p.diagnostico,
                    e.clinica_profissional_id, cp.nome
             FROM agenda_eventos e
             LEFT JOIN pacientes p ON p.id=e.paciente_id
             LEFT JOIN clinica_profissionais cp ON cp.id=e.clinica_profissional_id
             WHERE e.user_id=? AND e.data BETWEEN ? AND ?"""
    params = [gestor_id, data_inicio.strftime("%Y-%m-%d"), data_fim.strftime("%Y-%m-%d")]
    if profissional_id:
        sql += " AND e.clinica_profissional_id=?"
        params.append(profissional_id)
    sql += " ORDER BY e.data, e.hora_inicio"
    c.execute(sql, params)
    rows = c.fetchall(); conn.close(); return rows


def eventos_contadores_clinica(gestor_id, inicio, fim, profissional_id=None):
    eventos = listar_eventos_periodo_clinica(gestor_id, inicio, fim, profissional_id)
    realizados = sum(1 for e in eventos if e[6] == "evolucao_ok")
    pendentes = 0
    cancelados = 0
    horas = 0
    for e in eventos:
        try:
            data_evento = datetime.strptime(e[1], "%Y-%m-%d").date()
        except Exception:
            data_evento = date.today()
        if e[6] in ["previsto", "realizado_sem_evolucao"] and data_evento <= date.today():
            pendentes += 1
        if e[6] in ["cancelado", "falta_justificada", "remarcado"]:
            cancelados += 1
        if e[6] == "evolucao_ok":
            try:
                ini = datetime.strptime(e[2], "%H:%M"); fim_h = datetime.strptime(e[3], "%H:%M")
                horas += (fim_h - ini).seconds / 3600
            except Exception:
                pass
    return {"total": len(eventos), "realizados": realizados, "pendentes": pendentes, "cancelados": cancelados, "horas": round(horas, 1)}


def render_mini_evento_clinica(e, prefix):
    cor, label = cor_status(e[6], e[1])
    prof = e[10] or "Sem profissional"
    st.markdown(f"""
    <div style="border-left:5px solid {cor}; background:rgba(15,23,42,.68); border:1px solid rgba(148,163,184,.14); border-radius:14px; padding:10px; margin:8px 0;">
        <strong>{e[2]} - {e[3]}</strong> • {e[4]}<br>
        <span style="color:#9fb0cf;font-size:.78rem;">{prof} • {e[5]} • {label}</span>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Evoluir", key=f"{prefix}_evoluir_clin_{e[0]}", use_container_width=True):
        st.session_state.evento_evolucao_id = e[0]
        st.session_state.evento_evolucao_paciente_id = e[7]
        st.session_state.evento_evolucao_nome = e[4]
        st.session_state.evento_evolucao_data = e[1]
        st.rerun()


def tela_gestao_clinica():
    st.markdown("## Modo Clínica")
    st.caption("Painel do gestor/admin master: equipe, setores, carga horária, avaliações padrão por área, pacientes, remanejamento e indicadores.")
    if not modo_clinica_liberado(USER):
        st.info("Esse painel é exclusivo do admin master ou dos planos Clinic/Enterprise.")
        return

    ensure_clinica_upgrade_schema()
    tab_s, tab_p, tab_ch, tab_pac, tab_av, tab_ind = st.tabs([
        "Setores e equipe", "Profissionais", "Carga horária", "Pacientes e remanejamento", "Avaliações padrão", "Atendimentos"
    ])

    with tab_s:
        card_inicio("Setores por área", "Cadastre as áreas da clínica. Depois vincule profissionais, pacientes e avaliações padrão a cada setor.")
        c1, c2 = st.columns(2)
        with c1:
            nome_setor = st.text_input("Nome do setor", placeholder="Ex.: Musicoterapia Infantil", key="clin2_setor_nome")
            area_setor = st.selectbox("Área profissional", PROFISSOES, key="clin2_setor_area")
        with c2:
            desc_setor = st.text_area("Descrição / foco do setor", height=110, key="clin2_setor_desc")
        if st.button("Criar setor", key="clin2_criar_setor"):
            if not nome_setor.strip():
                st.warning("Informe o nome do setor.")
            else:
                salvar_setor_clinica(USER_ID, nome_setor.strip(), area_setor, desc_setor.strip())
                st.success("Setor cadastrado."); st.rerun()
        card_fim()
        setores = listar_setores_clinica(USER_ID, apenas_ativos=False)
        if setores:
            dados = pd.DataFrame(setores, columns=["ID","Setor","Área","Descrição","Ativo","Criado em"])
            dados["Ativo"] = dados["Ativo"].map(lambda x: "Sim" if x else "Não")
            st.dataframe(dados, use_container_width=True, hide_index=True)
            for sid,nome,area,desc,ativo,criado in setores:
                with st.container(border=True):
                    c1,c2,c3 = st.columns([0.48,0.30,0.22])
                    c1.markdown(f"**{nome}**  \n{area}  \n<span style='color:#94a3b8'>{desc or ''}</span>", unsafe_allow_html=True)
                    c2.caption("Ativo" if ativo else "Inativo")
                    novo = c3.toggle("Ativar", value=bool(ativo), key=f"clin2_setor_ativo_{sid}")
                    if novo != bool(ativo):
                        atualizar_status_setor(USER_ID, sid, novo); st.rerun()
        else:
            st.info("Nenhum setor cadastrado ainda.")

    with tab_p:
        setores = listar_setores_clinica(USER_ID)
        mapa_setores = {f"{r[1]} — {r[2]}": r[0] for r in setores}
        card_inicio("Cadastrar profissional", "O gestor/admin pode organizar cada profissional por área/setor e depois controlar carga horária e atendimentos.")
        if not setores:
            st.warning("Cadastre pelo menos um setor antes de cadastrar profissionais.")
        else:
            c1,c2,c3 = st.columns(3)
            with c1:
                nome_prof = st.text_input("Nome do profissional", key="clin2_prof_nome")
                email_prof = st.text_input("E-mail", key="clin2_prof_email")
            with c2:
                prof_area = st.selectbox("Profissão", PROFISSOES, key="clin2_prof_area")
                registro_prof = st.text_input("Registro / conselho", key="clin2_prof_reg")
            with c3:
                setor_label = st.selectbox("Setor", list(mapa_setores.keys()), key="clin2_prof_setor")
                obs_prof = st.text_area("Observações", height=92, key="clin2_prof_obs")
            if st.button("Cadastrar profissional", key="clin2_prof_salvar"):
                if not nome_prof.strip():
                    st.warning("Informe o nome do profissional.")
                else:
                    salvar_profissional_clinica(USER_ID, nome_prof.strip(), email_prof.strip(), prof_area, registro_prof.strip(), mapa_setores[setor_label], obs_prof.strip())
                    st.success("Profissional cadastrado."); st.rerun()
        card_fim()

        profissionais = listar_profissionais_clinica(USER_ID)
        if profissionais:
            st.markdown("#### Equipe cadastrada")
            dfp = pd.DataFrame(profissionais, columns=["ID","Nome","E-mail","Profissão","Registro","Status","Setor","Área","Observações"])
            st.dataframe(dfp, use_container_width=True, hide_index=True)
            for pid,nome,email,prof,reg,status,setor,area,obs in profissionais:
                with st.container(border=True):
                    c1,c2,c3 = st.columns([0.45,0.30,0.25])
                    c1.markdown(f"**{nome}**  \n{email or 'sem e-mail'}  \n{prof} • {reg or 'sem registro'}")
                    c2.markdown(f"Setor: **{setor or 'não definido'}**  \nÁrea: {area or '-'}")
                    novo_status = c3.selectbox("Status", ["ativo","pausado","desligado"], index=["ativo","pausado","desligado"].index(status if status in ["ativo","pausado","desligado"] else "ativo"), key=f"clin2_prof_status_{pid}")
                    if novo_status != status:
                        atualizar_status_profissional(USER_ID, pid, novo_status); st.rerun()
        else:
            st.info("Nenhum profissional cadastrado ainda.")

    with tab_ch:
        profissionais = listar_profissionais_clinica(USER_ID)
        if not profissionais:
            st.info("Cadastre profissionais antes de editar carga horária.")
        else:
            st.markdown("### Carga horária por profissional")
            dias_label = {v:k for k,v in DIAS_SEMANA.items()}
            for pid,nome,email,prof,reg,status,setor,area,obs in profissionais:
                cfgp = get_carga_profissional(USER_ID, pid)
                with st.container(border=True):
                    st.markdown(f"#### {nome}")
                    st.caption(f"{prof} • {setor or 'sem setor'} • status: {status}")
                    selecionados = st.multiselect("Dias de atendimento", list(dias_label.keys()), default=[DIAS_SEMANA[d] for d in cfgp["dias_semana"] if d in DIAS_SEMANA], key=f"clin2_carga_dias_{pid}")
                    c1,c2,c3,c4 = st.columns(4)
                    hi = c1.time_input("Início", value=parse_time_str(cfgp["hora_inicio"]), key=f"clin2_carga_hi_{pid}")
                    hf = c2.time_input("Fim", value=parse_time_str(cfgp["hora_fim"]), key=f"clin2_carga_hf_{pid}")
                    ai = c3.time_input("Início almoço", value=parse_time_str(cfgp["almoco_inicio"], "12:00"), key=f"clin2_carga_ai_{pid}")
                    af = c4.time_input("Fim almoço", value=parse_time_str(cfgp["almoco_fim"], "13:00"), key=f"clin2_carga_af_{pid}")
                    c5,c6,c7 = st.columns([0.25,0.25,0.50])
                    dur = c5.number_input("Duração sessão", 20, 180, int(cfgp["duracao_padrao"] or 50), step=5, key=f"clin2_carga_dur_{pid}")
                    inter = c6.number_input("Intervalo", 0, 60, int(cfgp["intervalo_padrao"] or 10), step=5, key=f"clin2_carga_inter_{pid}")
                    obs_carga = c7.text_input("Observações de agenda", value=cfgp.get("observacoes", ""), key=f"clin2_carga_obs_{pid}")
                    if st.button("Salvar carga horária", key=f"clin2_salvar_carga_{pid}"):
                        salvar_carga_profissional(USER_ID, pid, [dias_label[d] for d in selecionados], time_to_str(hi), time_to_str(hf), time_to_str(ai), time_to_str(af), dur, inter, obs_carga)
                        st.success("Carga horária atualizada."); st.rerun()

    with tab_pac:
        profissionais = get_profissionais_map(USER_ID, ativos=True)
        setores = listar_setores_clinica(USER_ID)
        prof_label_to_id = {label: row[0] for label,row in profissionais.items()}
        setor_label_to_id = {f"{r[1]} — {r[2]}": r[0] for r in setores}
        card_inicio("Adicionar paciente da clínica", "Cadastre pacientes e vincule cada um a um profissional/setor. Depois é possível remanejar, pausar ou remover da agenda.")
        if not profissionais:
            st.warning("Cadastre pelo menos um profissional ativo antes de adicionar pacientes.")
        else:
            c1,c2,c3 = st.columns(3)
            with c1:
                nome_p = st.text_input("Nome do paciente", key="clin2_pac_nome")
                idade_p = st.number_input("Idade", 0, 120, 0, key="clin2_pac_idade")
            with c2:
                nasc_p = st.text_input("Nascimento", placeholder="dd/mm/aaaa", key="clin2_pac_nasc")
                diag_p = st.text_input("Diagnóstico / hipótese", key="clin2_pac_diag")
            with c3:
                prof_label = st.selectbox("Profissional responsável", list(profissionais.keys()), key="clin2_pac_prof")
                setor_id_default = obter_setor_profissional(USER_ID, prof_label_to_id[prof_label])
                setor_labels = list(setor_label_to_id.keys())
                idx_setor = 0
                for i, sl in enumerate(setor_labels):
                    if setor_label_to_id[sl] == setor_id_default:
                        idx_setor = i
                setor_label = st.selectbox("Setor", setor_labels, index=idx_setor if setor_labels else 0, key="clin2_pac_setor")
            escola_p = st.text_input("Escolaridade / ocupação", key="clin2_pac_escola")
            resp_p = st.text_area("Responsáveis / contato", height=90, key="clin2_pac_resp")
            if st.button("Adicionar paciente", key="clin2_add_paciente"):
                if not nome_p.strip():
                    st.warning("Informe o nome do paciente.")
                else:
                    salvar_paciente_clinica(USER_ID, nome_p.strip(), int(idade_p), nasc_p.strip(), diag_p.strip(), escola_p.strip(), resp_p.strip(), prof_label_to_id[prof_label], setor_label_to_id.get(setor_label))
                    st.success("Paciente adicionado à clínica."); st.rerun()
        card_fim()

        pacientes = listar_pacientes_clinica(USER_ID)
        if pacientes:
            st.markdown("### Pacientes, agenda e remanejamento")
            dfpac = pd.DataFrame(pacientes, columns=["ID","Nome","Idade","Nascimento","Diagnóstico","Escolaridade","Responsáveis","Dia","Hora","Duração","Frequência","Modalidade","Status","Profissional ID","Setor ID","Profissional","Profissão","Setor","Área"])
            st.dataframe(dfpac[["ID","Nome","Diagnóstico","Profissional","Setor","Dia","Hora","Frequência","Status"]], use_container_width=True, hide_index=True)
            prof_labels = list(prof_label_to_id.keys())
            setor_labels = list(setor_label_to_id.keys())
            for pac in pacientes:
                pid,nome,idade,nasc,diag,esc,resp,dia,hora,dur,freq,mod,status,prof_id,setor_id,prof_nome,profissao,setor_nome,area = pac
                with st.container(border=True):
                    st.markdown(f"#### {nome}")
                    st.caption(f"{diag or 'sem diagnóstico'} • Profissional atual: {prof_nome or 'não definido'} • Setor: {setor_nome or 'não definido'}")
                    c1,c2,c3,c4 = st.columns(4)
                    idx_prof = 0
                    for i, label in enumerate(prof_labels):
                        if prof_label_to_id[label] == prof_id:
                            idx_prof = i
                    novo_prof_label = c1.selectbox("Remanejar para", prof_labels, index=idx_prof if prof_labels else 0, key=f"clin2_rem_prof_{pid}")
                    idx_set = 0
                    for i, label in enumerate(setor_labels):
                        if setor_label_to_id[label] == setor_id:
                            idx_set = i
                    novo_setor_label = c2.selectbox("Setor", setor_labels, index=idx_set if setor_labels else 0, key=f"clin2_rem_setor_{pid}")
                    novo_status = c3.selectbox("Status", ["ativo", "pausado", "removido"], index=["ativo","pausado","removido"].index(status if status in ["ativo","pausado","removido"] else "ativo"), key=f"clin2_rem_status_{pid}")
                    if c4.button("Salvar remanejamento", key=f"clin2_salvar_rem_{pid}"):
                        atualizar_paciente_clinica(USER_ID, pid, prof_label_to_id.get(novo_prof_label), setor_label_to_id.get(novo_setor_label), novo_status)
                        st.success("Paciente atualizado."); st.rerun()

                    c5,c6,c7,c8,c9 = st.columns(5)
                    dia_nome = c5.selectbox("Dia", list(DIAS_SEMANA.values()), index=int(dia) if dia is not None else 0, key=f"clin2_ag_dia_{pid}")
                    hora_p = c6.time_input("Horário", value=parse_time_str(hora or "08:00"), key=f"clin2_ag_hora_{pid}")
                    dur_p = c7.number_input("Duração", 20, 180, int(dur or 50), step=5, key=f"clin2_ag_dur_{pid}")
                    freq_p = c8.selectbox("Frequência", ["Semanal","Quinzenal","Mensal"], index=["Semanal","Quinzenal","Mensal"].index(freq if freq in ["Semanal","Quinzenal","Mensal"] else "Semanal"), key=f"clin2_ag_freq_{pid}")
                    mod_p = c9.selectbox("Modalidade", ["Presencial","Online","Domiciliar"], index=["Presencial","Online","Domiciliar"].index(mod if mod in ["Presencial","Online","Domiciliar"] else "Presencial"), key=f"clin2_ag_mod_{pid}")
                    if st.button("Salvar agenda do paciente", key=f"clin2_salvar_ag_{pid}"):
                        dia_idx = [k for k,v in DIAS_SEMANA.items() if v == dia_nome][0]
                        novo_prof_id = prof_label_to_id.get(novo_prof_label) or prof_id
                        novo_setor_id = setor_label_to_id.get(novo_setor_label) or setor_id
                        atualizar_agenda_paciente_clinica(USER_ID, pid, dia_idx, time_to_str(hora_p), dur_p, freq_p, mod_p, novo_prof_id, novo_setor_id)
                        st.success("Agenda atualizada."); st.rerun()
        else:
            st.info("Nenhum paciente cadastrado no Modo Clínica ainda.")

    with tab_av:
        setores = listar_setores_clinica(USER_ID)
        if not setores:
            st.warning("Cadastre setores antes de definir avaliações padrão por área.")
        else:
            mapa_setores = {f"{r[1]} — {r[2]}": r for r in setores}
            card_inicio("Avaliações padrão por área/setor", "O gestor/admin define quais avaliações cada área usará. Os modelos aparecem na construção de relatórios.")
            setor_label = st.selectbox("Setor/área que usará esta avaliação", list(mapa_setores.keys()), key="clin2_av_setor")
            setor = mapa_setores[setor_label]
            sid, setor_nome, setor_area = setor[0], setor[1], setor[2]
            c1,c2 = st.columns([0.55,0.45])
            with c1:
                arquivos = st.file_uploader("Enviar avaliação padrão", type=["pdf","docx","xlsx","xls","csv","txt","md","png","jpg","jpeg"], accept_multiple_files=True, key="clin2_av_upload")
            with c2:
                fonte = st.text_input("Fonte/referência do modelo", key="clin2_av_fonte")
                nota = st.text_area("Observação para a equipe", height=110, key="clin2_av_nota")
            previews = []
            if arquivos:
                for i, arq in enumerate(arquivos):
                    texto = texto_arquivo_upload(arq)
                    av_detectada, secoes = montar_modelo_automatizado(arq.name, texto, setor_area)
                    grafico = especificacao_grafico_avaliacao(av_detectada, texto, setor_area)
                    with st.container(border=True):
                        st.markdown(f"#### {arq.name}")
                        c1,c2 = st.columns(2)
                        nome_modelo = c1.text_input("Nome no sistema", value=av_detectada, key=f"clin2_av_nome_{i}")
                        av_conf = c1.text_input("Avaliação reconhecida", value=av_detectada, key=f"clin2_av_rec_{i}")
                        desc = c2.text_area("Justificativa/descrição", value=texto_justificativa_avaliacao(av_detectada, setor_area) + ("\n\n" + nota if nota else ""), height=135, key=f"clin2_av_desc_{i}")
                        st.caption("Domínios do gráfico: " + ", ".join([str(d.get("dominio", d.get("nome", ""))) for d in grafico.get("dominios", [])]))
                        previews.append((arq,nome_modelo,av_conf,desc,secoes,grafico))
                if st.button("Salvar como avaliações padrão deste setor", key="clin2_av_salvar"):
                    for arq,nome_modelo,av_conf,desc,secoes,grafico in previews:
                        safe_name = f"clinica_{USER_ID}_{sid}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{arq.name}".replace("/","_").replace("\\","_")
                        caminho = os.path.join(UPLOAD_DIR, safe_name)
                        with open(caminho,"wb") as f:
                            f.write(arq.getbuffer())
                        salvar_modelo_avaliacao(USER_ID, nome_modelo or av_conf, setor_area, av_conf, desc, secoes, arq.name, caminho, grafico, fonte, gestor_id=USER_ID, setor_id=sid, escopo="clinica", area_alvo=setor_area)
                    st.success("Avaliações padrão salvas para este setor."); st.rerun()
            card_fim()
            avals = listar_avaliacoes_clinica(USER_ID)
            if avals:
                st.markdown("#### Biblioteca padrão da clínica")
                dfav = pd.DataFrame(avals, columns=["ID","Nome","Profissão","Avaliação","Descrição","Criado em","Setor","Área","Fonte"])
                st.dataframe(dfav[["ID","Nome","Avaliação","Setor","Área","Fonte"]], use_container_width=True, hide_index=True)
                for mid,nome,prof,av,desc,criado,setor_nome,area,fonte_ref in avals:
                    with st.container(border=True):
                        c1,c2,c3 = st.columns([0.55,0.25,0.20])
                        c1.markdown(f"**{nome}**  \n{av}  \n<span style='color:#94a3b8'>{(desc or '')[:220]}</span>", unsafe_allow_html=True)
                        c2.markdown(f"Setor: **{setor_nome or '-'}**  \nÁrea: {area or prof}  \nFonte: {fonte_ref or '-'}")
                        if c3.button("Desativar", key=f"clin2_av_del_{mid}"):
                            desativar_avaliacao_clinica(USER_ID, mid); st.rerun()

    with tab_ind:
        st.markdown("### Indicadores e produção clínica")
        hoje = date.today(); inicio_mes = hoje.replace(day=1); fim_mes = (inicio_mes + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        c1,c2,c3 = st.columns(3)
        inicio = c1.date_input("Início", value=inicio_mes, key="clin2_ind_inicio")
        fim = c2.date_input("Fim", value=fim_mes, key="clin2_ind_fim")
        profissionais = get_profissionais_map(USER_ID, ativos=False)
        filtro_opts = ["Todos"] + list(profissionais.keys())
        filtro = c3.selectbox("Profissional", filtro_opts, key="clin2_ind_prof")
        prof_id = None if filtro == "Todos" else profissionais[filtro][0]
        if inicio > fim:
            st.warning("A data inicial não pode ser maior que a final.")
            return
        gerar_eventos_periodo_clinica(USER_ID, inicio, fim)
        cont = eventos_contadores_clinica(USER_ID, inicio, fim, prof_id)
        m1,m2,m3,m4,m5 = st.columns(5)
        m1.metric("Atendimentos previstos", cont["total"])
        m2.metric("Atendimentos feitos", cont["realizados"])
        m3.metric("Evoluções pendentes", cont["pendentes"])
        m4.metric("Cancelamentos/faltas", cont["cancelados"])
        m5.metric("Horas realizadas", cont["horas"])
        st.markdown("#### Produção individual")
        linhas = []
        for label, row in profissionais.items():
            pid = row[0]
            c = eventos_contadores_clinica(USER_ID, inicio, fim, pid)
            linhas.append({"Profissional": row[1], "Área": row[3] or row[7] or "", "Previstos": c["total"], "Feitos": c["realizados"], "Pendentes": c["pendentes"], "Cancelamentos/Faltas": c["cancelados"], "Horas": c["horas"]})
        if linhas:
            st.dataframe(pd.DataFrame(linhas), use_container_width=True, hide_index=True)
        eventos = listar_eventos_periodo_clinica(USER_ID, inicio, fim, prof_id)
        if eventos:
            st.markdown("#### Atendimentos do período")
            dfe = pd.DataFrame(eventos, columns=["ID","Data","Início","Fim","Paciente","Modalidade","Status","Paciente ID","Diagnóstico","Profissional ID","Profissional"])
            st.dataframe(dfe[["Data","Início","Fim","Paciente","Profissional","Modalidade","Status"]], use_container_width=True, hide_index=True)
        else:
            st.info("Nenhum atendimento encontrado no período.")


def tela_agenda_clinica():
    if not modo_clinica_liberado(USER):
        st.info("Modo Clínica disponível para admin master ou plano Clinic/Enterprise.")
        return
    tela_gestao_clinica()


# ==========================================================
# UPGRADE FINAL — WIZARD DE AVALIAÇÕES + INTELIGÊNCIA CLÍNICA
# ==========================================================
FORMATOS_AVALIACAO_APP = [
    "Escala por domínios",
    "Checklist clínico",
    "Formulário descritivo",
    "Escala Likert",
    "Avaliação híbrida",
    "Relatório narrativo",
]
TIPOS_RESPOSTA_AVALIACAO = [
    "Pontuação + observação",
    "Somente pontuação",
    "Somente texto clínico",
    "Sim/Não",
    "Frequência",
]
NIVEIS_FIDELIDADE_AVALIACAO = [
    "Equilibrada e segura",
    "Mais fiel aos enunciados",
    "Mais resumida juridicamente",
    "Mais clínica e interpretativa",
]


def texto_explicativo_avaliacao_segura(nome_avaliacao, profissao, formato, texto_extraido=""):
    """Gera descrição clínica segura sem alegar validação psicométrica/normativa."""
    nome = nome_avaliacao or "Avaliação estruturada personalizada"
    perfil = perfil_profissional(profissao or "Outro")
    texto = str(texto_extraido or "")
    pistas = []
    for termo in ["qualidade de vida", "comunicação", "linguagem", "sensorial", "autonomia", "comportamento", "atenção", "interação", "motor", "emocional", "musical", "família", "rotina"]:
        if termo in texto.lower():
            pistas.append(termo)
    foco_detectado = ", ".join(pistas[:5]) if pistas else perfil.get("foco", "indicadores clínico-funcionais")
    return {
        "o_que_e": (
            f"{nome} foi organizada no NEXO como uma avaliação estruturada de apoio clínico para a área de {profissao}. "
            f"Ela transforma informações do arquivo enviado em domínios, indicadores e campos editáveis, com foco em {foco_detectado}."
        ),
        "para_que_serve": (
            "Serve para padronizar a coleta de dados, registrar observações, atribuir pontuações quando aplicável, acompanhar evolução "
            "e alimentar relatórios clínicos automáticos. O modelo não substitui julgamento profissional, anamnese, observação clínica "
            "nem instrumentos formais quando estes exigirem aplicação específica."
        ),
        "como_foi_criada": (
            f"Formato escolhido: {formato}. O app resumiu e reorganizou a estrutura do arquivo enviado em linguagem clínica editável, "
            "priorizando coerência com os domínios encontrados e evitando reprodução integral de conteúdo protegido."
        ),
        "proteção": (
            "Aviso de proteção: este é um modelo derivado e editável criado a partir de material fornecido pelo próprio profissional. "
            "O NEXO não declara propriedade sobre instrumentos originais, não reproduz manuais/tabelas normativas proprietárias e não garante validade psicométrica. "
            "Antes do uso, revise o conteúdo, confirme autorização/licença quando houver instrumento protegido e valide clinicamente a estrutura final."
        ),
    }


def transformar_secoes_por_formato(secoes, formato, tipo_resposta, nivel_fidelidade):
    """Adapta a avaliação detectada para o formato escolhido pelo usuário."""
    secoes = secoes or []
    novas = []
    for sec in secoes:
        titulo = normalizar_nome_campo(sec.get("titulo", "Seção")) or "Seção"
        if titulo.lower().startswith("identificação"):
            novas.append({"titulo": "Identificação", "campos": ["Data da avaliação", "Contexto de aplicação", "Responsável pelo preenchimento"]})
            continue
        if titulo.lower().startswith("síntese"):
            continue
        campos = campos_unicos(sec.get("campos", []), limite=135)
        if not campos:
            continue
        campos_transformados = []
        for campo in campos:
            base = re.sub(r"^Item\s*\d+\s*[—-]\s*", "", str(campo)).strip()
            if nivel_fidelidade == "Mais resumida juridicamente":
                base = resumir_enunciado_seguro(base, max_palavras=10)
            elif nivel_fidelidade == "Mais clínica e interpretativa":
                base = f"Indicador clínico — {resumir_enunciado_seguro(base, max_palavras=12)}"
            elif nivel_fidelidade == "Mais fiel aos enunciados":
                base = resumir_enunciado_seguro(base, max_palavras=20)
            else:
                base = resumir_enunciado_seguro(base, max_palavras=14)
            campos_transformados.append(base)
        novas.append({"titulo": titulo[:80], "campos": campos_unicos(campos_transformados, limite=135)[:45]})

    if not novas:
        novas = [{"titulo": "Indicadores clínicos", "campos": ["Indicador principal", "Resposta observada", "Nível de suporte", "Conduta sugerida"]}]

    if formato == "Relatório narrativo":
        novas = [
            {"titulo": "Descrição clínica", "campos": ["Queixa ou objetivo da avaliação", "Observações clínicas principais", "Funcionamento atual", "Fatores contextuais"]},
            {"titulo": "Síntese e conduta", "campos": ["Síntese interpretativa", "Prioridades terapêuticas", "Recomendações", "Encaminhamentos quando necessários"]},
        ]
    elif formato == "Formulário descritivo":
        novas.append({"titulo": "Formulação clínica", "campos": ["Pontos fortes", "Necessidades de suporte", "Impacto funcional", "Plano terapêutico sugerido"]})
    elif formato == "Checklist clínico":
        novas.append({"titulo": "Conclusão do checklist", "campos": ["Itens presentes", "Itens ausentes", "Itens parcialmente observados", "Observações complementares"]})
    elif formato == "Escala Likert":
        novas.append({"titulo": "Interpretação da escala", "campos": ["Domínios de maior pontuação", "Domínios de menor pontuação", "Hipótese funcional", "Conduta sugerida"]})
    elif formato == "Avaliação híbrida":
        novas.append({"titulo": "Síntese híbrida", "campos": ["Pontuação global", "Observação clínica integrada", "Prioridades", "Objetivos sugeridos"]})

    return [
        {"titulo": "Identificação", "campos": ["Data da avaliação", "Contexto de aplicação", "Responsável pelo preenchimento"]},
        *[s for s in novas if not str(s.get("titulo", "")).lower().startswith("identificação")],
        {"titulo": "Síntese final", "campos": ["Interpretação clínica", "Impacto funcional", "Prioridades terapêuticas", "Conduta sugerida"]},
    ]


def montar_grafico_por_formato(avaliacao, texto, profissao, formato):
    spec = especificacao_grafico_avaliacao(avaliacao, texto or "", profissao)
    if formato in ["Checklist clínico", "Formulário descritivo", "Relatório narrativo"]:
        spec["tipo"] = "barras_percentual"
    elif formato == "Escala Likert":
        minimo, maximo = detectar_intervalo_pontuacao(texto or "")
        if maximo == 10:
            maximo = 4
        spec["dominios"] = [{**d, "minimo": minimo, "maximo": maximo} for d in spec.get("dominios", [])]
    spec["fonte"] = (
        f"Modelo estruturado no formato '{formato}' a partir de upload profissional. "
        "Domínios e limites devem ser revisados pelo profissional antes do uso clínico."
    )
    return spec


def render_resumo_avaliacao_card(nome, profissao, formato, explicacao, texto_extraido):
    caracteres = len(str(texto_extraido or ""))
    st.markdown(f"""
    <div class="stat-card" style="margin:8px 0 18px 0;">
      <div class="stat-icon">{svg_icon('clipboard',26,'#60a5fa')}</div>
      <div>
        <div class="stat-label">Avaliação detectada</div>
        <div class="stat-number" style="font-size:1.45rem;">{nome}</div>
        <div class="stat-sub">Área: {profissao} • Formato: {formato} • Texto extraído: {caracteres} caracteres</div>
      </div>
    </div>
    """, unsafe_allow_html=True)
    a,b = st.columns(2)
    with a:
        st.markdown("#### O que é")
        st.write(explicacao["o_que_e"])
        st.markdown("#### Para que serve")
        st.write(explicacao["para_que_serve"])
    with b:
        st.markdown("#### Como foi estruturada")
        st.write(explicacao["como_foi_criada"])
        with st.expander("Proteção jurídica e autoral"):
            st.warning(explicacao["proteção"])


def _salvar_bytes_upload_final(uploaded_file, user_id):
    nome = nome_arquivo_seguro(getattr(uploaded_file, "name", "arquivo"))
    try:
        data = uploaded_file.getvalue() or b""
    except Exception:
        data = b""
    stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    caminho = os.path.join(UPLOAD_DIR, f"{user_id}_{stamp}_{nome}".replace("/", "_").replace("\\", "_"))
    with open(caminho, "wb") as f:
        f.write(data)
    return caminho, data


def tela_upload():
    """Wizard premium para transformar qualquer upload em avaliação estruturada configurável."""
    pacientes = listar_pacientes(USER_ID)
    card_inicio("Upload inteligente de avaliações", "Transforme arquivos em avaliações estruturadas com escolha de formato, revisão clínica e proteção autoral.")

    st.markdown("""
    <div class="stats-grid" style="grid-template-columns:repeat(5,1fr);">
      <div class="stat-card"><div class="flow-num">1</div><div><strong>Upload</strong><br><span class="small-muted">Arquivo clínico</span></div></div>
      <div class="stat-card"><div class="flow-num">2</div><div><strong>Leitura</strong><br><span class="small-muted">Texto e domínios</span></div></div>
      <div class="stat-card"><div class="flow-num">3</div><div><strong>Formato</strong><br><span class="small-muted">Como usar no app</span></div></div>
      <div class="stat-card"><div class="flow-num">4</div><div><strong>Revisão</strong><br><span class="small-muted">Validação profissional</span></div></div>
      <div class="stat-card"><div class="flow-num">5</div><div><strong>Relatório</strong><br><span class="small-muted">Avaliação salva</span></div></div>
    </div>
    """, unsafe_allow_html=True)

    c1, c2, c3 = st.columns([0.34, 0.33, 0.33])
    with c1:
        profissao_upload = st.selectbox(
            "Área profissional",
            PROFISSOES,
            index=PROFISSOES.index(USER.get("profissao")) if USER.get("profissao") in PROFISSOES else 0,
            key="wizard_profissao_upload"
        )
    with c2:
        categoria = st.selectbox("Categoria", ["Avaliação", "Modelo de avaliação", "Anamnese", "Escala", "Checklist", "Relatório", "Outro"], key="wizard_categoria_upload")
    with c3:
        paciente_id = None
        if pacientes:
            opts = {f"{p[1]} — {p[4] or 'sem diagnóstico'}": p[0] for p in pacientes}
            paciente_id = opts[st.selectbox("Vincular ao paciente", list(opts.keys()), key="wizard_paciente_upload")]
        else:
            st.caption("Sem paciente vinculado.")

    arquivos = st.file_uploader(
        "Envie PDF, DOCX, XLSX, CSV, TXT ou imagem/documento de apoio",
        type=None,
        accept_multiple_files=True,
        key="wizard_upload_files"
    )
    if not arquivos:
        st.info("Envie um arquivo para iniciar. O sistema vai sugerir uma estrutura, mas o profissional escolhe o formato final.")
        card_fim()
        return

    for idx, arquivo in enumerate(arquivos):
        nome_arq = nome_arquivo_seguro(getattr(arquivo, "name", f"arquivo_{idx+1}"))
        key_base = slug_streamlit_key(hashlib.sha1(f"{USER_ID}_{idx}_{nome_arq}".encode()).hexdigest()[:12])
        with st.container():
            st.markdown(f"### Arquivo {idx+1}: {nome_arq}")
            texto = texto_arquivo_upload(arquivo)
            avaliacao_auto, secoes_auto = montar_modelo_automatizado(nome_arq, texto, profissao_upload)

            f1, f2, f3 = st.columns(3)
            with f1:
                formato = st.selectbox("Como deseja montar no app?", FORMATOS_AVALIACAO_APP, index=4, key=f"fmt_{key_base}")
            with f2:
                tipo_resposta = st.selectbox("Tipo de resposta", TIPOS_RESPOSTA_AVALIACAO, key=f"tipo_resp_{key_base}")
            with f3:
                fidelidade = st.selectbox("Nível de fidelidade", NIVEIS_FIDELIDADE_AVALIACAO, key=f"fid_{key_base}")

            nome_modelo = st.text_input("Nome final da avaliação", value=avaliacao_auto, key=f"nome_modelo_{key_base}")
            secoes_final = transformar_secoes_por_formato(secoes_auto, formato, tipo_resposta, fidelidade)
            grafico_spec = montar_grafico_por_formato(nome_modelo, texto, profissao_upload, formato)
            explicacao = texto_explicativo_avaliacao_segura(nome_modelo, profissao_upload, formato, texto)
            render_resumo_avaliacao_card(nome_modelo, profissao_upload, formato, explicacao, texto)

            tab_prev, tab_dom, tab_graf, tab_salvar = st.tabs(["Prévia prática", "Domínios e itens", "Gráfico", "Salvar"])
            with tab_prev:
                st.markdown("#### Como ficará para preencher")
                for si, sec in enumerate(secoes_final[:6]):
                    st.markdown(f"##### {sec.get('titulo','Seção')}")
                    for ci, campo in enumerate((sec.get("campos") or [])[:8]):
                        cols = st.columns([0.50, 0.20, 0.30]) if tipo_resposta == "Pontuação + observação" else st.columns([0.65,0.35])
                        with cols[0]:
                            st.caption(campo)
                        if tipo_resposta == "Pontuação + observação":
                            with cols[1]: st.caption("Pontuação")
                            with cols[2]: st.caption("Observação")
                        else:
                            with cols[1]: st.caption(tipo_resposta)
                    st.markdown("---")
            with tab_dom:
                st.markdown("#### Revise os domínios e campos antes de salvar")
                secoes_editadas = []
                for si, sec in enumerate(secoes_final):
                    titulo_edit = st.text_input("Título da seção", value=sec.get("titulo","Seção"), key=f"sec_t_{key_base}_{si}")
                    campos_texto = "\n".join(sec.get("campos", []))
                    campos_editados = st.text_area("Campos/itens — um por linha", value=campos_texto, height=150, key=f"sec_c_{key_base}_{si}")
                    secoes_editadas.append({"titulo": titulo_edit, "campos": campos_unicos([x.strip() for x in campos_editados.splitlines() if x.strip()], limite=135)})
            with tab_graf:
                st.markdown("#### Domínios do gráfico")
                dominios = grafico_spec.get("dominios", []) or []
                if dominios:
                    df_g = pd.DataFrame(dominios)
                    df_g = st.data_editor(df_g, use_container_width=True, num_rows="dynamic", key=f"graf_editor_{key_base}")
                    grafico_spec = normalizar_grafico_spec(df_g.rename(columns={"dominio":"Domínio", "maximo":"Máximo", "minimo":"Mínimo"}), tipo=grafico_spec.get("tipo","barras_percentual"), fonte=grafico_spec.get("fonte",""))
                else:
                    st.warning("Nenhum domínio gráfico foi detectado. Você poderá usar texto clínico sem gráfico ou adicionar domínios manualmente.")
            with tab_salvar:
                descricao = st.text_area(
                    "Descrição segura da avaliação",
                    value="\n\n".join([explicacao["o_que_e"], explicacao["para_que_serve"], explicacao["como_foi_criada"], explicacao["proteção"]]),
                    height=210,
                    key=f"desc_segura_{key_base}"
                )
                if st.button("Salvar como avaliação estruturada", key=f"salvar_modelo_wizard_{key_base}"):
                    caminho, data = _salvar_bytes_upload_final(arquivo, USER_ID)
                    registrar_upload(USER_ID, paciente_id, nome_arq, caminho, extensao_arquivo(nome_arq), categoria, nome_modelo, descricao[:900])
                    salvar_modelo_avaliacao(
                        USER_ID,
                        nome_modelo,
                        profissao_upload,
                        nome_modelo,
                        descricao,
                        secoes_editadas if 'secoes_editadas' in locals() else secoes_final,
                        arquivo_origem=nome_arq,
                        caminho_origem=caminho,
                        grafico_spec=grafico_spec,
                        fonte_referencia=explicacao["proteção"],
                        escopo="perfil",
                        area_alvo=profissao_upload,
                    )
                    log_uso(USER_ID, "upload")
                    st.success("Avaliação estruturada salva. Ela já aparece em Avaliação e relatório > Avaliações automatizadas do Upload.")
                    st.rerun()
    card_fim()




# ==========================================================
# GRÁFICOS AUTOMÁTICOS PARA AVALIAÇÕES CRIADAS POR UPLOAD
# ==========================================================
def eh_campo_pontuavel(campo, titulo_secao=""):
    """Define se um campo deve receber pontuação automática para compor gráficos."""
    txt = f"{titulo_secao} {campo}".lower()
    bloqueios = [
        "data da avaliação", "data da avaliacao", "contexto de aplicação", "contexto de aplicacao",
        "responsável", "responsavel", "identificação", "identificacao", "síntese", "sintese",
        "interpretação", "interpretacao", "impacto funcional", "prioridades terapêuticas", "prioridades terapeuticas",
        "conduta", "responsável pelo preenchimento", "responsavel pelo preenchimento"
    ]
    if any(b in txt for b in bloqueios):
        return False
    gatilhos = [
        "item", "indicador", "domínio", "dominio", "habilidade", "comunica", "aten", "intera",
        "regula", "motor", "sensor", "linguagem", "autonomia", "comport", "participa", "responde",
        "inicia", "mantém", "mantem", "realiza", "explora", "percep", "express"
    ]
    return any(g in txt for g in gatilhos) or len(txt.split()) >= 3


def _tokens_clinicos(txt):
    txt = re.sub(r"[^a-zA-ZÀ-ÿ0-9 ]+", " ", str(txt or "").lower())
    stop = {"item", "indicador", "de", "da", "do", "das", "dos", "em", "a", "o", "e", "ou", "para", "com", "por", "um", "uma"}
    return [t for t in txt.split() if len(t) >= 4 and t not in stop]


def calcular_grafico_automatico(secoes, pontuacoes_campos, grafico_spec):
    """Calcula valores de gráfico sem exigir preenchimento manual separado.

    Prioridade:
    1) associa domínio ao título da seção;
    2) associa domínio a itens/campos com palavras semelhantes;
    3) usa a média geral dos itens pontuados quando não há associação clara.
    """
    dominios = (grafico_spec or {}).get("dominios", []) or []
    if not dominios or not pontuacoes_campos:
        return {}

    mapa_secao = []
    for sec in secoes or []:
        titulo = normalizar_nome_campo(sec.get("titulo", ""))
        campos = [normalizar_nome_campo(c) for c in (sec.get("campos") or [])]
        mapa_secao.append((titulo, campos))

    valores = {}
    todos_scores = [float(v) for v in pontuacoes_campos.values() if isinstance(v, (int, float))]
    media_geral = sum(todos_scores) / len(todos_scores) if todos_scores else 0.0

    for i, dominio in enumerate(dominios):
        nome_dom = normalizar_nome_campo(dominio.get("dominio", f"Domínio {i+1}")) or f"Domínio {i+1}"
        min_dom = float(dominio.get("minimo", 0) or 0)
        max_dom = float(dominio.get("maximo", 10) or 10)
        tokens_dom = set(_tokens_clinicos(nome_dom))
        candidatos = []

        for titulo, campos in mapa_secao:
            tokens_sec = set(_tokens_clinicos(titulo))
            sec_match = bool(tokens_dom and tokens_dom.intersection(tokens_sec)) or nome_dom.lower() in titulo.lower() or titulo.lower() in nome_dom.lower()
            for campo in campos:
                score = pontuacoes_campos.get(campo)
                if score is None:
                    continue
                tokens_campo = set(_tokens_clinicos(campo))
                campo_match = bool(tokens_dom and tokens_dom.intersection(tokens_campo)) or nome_dom.lower() in campo.lower()
                if sec_match or campo_match:
                    candidatos.append(float(score))

        if not candidatos:
            candidatos = todos_scores

        valor = sum(candidatos) / len(candidatos) if candidatos else media_geral
        valor = max(min_dom, min(max_dom, valor))
        valores[nome_dom] = round(valor, 2)
    return valores


def widget_pontuacao_automatica(campo, sec_titulo, minimo, maximo, key):
    """Widget numérico discreto usado para gerar gráfico automaticamente."""
    if maximo <= minimo:
        maximo = minimo + 10
    default = minimo
    return st.slider(
        "Pontuação",
        min_value=float(minimo),
        max_value=float(maximo),
        value=float(default),
        step=1.0,
        key=key,
        help="Essa pontuação alimenta automaticamente o gráfico do relatório."
    )

def gerar_objetivos_terapeuticos(dados_base, profissao):
    perfil = perfil_profissional(profissao or "Outro")
    objetivos = []
    for area in perfil.get("areas", [])[:5]:
        objetivos.append({
            "Domínio": area.capitalize(),
            "Objetivo SMART": f"Ampliar {area} em contextos funcionais, com critérios observáveis e registro evolutivo periódico.",
            "Meta GAS -2": "Resposta ausente ou dependente de suporte máximo.",
            "Meta GAS 0": "Resposta funcional esperada com suporte moderado e consistência clínica.",
            "Meta GAS +2": "Resposta espontânea, generalizada e sustentada em diferentes contextos.",
        })
    return objetivos


def biblioteca_intervencoes_por_profissao(profissao):
    base = {
        "Musicoterapeuta": ["Improvisação clínica responsiva", "Turn taking musical", "Jogos rítmicos estruturados", "Canções de rotina", "Escuta ativa com regulação"],
        "Fonoaudiólogo": ["Modelagem de comunicação funcional", "Rotinas de troca comunicativa", "Expansão de fala", "Pragmática em contexto natural", "Treino de compreensão"],
        "Terapeuta Ocupacional": ["Circuitos sensório-motores", "Treino de AVDs", "Atividades de modulação sensorial", "Coordenação bimanual", "Planejamento motor funcional"],
        "Psicólogo": ["Psicoeducação emocional", "Treino de autorregulação", "Análise funcional", "Habilidades sociais", "Estratégias de enfrentamento"],
        "Aplicador ABA": ["Ensino por tentativas discretas", "NET", "Reforçamento diferencial", "Prompt fading", "Generalização programada"],
    }
    return base.get(profissao, ["Observação funcional", "Intervenção estruturada", "Mediação terapêutica", "Registro evolutivo", "Revisão de objetivos"])


def tela_central_clinica_inteligente():
    card_inicio("Central clínica inteligente", "Ferramentas para evolução comparativa, objetivos terapêuticos, biblioteca de intervenções, dashboard e segurança documental.")
    pacientes = listar_pacientes(USER_ID)
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Evolução comparativa", "Objetivos terapêuticos", "Biblioteca de intervenções", "Dashboard clínico", "Segurança"])

    with tab1:
        st.markdown("### Evolução clínica inteligente")
        if not pacientes:
            st.info("Cadastre pacientes para comparar evolução.")
        else:
            opts = {f"{p[1]} — {p[4] or 'sem diagnóstico'}": p for p in pacientes}
            psel = opts[st.selectbox("Paciente", list(opts.keys()), key="central_paciente_comp")]
            evols = listar_evolucoes(USER_ID, psel[0])
            respostas = listar_respostas_relatorio(USER_ID, psel[0])
            c1,c2,c3 = st.columns(3)
            c1.metric("Evoluções registradas", len(evols))
            c2.metric("Avaliações no relatório", len(respostas))
            c3.metric("Status", "Em acompanhamento" if evols or respostas else "Sem dados")
            if evols:
                st.markdown("#### Síntese automática")
                ultimas = evols[:4]
                media_eng = sum([int(e[4] or 0) for e in ultimas]) / max(1, len(ultimas))
                media_prog = sum([int(e[5] or 0) for e in ultimas]) / max(1, len(ultimas))
                st.write(f"Nas últimas evoluções, o paciente apresenta média de engajamento {media_eng:.1f}/10 e progresso {media_prog:.1f}/10. Recomenda-se comparar esses dados com os domínios das avaliações estruturadas e revisar objetivos terapêuticos a cada ciclo de reavaliação.")
                st.dataframe(pd.DataFrame(evols, columns=["Data","Título","Descrição","Estado","Engajamento","Progresso"]), use_container_width=True, hide_index=True)
            else:
                st.info("Ainda não há evoluções suficientes para análise comparativa.")

    with tab2:
        st.markdown("### Gerador de objetivos terapêuticos")
        profissao = st.selectbox("Área", PROFISSOES, index=PROFISSOES.index(USER.get("profissao")) if USER.get("profissao") in PROFISSOES else 0, key="central_obj_prof")
        objetivos = gerar_objetivos_terapeuticos({}, profissao)
        df_obj = pd.DataFrame(objetivos)
        st.dataframe(df_obj, use_container_width=True, hide_index=True)
        st.download_button("Baixar objetivos em CSV", df_obj.to_csv(index=False).encode("utf-8"), file_name="objetivos_terapeuticos_nexo.csv", mime="text/csv")

    with tab3:
        st.markdown("### Biblioteca de intervenções")
        profissao_b = st.selectbox("Profissão", PROFISSOES, index=PROFISSOES.index(USER.get("profissao")) if USER.get("profissao") in PROFISSOES else 0, key="central_bib_prof")
        dominio = st.text_input("Domínio ou necessidade clínica", value="comunicação funcional", key="central_bib_dom")
        for item in biblioteca_intervencoes_por_profissao(profissao_b):
            st.markdown(f"- **{item}** — pode ser adaptada para {dominio}, com registro de resposta, nível de suporte e generalização.")

    with tab4:
        st.markdown("### Dashboard clínico")
        n_pac = contar_pacientes(USER_ID)
        uploads = uso_mes(USER_ID, "upload")
        rels = uso_mes(USER_ID, "relatorio")
        c1,c2,c3,c4 = st.columns(4)
        c1.metric("Pacientes", n_pac)
        c2.metric("Uploads no mês", uploads)
        c3.metric("Relatórios no mês", rels)
        c4.metric("Plano", USER.get("plano") or "sem plano")
        st.info("Este painel pode evoluir para BI clínico com faltas, horas, evolução por domínio, produtividade por profissional e alertas de reavaliação.")

    with tab5:
        st.markdown("### Segurança, versão e rastreabilidade")
        st.write("Recursos adicionados conceitualmente ao fluxo: modelos derivados, validação profissional, aviso de proteção autoral e histórico por upload.")
        texto_hash = f"{USER_ID}-{datetime.now().isoformat()}-{USER.get('email','')}"
        h = hashlib.sha256(texto_hash.encode()).hexdigest()
        st.code(f"Hash de exemplo para rastreabilidade documental: {h}")
        st.warning("Para produção comercial, o próximo passo é implementar histórico de versões, log imutável e assinatura digital dos relatórios.")
    card_fim()


# ==========================================================
# LAYOUT FINAL SEGURO — SIDEBAR PRÓPRIA REAL, SEM st.sidebar
# ==========================================================
st.markdown("""
<style>
/* Remove totalmente a sidebar nativa e o botão de colapsar do Streamlit.
   A navegação do app passa a ser 100% própria dentro da primeira coluna. */
[data-testid="stSidebar"],
[data-testid="collapsedControl"],
section[data-testid="stSidebar"],
button[title="View fullscreen"],
button[kind="header"] {
    display: none !important;
    visibility: hidden !important;
    width: 0 !important;
    min-width: 0 !important;
}

[data-testid="stAppViewContainer"] > .main {
    margin-left: 0 !important;
}

.block-container {
    padding-top: 1.15rem !important;
    padding-left: 1.65rem !important;
    padding-right: 1.65rem !important;
    max-width: 1520px !important;
}

/* Correção definitiva do bug arrow / keyboard_arrow_down */
div[data-testid="stExpander"] summary span[class*="material-symbols"],
div[data-testid="stExpander"] summary span[class*="material-icons"],
div[data-testid="stExpander"] summary i[class*="material-symbols"],
div[data-testid="stExpander"] summary i[class*="material-icons"] {
    color: transparent !important;
    font-size: 0 !important;
    line-height: 0 !important;
    max-width: 0 !important;
    width: 0 !important;
    overflow: hidden !important;
}
div[data-testid="stExpander"] summary::before {
    content: "▾";
    display: inline-flex;
    width: 22px;
    height: 22px;
    align-items: center;
    justify-content: center;
    margin-right: 8px;
    color: #9fb7d8;
    font-size: 16px;
}

/* Sidebar própria aplicada na coluna que contém o sentinela.
   Não usamos mais um <div> aberto envolvendo widgets, porque isso gerava
   a caixa preta vazia e empurrava os botões para baixo. */
[data-testid="column"]:has(.nexo-sidebar-sentinel) {
    min-height: calc(100vh - 38px);
    padding: 22px 14px 18px 14px !important;
    border: 1px solid rgba(148,163,184,.16);
    border-radius: 22px;
    background: linear-gradient(180deg,#040a15,#07101d 56%,#060b17);
    box-shadow: 0 24px 60px rgba(2,6,23,.32);
    position: sticky;
    top: 14px;
    align-self: flex-start;
}

.sidebar-brand { padding: 0 4px 18px !important; }
.sidebar-brand-row { display:flex; align-items:center; gap:12px; }
.sidebar-brand h1 { margin:0 !important; font-size:1.8rem !important; font-weight:950 !important; letter-spacing:-.08em; }
.sidebar-brand p { margin:0 !important; color:#aab7cf !important; font-size:.66rem !important; letter-spacing:.06em; text-transform:uppercase; }
.sidebar-section { color:#7587a6; font-size:.72rem; letter-spacing:.08em; text-transform:uppercase; margin:18px 0 8px; padding-left:8px; }
.sidebar-actions-note { color:#7587a6; font-size:.72rem; letter-spacing:.08em; text-transform:uppercase; margin:20px 0 8px; padding-left:8px; }
.sidebar-user-card {
    margin-top: 18px;
    border: 1px solid rgba(148,163,184,.14);
    border-radius: 16px;
    padding: 14px;
    background: linear-gradient(145deg,rgba(15,23,42,.82),rgba(30,41,59,.38));
}

/* Botões da lateral: simétricos, mesmo tamanho e sem quebrar texto */
[data-testid="column"]:has(.nexo-sidebar-sentinel) .stButton > button {
    min-height: 48px !important;
    height: auto !important;
    margin: 3px 0 !important;
    justify-content: flex-start !important;
    text-align: left !important;
    padding: 0 15px !important;
    border-radius: 14px !important;
    background: linear-gradient(90deg,#7c3aed,#0ea5e9) !important;
    border: 1px solid rgba(255,255,255,.13) !important;
    box-shadow: 0 10px 24px rgba(14,165,233,.12) !important;
    font-weight: 800 !important;
    white-space: normal !important;
}
[data-testid="column"]:has(.nexo-sidebar-sentinel) .stButton > button p {
    font-size: .93rem !important;
    line-height: 1.15 !important;
}
.nexo-nav-active-marker + div button {
    background: linear-gradient(90deg,#4f46e5,#7c3aed) !important;
    border: 1px solid rgba(255,255,255,.18) !important;
    box-shadow: 0 14px 36px rgba(79,70,229,.34) !important;
}

/* Mantém a área principal fluida */
.topbar { margin-top: 0 !important; }
.layout-grid { align-items: stretch; }
.hero { min-height: 300px !important; }
.stats-grid { margin-top: 18px !important; }

@media(max-width:1200px){
    [data-testid="column"]:has(.nexo-sidebar-sentinel) {
        position: relative;
        top: 0;
        min-height: auto;
        margin-bottom: 18px;
    }
}

/* Ajustes startup NEXO: logo não sobrepõe texto, busca funcional e eliminação visual do bug _arrow */
.hero { position: relative !important; overflow: hidden !important; }
.hero-copy { position: relative !important; z-index: 2 !important; max-width: 62% !important; padding-right: 18px !important; }
.hero-brain { right: 34px !important; top: 50% !important; transform: translateY(-50%) !important; width: 190px !important; height: 190px !important; opacity: .58 !important; z-index: 1 !important; pointer-events: none !important; }
.hero-brain svg { max-width: 132px !important; max-height: 132px !important; }
.search-results-card { margin: -8px 0 18px 0; padding: 14px; border-radius: 16px; border: 1px solid rgba(148,163,184,.14); background: linear-gradient(145deg,rgba(13,22,39,.94),rgba(17,25,42,.72)); }
div[data-testid="stTextInput"] input { min-height: 48px !important; }
/* Oculta qualquer resíduo de ícone textual do Streamlit em headers/expanders */
span, i { text-decoration: none !important; }
@media(max-width:1200px){ .hero-copy{max-width:100% !important;} .hero-brain{display:none !important;} }

/* Accordion premium próprio — sem st.expander, sem Material Icons, sem _arrow_down */
.nexo-accordion-head {
    display:flex;
    align-items:center;
    gap:14px;
    margin:18px 0 10px 0;
    padding:16px 18px;
    border-radius:18px;
    border:1px solid rgba(148,163,184,.18);
    background:linear-gradient(145deg,rgba(13,22,39,.96),rgba(17,25,42,.70));
    box-shadow:0 14px 38px rgba(2,6,23,.24);
}
.nexo-accordion-icon {
    width:32px;
    height:32px;
    border-radius:11px;
    display:flex;
    align-items:center;
    justify-content:center;
    background:linear-gradient(135deg,#7c3aed,#0ea5e9);
    color:#fff;
    font-weight:900;
    font-size:22px;
    line-height:1;
    flex:0 0 32px;
}
.nexo-accordion-head strong {
    color:#f8fafc;
    font-size:1.02rem;
    letter-spacing:-.02em;
}
.nexo-accordion-head span {
    color:#9fb0cf;
    font-size:.88rem;
}
.nexo-accordion-body {
    margin:10px 0 18px 0;
    padding:18px;
    border-radius:18px;
    border:1px solid rgba(148,163,184,.14);
    background:rgba(15,23,42,.46);
}
</style>
""", unsafe_allow_html=True)

# Menu lateral principal + ações rápidas. O admin aparece somente para administradores.
opcoes = [
    "Dashboard",
    "Perfil profissional",
    "Pacientes e evolução",
    "Upload inteligente",
    "Agenda e calendário",
    "Avaliação e relatório",
    "Central clínica inteligente",
]

# Modo Clínica aparece para o admin master e para usuários dos planos Clinic/Enterprise.
if modo_clinica_liberado(USER):
    opcoes += ["Modo Clínica"]

opcoes += ["Planos e acesso"]

if str(USER.get("role") or "").lower() == "admin" or str(USER.get("email") or "").lower() == ADMIN_EMAIL.lower():
    opcoes += ["Admin — Controle de Acessos"]

alias_paginas = {
    "Pacientes": "Pacientes e evolução",
    "Evolução do paciente": "Pacientes e evolução",
    "Avaliações": "Avaliação e relatório",
    "Relatórios": "Avaliação e relatório",
    "Perfil e logo": "Perfil profissional",
    "Usuários": "Admin — Controle de Acessos",
    "Assinatura e cobrança": "Admin — Controle de Acessos",
    "Logs do sistema": "Admin — Controle de Acessos",
    "Configurações": "Admin — Controle de Acessos",
}
if st.session_state.page in alias_paginas:
    st.session_state.page = alias_paginas[st.session_state.page]
if st.session_state.page not in opcoes:
    st.session_state.page = "Dashboard"

left_col, right_col = st.columns([0.165, 0.835], gap="medium")

with left_col:
    st.markdown('<span class="nexo-sidebar-sentinel"></span>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="sidebar-brand">
      <div class="sidebar-brand-row">{brain_logo(54)}<div><h1>NEXO</h1><p>{APP_SLOGAN}</p></div></div>
    </div>
    <div class="sidebar-section">Principal</div>
    """, unsafe_allow_html=True)

    icon_map = {
        "Dashboard": "🏠",
        "Perfil profissional": "👤",
        "Pacientes e evolução": "🧑‍⚕️",
        "Upload inteligente": "📁",
        "Agenda e calendário": "📅",
        "Avaliação e relatório": "🧾",
        "Central clínica inteligente": "🧠",
        "Planos e acesso": "💳",
        "Modo Clínica": "🏥",
        "Admin — Controle de Acessos": "👑",
    }

    for item in opcoes:
        if item == "Admin — Controle de Acessos":
            st.markdown('<div class="sidebar-section">Administração</div>', unsafe_allow_html=True)
        if item == st.session_state.page:
            st.markdown('<div class="nexo-nav-active-marker"></div>', unsafe_allow_html=True)
        if st.button(f"{icon_map.get(item, '•')}  {item}", key=f"nav_fixed_{item}", use_container_width=True):
            go_to(item)

    st.markdown('<div class="sidebar-actions-note">Ações rápidas</div>', unsafe_allow_html=True)
    if st.button("＋  Cadastrar paciente", key="side_action_paciente", use_container_width=True):
        go_to("Pacientes e evolução")
    if st.button("📊  Nova avaliação", key="side_action_avaliacao", use_container_width=True):
        go_to("Avaliação e relatório")
    if st.button("📁  Upload inteligente", key="side_action_upload", use_container_width=True):
        go_to("Upload inteligente")
    if st.button("📅  Agenda", key="side_action_agenda", use_container_width=True):
        go_to("Agenda e calendário")
    if st.button("🧾  Gerar relatório", key="side_action_relatorio", use_container_width=True):
        go_to("Avaliação e relatório")
    if st.button("🧠  Central clínica", key="side_action_central_clinica", use_container_width=True):
        go_to("Central clínica inteligente")
    if modo_clinica_liberado(USER):
        if st.button("🏥  Modo Clínica", key="side_action_modo_clinica", use_container_width=True):
            go_to("Modo Clínica")
    if st.button("💳  Planos", key="side_action_planos", use_container_width=True):
        go_to("Planos e acesso")

    st.markdown(f"""
    <div class="sidebar-user-card">
      <div class="user-line"><div class="avatar">{initials(USER.get('nome'))}</div>
      <div><strong>{USER.get('nome')}</strong><br><span style="color:#9fb0cf;font-size:.84rem;">{USER.get('profissao') or 'Profissional'}</span><br><span style="color:#34d399;font-size:.78rem;">Online</span></div></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="nexo-plan-card">
      <div style="display:flex;align-items:center;gap:8px;font-weight:900;color:#f8fafc;">👑 Plano {USER.get('plano') or 'não definido'}</div>
      <div class="nexo-muted" style="font-size:.78rem;margin-top:5px;">Válido até {USER.get('data_expiracao') or 'sem data'}</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("↪  Sair", key="logout_fixed", use_container_width=True):
        st.session_state.clear()
        st.rerun()

with right_col:
    topbar(USER)
    render_back_button()
    startup_header(USER)

    menu = st.session_state.page
    if menu == "Dashboard":
        tela_dashboard()
    elif menu == "Perfil profissional":
        tela_perfil()
    elif menu == "Pacientes e evolução":
        tela_pacientes_e_evolucao()
    elif menu == "Upload inteligente":
        tela_upload()
    elif menu == "Agenda e calendário":
        tela_agenda_terapeuta()
    elif menu == "Avaliação e relatório":
        tela_avaliacoes_relatorios()
    elif menu == "Central clínica inteligente":
        tela_central_clinica_inteligente()
    elif menu == "Modo Clínica":
        tela_agenda_clinica()
    elif menu == "Planos e acesso":
        tela_planos()
    elif menu == "Admin — Controle de Acessos":
        tela_admin()
